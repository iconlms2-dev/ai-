"""마케팅 자동화 대시보드 — 백엔드 서버"""
import asyncio, json, re, time, os, io, csv, hmac, hashlib, base64, uuid, shutil, subprocess, threading, random
from datetime import datetime, timedelta
from urllib.parse import quote, quote_plus
from concurrent.futures import ThreadPoolExecutor
from PIL import Image, ImageDraw, ImageFont

from fastapi import FastAPI, Request, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse, Response
import requests as req
from bs4 import BeautifulSoup

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

import openpyxl

# ───────────────────────────── CONFIG ─────────────────────────────
from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'), override=True)

NAVER_AD_API_KEY  = os.environ.get('NAVER_AD_API_KEY', '')
NAVER_AD_SECRET   = os.environ.get('NAVER_AD_SECRET', '')
NAVER_AD_CUSTOMER = os.environ.get('NAVER_AD_CUSTOMER', '')
NOTION_TOKEN      = os.environ.get('NOTION_TOKEN', '')
KEYWORD_DB_ID     = os.environ.get('KEYWORD_DB_ID', '')
PROGRESS_FILE     = os.path.join(os.path.dirname(__file__), "keyword_progress.json")
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
GEMINI_API_KEY    = os.environ.get('GEMINI_API_KEY', '')
CONTENT_DB_ID     = os.environ.get('CONTENT_DB_ID', '')
CAFE24_CLIENT_ID  = os.environ.get('CAFE24_CLIENT_ID', '')
CAFE24_CLIENT_SECRET = os.environ.get('CAFE24_CLIENT_SECRET', '')
CAFE24_MALL_ID    = os.environ.get('CAFE24_MALL_ID', '')
CAFE24_TOKEN_FILE = os.path.join(os.path.dirname(__file__), "cafe24_token.json")
ELEVENLABS_API_KEY = os.environ.get('ELEVENLABS_API_KEY', '')
SHORTS_DIR = os.path.join(os.path.dirname(__file__), "shorts_output")
os.makedirs(SHORTS_DIR, exist_ok=True)
THREADS_APP_ID     = os.environ.get('THREADS_APP_ID', '')
THREADS_APP_SECRET = os.environ.get('THREADS_APP_SECRET', '')
THREADS_ACCOUNTS_FILE = os.path.join(os.path.dirname(__file__), "threads_accounts.json")
THREADS_QUEUE_FILE = os.path.join(os.path.dirname(__file__), "threads_queue.json")
REDIRECT_BASE_URL = os.environ.get('REDIRECT_BASE_URL', 'http://localhost:8000')

SECTION_MAP = {
    'pwl_nop':'파워링크','shp_gui':'쇼핑','shp_dui':'네이버가격비교',
    'shs_lis':'네이버플러스스토어','urB_coR':'신뢰도통합','urB_imM':'이미지',
    'urB_boR':'VIEW/블로그','ugB_adR':'브랜드콘텐츠','ugB_pkR':'브랜드콘텐츠',
    'ugB_bsR':'인기글','ugB_b1R':'신뢰도통합','ugB_b2R':'신뢰도통합',
    'ugB_b3R':'신뢰도통합','ugB_ipR':'인플루언서','ugB_qpR':'기타',
    'heL_htX':'AI브리핑','heB_ceR':'관련경험카페글','nws_all':'뉴스',
    'web_gen':'웹사이트','kwX_ndT':'함께많이찾는','exB_soT':'함께보면좋은',
    'kwL_ssT':'연관검색어','ldc_btm':'지식백과','bok_lst':'도서',
    'nmb_hpl':'플레이스','sit_4po':'웹사이트내검색','brd_brd':'브랜드서치',
    'abL_baX':'AI브리핑','abL_rtX':'AI브리핑','rrB_hdR':'리랭킹',
    'rrB_bdR':'리랭킹','nco_x58':'기타','ink_mik':'기타','nmb_rnk':'기타','ink_kid':'기타',
}
CONTENT_CODES = {'urB_coR','urB_boR','ugB_b1R','ugB_b2R','ugB_b3R','heB_ceR','ink_kid','ugB_bsR'}
NOISE_WORDS = ['더보기','클릭','전체보기','닫기','FAQ','인플루언서 참여','콘텐츠더보기','바로가기','자세히보기','관련검색어']

executor = ThreadPoolExecutor(max_workers=3)
_selenium_semaphore = asyncio.Semaphore(1)   # 동시 Selenium driver 1개로 제한
app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["http://localhost:8000"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

# ───────────────────────────── HELPERS ─────────────────────────────
def _error_response(message, status=500, details=None):
    """표준 에러 응답"""
    body = {'ok': False, 'error': message}
    if details: body['details'] = str(details)
    return JSONResponse(body, status_code=status)

def _valid_kw(text):
    text = text.strip()
    if not text or len(text) < 2 or len(text) > 40:
        return False
    if not re.search(r'[가-힣]', text):
        return False
    for n in NOISE_WORDS:
        if n in text:
            return False
    return True

def _create_driver():
    opts = Options()
    opts.add_argument('--headless=new')
    opts.add_argument('--no-sandbox')
    opts.add_argument('--disable-dev-shm-usage')
    opts.add_argument('--disable-gpu')
    opts.add_argument('--window-size=1920,1080')
    opts.add_argument('user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36')
    svc = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=svc, options=opts)

# ── 자동완성 ──
def _autocomplete(keyword):
    url = f"https://ac.search.naver.com/nx/ac?q={quote(keyword)}&con=1&frm=nv&ans=2&t_koreng=1&q_enc=UTF-8&st=100&r_format=json&r_enc=UTF-8&r_unicode=0&type=1"
    try:
        r = req.get(url, timeout=5)
        data = r.json()
        items = data.get('items', [[]])
        results = []
        for group in items:
            for item in group:
                if isinstance(item, list) and len(item) > 0:
                    results.append(item[0])
                elif isinstance(item, str):
                    results.append(item)
        return results
    except Exception:
        return []

# ── Selenium 키워드 확장 ──
def _expand_selenium(driver, keyword):
    """네이버 검색 페이지에서 연관검색어/함께많이찾는/함께보면좋은 영역만 정확히 파싱"""
    kws = {}
    try:
        url = f"https://search.naver.com/search.naver?query={quote(keyword)}&where=nexearch"
        driver.get(url)
        time.sleep(2.5)
        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')

        # ── 방법 1: 섹션 제목(h2) 기반 타겟팅 ──
        for section in soup.select('.api_subject_bx'):
            title_el = section.select_one('h2, .api_title, .tit')
            if not title_el:
                continue
            title_text = title_el.get_text(strip=True)

            source = None
            if '연관검색어' in title_text:
                source = '연관검색어'
            elif '함께 많이 찾는' in title_text or '함께많이찾는' in title_text:
                source = '함께많이찾는'
            elif '함께 보면 좋은' in title_text or '함께보면좋은' in title_text:
                source = '함께보면좋은'

            if source:
                for a in section.find_all('a'):
                    span = a.select_one('span')
                    t = span.get_text(strip=True) if span else a.get_text(strip=True)
                    if _valid_kw(t) and t not in kws:
                        kws[t] = source

        # ── 방법 2: URL 파라미터 기반 (fallback) ──
        if not kws:
            for a in soup.find_all('a', href=True):
                href = a.get('href', '')
                source = None
                if 'sm=tab_rel' in href:
                    source = '연관검색어'
                elif 'sm=tab_clk.ndT' in href:
                    source = '함께많이찾는'
                elif 'sm=tab_clk.ssT' in href:
                    source = '함께보면좋은'

                if source:
                    span = a.select_one('span')
                    t = span.get_text(strip=True) if span else a.get_text(strip=True)
                    if _valid_kw(t) and t not in kws:
                        kws[t] = source

    except Exception as e:
        print(f"Selenium error for '{keyword}': {e}")
    return kws

# ── 검색광고 API ──
def _ad_signature(timestamp):
    msg = f"{timestamp}.GET./keywordstool"
    sig = hmac.new(NAVER_AD_SECRET.encode(), msg.encode(), hashlib.sha256).digest()
    return base64.b64encode(sig).decode()

def _search_volume(keywords_list):
    """여러 키워드를 배치로 조회. dict 반환 {keyword: {pc, mo}}"""
    result = {}
    batch_size = 5
    for i in range(0, len(keywords_list), batch_size):
        batch = keywords_list[i:i+batch_size]
        hint = ','.join(kw.replace(' ','') for kw in batch)
        ts = str(int(time.time() * 1000))
        headers = {
            'Content-Type': 'application/json',
            'X-Timestamp': ts,
            'X-API-KEY': NAVER_AD_API_KEY,
            'X-Customer': NAVER_AD_CUSTOMER,
            'X-Signature': _ad_signature(ts),
        }
        try:
            r = req.get(f"https://api.searchad.naver.com/keywordstool?hintKeywords={quote_plus(hint)}&showDetail=1",
                       headers=headers, timeout=10)
            if r.status_code == 200:
                data = r.json()
                for item in data.get('keywordList', []):
                    rk = item.get('relKeyword','').replace(' ','').upper()
                    pc = item.get('monthlyPcQcCnt', 0)
                    mo = item.get('monthlyMobileQcCnt', 0)
                    if isinstance(pc, str): pc = 10 if '<' in pc else int(pc.replace(',','') or '0')
                    if isinstance(mo, str): mo = 10 if '<' in mo else int(mo.replace(',','') or '0')
                    # 배치 키워드들과 매칭
                    for kw in batch:
                        if kw.replace(' ','').upper() == rk:
                            result[kw] = {'pc': pc, 'mo': mo}
                            break
            else:
                print(f"Search Ad API error {r.status_code}: {r.text[:200]}")
        except Exception as e:
            print(f"Search volume error: {e}")
        time.sleep(0.3)

    # 조회 안 된 키워드는 0으로
    for kw in keywords_list:
        if kw not in result:
            result[kw] = {'pc': 0, 'mo': 0}
    return result

# ── SERP 분석 ──
def _parse_date(text, today):
    if not text: return None
    m = re.search(r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})', text)
    if m: return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    m = re.search(r'(\d+)일\s*전', text)
    if m: return today - timedelta(days=int(m.group(1)))
    m = re.search(r'(\d+)주\s*전', text)
    if m: return today - timedelta(weeks=int(m.group(1)))
    m = re.search(r'(\d+)개월\s*전', text)
    if m: return today - timedelta(days=int(m.group(1))*30)
    if '어제' in text: return today - timedelta(days=1)
    return None

def _analyze_serp(keyword, today):
    """SERP HTML을 requests로 가져와서 분석"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    result = {'content_tab_rank': '-', 'content_format': '-', 'top6_tabs': '-', 'articles': []}
    try:
        r = req.get(f"https://search.naver.com/search.naver?query={quote(keyword)}&where=nexearch", headers=headers, timeout=10)
        html = r.text
        soup = BeautifulSoup(html, 'html.parser')

        # nx_cr_area_info 파싱
        m = re.search(r'var\s+nx_cr_area_info\s*=\s*(\[.*?\])\s*;', html, re.DOTALL)
        areas = []
        if m:
            try: areas = json.loads(m.group(1))
            except (json.JSONDecodeError, ValueError): pass

        sorted_areas = sorted(areas, key=lambda x: x.get('r', 999))

        # 콘텐츠탭 순위, 포맷
        for a in sorted_areas:
            code = a.get('n','')
            if code in CONTENT_CODES:
                result['content_tab_rank'] = a.get('r', '-')
                result['content_format'] = SECTION_MAP.get(code, code)
                break

        # 상위 6개 탭순서
        top6 = []
        seen = set()
        for a in sorted_areas:
            name = SECTION_MAP.get(a.get('n',''), a.get('n',''))
            if name not in seen:
                seen.add(name)
                top6.append(name)
            if len(top6) >= 6: break
        result['top6_tabs'] = ' > '.join(top6) if top6 else '-'

        # 상위글 추출 (포맷 + 날짜)
        articles = []
        # 날짜 요소 찾기
        date_elements = soup.find_all(class_=re.compile(r'sds-comps-profile-info-subtext|sub_txt|date|upload_time|info_item'))
        # 링크+날짜 쌍 추출
        seen_urls = set()
        for container in soup.select('.view_wrap, .total_wrap, .api_subject_bx, .sp_tit, .total_tit, [class*="content_area"], [class*="view"]'):
            links = container.find_all('a', href=re.compile(r'blog\.naver|cafe\.naver|kin\.naver|post\.naver|in\.naver'))
            for link in links:
                href = link.get('href','')
                if href in seen_urls: continue
                seen_urls.add(href)
                fmt = '기타'
                if 'blog.naver' in href: fmt = '블로그'
                elif 'cafe.naver' in href: fmt = '카페'
                elif 'kin.naver' in href: fmt = '지식인'
                elif 'post.naver' in href: fmt = '포스트'
                elif 'in.naver' in href: fmt = '인플루언서'

                # 날짜 찾기 — 부모 컨테이너에서
                date_str = ''
                parent = link
                for _ in range(8):
                    parent = parent.parent
                    if not parent: break
                    de = parent.find(class_=re.compile(r'sub_txt|date|time|profile.*sub|info_item|sds-comps-profile-info-subtext|upload_time'))
                    if de:
                        date_str = de.get_text(strip=True)
                        break

                pub = _parse_date(date_str, today)
                days_ago = (today - pub).days if pub else None
                articles.append({
                    'format': fmt,
                    'date': pub.strftime('%Y.%m.%d') if pub else '',
                    'days_ago': days_ago
                })
                if len(articles) >= 5: break
            if len(articles) >= 5: break

        # fallback: 모든 블로그/카페 링크에서
        if len(articles) < 3:
            for link in soup.find_all('a', href=re.compile(r'blog\.naver|cafe\.naver|kin\.naver')):
                href = link.get('href','')
                if href in seen_urls: continue
                seen_urls.add(href)
                fmt = '블로그' if 'blog' in href else '카페' if 'cafe' in href else '지식인'
                articles.append({'format': fmt, 'date': '', 'days_ago': None})
                if len(articles) >= 5: break

        result['articles'] = articles

        # 경쟁강도 계산
        days_list = [a['days_ago'] for a in articles[:3] if a['days_ago'] is not None]
        if days_list:
            avg_days = sum(days_list) / len(days_list)
            if avg_days < 14: result['competition'] = '상'
            elif avg_days <= 28: result['competition'] = '중'
            else: result['competition'] = '하'
        else:
            result['competition'] = '-'

    except Exception as e:
        print(f"SERP analysis error for '{keyword}': {e}")
        result['competition'] = '-'
    return result

# ── 노션 저장 ──
# ── 블로그 상위글 분석 ──
def _analyze_blog_article(url, keyword):
    """개별 블로그 글 분석: 사진수, 키워드반복수"""
    try:
        m = re.search(r'blog\.naver\.com/([^/?]+)/(\d+)', url)
        if not m:
            return None
        blog_id, log_no = m.group(1), m.group(2)
        mobile_url = f"https://m.blog.naver.com/{blog_id}/{log_no}"
        headers = {'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1'}
        r = req.get(mobile_url, headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        content = soup.select_one('.se-main-container, .post_ct, #postViewArea')
        if content:
            imgs = content.find_all('img', src=re.compile(r'http'))
            photo_count = len(imgs)
            text = content.get_text()
        else:
            imgs = soup.find_all('img', src=re.compile(r'blogfiles|postfiles|pstatic'))
            photo_count = len(imgs)
            text = soup.get_text()
        keyword_repeat = text.count(keyword)
        return {'photo_count': photo_count, 'keyword_repeat': keyword_repeat}
    except Exception:
        return None

def _analyze_top_for_blog(keyword):
    """상위글 분석 → 평균 사진수, 키워드반복수"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get(f"https://search.naver.com/search.naver?query={quote(keyword)}&where=nexearch", headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        urls = []
        for a in soup.find_all('a', href=re.compile(r'blog\.naver\.com/[^/]+/\d+')):
            href = a.get('href', '')
            if href not in urls:
                urls.append(href)
            if len(urls) >= 3:
                break
        if not urls:
            return {'photo_count': 8, 'keyword_repeat': 5}
        results = []
        for url in urls[:3]:
            a = _analyze_blog_article(url, keyword)
            if a:
                results.append(a)
            time.sleep(0.5)
        if not results:
            return {'photo_count': 8, 'keyword_repeat': 5}
        return {
            'photo_count': max(round(sum(r['photo_count'] for r in results) / len(results)), 3),
            'keyword_repeat': max(round(sum(r['keyword_repeat'] for r in results) / len(results)), 3)
        }
    except Exception:
        return {'photo_count': 8, 'keyword_repeat': 5}

# ── 카페 상위글 분석 ──
def _analyze_cafe_article(url, keyword):
    """개별 카페 글 분석: 사진수, 키워드반복수, 글자수"""
    try:
        mobile_url = url.replace('cafe.naver.com', 'm.cafe.naver.com')
        headers = {'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X)'}
        r = req.get(mobile_url, headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        body = soup.get_text()
        photo_count = len(soup.find_all('img', src=re.compile(r'cafeptthumb|postfiles|blogfiles|phinf')))
        kw_repeat = body.lower().count(keyword.lower())
        char_count = len(body.replace(' ', '').replace('\n', ''))
        return {'photo_count': max(photo_count, 1), 'keyword_repeat': max(kw_repeat, 1), 'char_count': char_count}
    except Exception:
        return None

def _analyze_top_for_cafe(keyword):
    """카페 상위글 3개 분석 → 평균 사진수, 키워드반복수, 글자수"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get(f"https://search.naver.com/search.naver?query={quote(keyword)}&where=article", headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        urls = []
        for a in soup.find_all('a', href=re.compile(r'cafe\.naver\.com.*ArticleRead')):
            href = a.get('href', '')
            if href not in urls:
                urls.append(href)
            if len(urls) >= 3:
                break
        if not urls:
            return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500}
        results = []
        for url in urls[:3]:
            a = _analyze_cafe_article(url, keyword)
            if a:
                results.append(a)
            time.sleep(0.5)
        if not results:
            return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500}
        return {
            'photo_count': max(round(sum(r['photo_count'] for r in results) / len(results)), 3),
            'keyword_repeat': max(round(sum(r['keyword_repeat'] for r in results) / len(results)), 3),
            'char_count': max(round(sum(r['char_count'] for r in results) / len(results)), 800)
        }
    except Exception:
        return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500}

# ── API 사용량 추적 ──
API_USAGE_FILE = os.path.join(os.path.dirname(__file__), "api_usage.json")
_usage_lock = threading.Lock()

# Sonnet 4 가격 ($ per 1M tokens)
PRICING = {
    "claude-sonnet-4-20250514": {"input": 3.0, "output": 15.0},
    "claude-haiku-4-5-20251001": {"input": 0.80, "output": 4.0},
}

def _track_usage(model, input_tokens, output_tokens, channel="unknown"):
    """API 사용량을 파일에 기록."""
    try:
        with _usage_lock:
            if os.path.exists(API_USAGE_FILE):
                with open(API_USAGE_FILE, "r") as f:
                    usage = json.load(f)
            else:
                usage = {"records": []}

            pricing = PRICING.get(model, {"input": 3.0, "output": 15.0})
            cost = (input_tokens / 1_000_000 * pricing["input"]) + (output_tokens / 1_000_000 * pricing["output"])

            usage["records"].append({
                "timestamp": datetime.now().isoformat(),
                "date": datetime.now().strftime("%Y-%m-%d"),
                "model": model,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "cost_usd": round(cost, 6),
                "channel": channel,
            })

            with open(API_USAGE_FILE, "w") as f:
                json.dump(usage, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[Usage tracking] 기록 실패: {e}")


# ── Claude API 호출 ──
def _call_claude(system_prompt, user_prompt, temperature=None, max_tokens=4096, channel="unknown"):
    """Claude API 호출 (비스트리밍). temperature: 0.0~1.0 (None이면 기본값). 429/5xx 시 최대 3회 재시도."""
    if not ANTHROPIC_API_KEY:
        return '[ERROR] ANTHROPIC_API_KEY가 설정되지 않았습니다. 환경변수로 설정하거나 server.py에서 직접 입력해주세요.'
    headers = {
        'Content-Type': 'application/json',
        'x-api-key': ANTHROPIC_API_KEY,
        'anthropic-version': '2023-06-01',
    }
    model = 'claude-sonnet-4-20250514'
    payload = {
        'model': model,
        'max_tokens': max_tokens,
        'system': system_prompt,
        'messages': [{'role': 'user', 'content': user_prompt}],
    }
    if temperature is not None:
        payload['temperature'] = temperature
    max_retries = 3
    for attempt in range(max_retries):
        try:
            r = req.post('https://api.anthropic.com/v1/messages', headers=headers, json=payload, timeout=60)
            if r.status_code == 200:
                data = r.json()
                # 사용량 추적
                usage = data.get('usage', {})
                _track_usage(model, usage.get('input_tokens', 0), usage.get('output_tokens', 0), channel)
                if data.get('content') and len(data['content']) > 0:
                    return data['content'][0]['text']
                return '[ERROR] Claude API 빈 응답'
            if r.status_code == 429 or r.status_code >= 500:
                wait = min(2 ** attempt * 5, 30)
                print(f"[Claude API] {r.status_code} 재시도 {attempt+1}/{max_retries} ({wait}초 대기)")
                time.sleep(wait)
                continue
            return f'[ERROR] Claude API {r.status_code}: {r.text[:300]}'
        except req.exceptions.Timeout:
            if attempt < max_retries - 1:
                print(f"[Claude API] 타임아웃 재시도 {attempt+1}/{max_retries}")
                time.sleep(3)
                continue
            return '[ERROR] Claude API 타임아웃 (60초 초과)'
        except Exception as e:
            return f'[ERROR] Claude API 호출 실패: {e}'
    return '[ERROR] Claude API 최대 재시도 횟수 초과'

def _build_blog_title_prompt(keyword, product):
    """블로그 제목 프롬프트 (STEP 1)"""
    system = """당신은 네이버 블로그의 상위 노출을 목표로 하는 제목 작성 전문가입니다.

주어진 키워드를 기반으로 클릭을 유도할 수 있는 '한 문장 제목'을 1개 생성해주세요.

[제목 작성 규칙]

1. 키워드는 제목에 반드시 1회 자연스럽게 포함
2. 후기처럼 보이게 작성 (예: "진짜 이틀 만에 해결됨", "저만 그런 줄 알았어요")
3. 기호, 과장 표현 금지: "대박템", "~강추", 특수문자 나열 등 제외
4. 말투는 자연스럽고 말하듯이, 궁금증이나 결과 중심 포맷이면 더 좋음
5. 한 문장으로 간결하게 출력 (제목만 출력, 다른 설명 없이)

[출력 예시]
- "{keyword} 진짜 이틀 만에 사라졌어요"
- "{keyword} 저만 겪는 줄 알았는데..."
- "이거 하나로 {keyword} 해결한 후기\""""

    user = f"상위 노출 키워드: {keyword}"
    return system, user

def _build_blog_body_prompt(keyword, stage, product, photo_count, keyword_repeat, title=''):
    """블로그 본문 프롬프트 (STEP 2) — {title}에 STEP 1 결과가 들어옴"""
    system = """📌 역할:
당신은 후기형 네이버 블로그 글의 오프닝부터 본문, 클로징까지 작성하는 작가입니다.

📏 필수 지침: 글자 수 강제 조건
- 이 글의 전체 분량은 공백 포함 2,200자 이상이어야 합니다.
- 단락 수는 최소 8단락 이상
- 각 단락은 200~350자 내외가 되도록 하며,
- 절대 압축 요약하지 말고 느낌, 사례, 생각, 비교 등을 풍부하게 풀어 써주세요.

---

🧩 작성 구조 (5단 구성):

| 구간 | 내용 | 비중 |
|------|------|------|
| 1. 오프닝 | 고민/문제 공감 → 독자가 "내 얘기잖아" 느끼게 | 15% |
| 2. 시도 & 실패 | 다른 방법 시도했지만 실패한 경험 → 객관성 확보 | 15% |
| 3. 발견 & 체험 | 제품을 알게 된 계기 + 초기 사용 과정 | 25% |
| 4. 변화 & 결과 | 구체적 수치로 변화 증명 (kg, 사이즈, 기간 등) | 30% |
| 5. 마무리 | 재구매 의사 + 조건부 추천 ("저한테는 맞았어요") | 15% |

핵심 전략: "공감 → 실패 → 해결 → 변화"의 4단 흐름
- 내 장점을 어필하지 말고, 차선책을 하나씩 지워가며 해결책으로 자연스럽게 연결
- 광고처럼 보이지 않고, "현실적인 공감 + 자연스러운 해결 흐름"을 만들 것

---

🎯 구매여정 단계별 오프닝 분기:

오프닝은 키워드를 보고 소비자가 어떤 심리로 클릭했을지 판단하여 작성합니다.
아래 단계 중 가장 적합한 것을 선택하세요.

● 문제 인식 (0_무지 ~ 1_인지):
"두피가 자꾸 가렵고 각질이 생기는데 왜 그런지 몰랐어요."
→ 공감 중심, 불편한 일상 묘사

● 해결 탐색 (2_호기심 ~ 3_정보습득):
"두피가려움 때문에 이것저것 알아보다가 스케일링, 오일, 샴푸까지 다 찾아봤어요."
→ 다양한 방법 탐색 + 문제 간략 언급

● 해결방법 선택 (3_정보습득):
"여러 가지 다 실패하다가 결국 ○○제품을 알게 됐어요."
→ 자사 제품 선택 계기 + 기대감 묘사

● 비교 검토 (4_의심):
"경쟁사 제품에는 ○○이 없어서 아쉬웠는데, 이 제품은 ○○이 있어서 만족했어요."
→ 경쟁 제품 언급 + 자사 USP 강조

● 구매 직전 (5_구매직전):
"후기들을 보다가 제 기준엔 이게 괜찮더라구요."
→ 사용 후기 중심 + 신뢰 강조 + 링크 클릭 심리 유도

---

📌 오프닝 작성 세부 규칙:
- 제품 이야기를 절대 먼저 꺼내지 않음. 독자의 고민에서 시작
- 나이, 직업, 가족 상황 등 구체적인 페르소나 설정
- 감정적 고통의 디테일 ("두려웠다", "포기하고 싶었다")
- 키워드를 1회 자연스럽게 포함

📌 시도 & 실패 세부 규칙:
- 다른 방법을 시도했지만 효과 없었던 경험을 솔직하게 서술
- 이 단계가 있어야 '광고'가 아닌 '진짜 후기'로 인식됨

📌 발견 & 체험 세부 규칙:
- 제품을 처음 알게 된 계기를 자연스럽게 서술
- "먹자마자 효과가 있었다" 같은 과장 표현 금지
- "~느낌이 들었다", "~것 같다" 수준으로 표현

📌 변화 & 결과 세부 규칙:
- 반드시 구체적인 숫자 포함 (체중, 사이즈, 체지방률, 기간 등)
- 사진과 함께 배치될 때 설득력이 극대화됨

📌 마무리 세부 규칙:
- 재구매 의사 + 조건부 추천
- 나만의 키워드를 자연스럽게 삽입 ("궁금하신 분들은 {나만의 키워드}로 검색해보세요")

---

📌 본문 전개 규칙:
- 경쟁 제품의 단점 서술
- 내가 선택한 제품 설명
- USP는 리뷰처럼 자연스럽게 녹이기 (예: "성분이 순해서 그런지 간지러움이 덜했어요.")

📌 이미지 삽입 규칙:
- 이미지는 이미지에 대한 설명을 쓰지 않고 순서만 적음
- [이미지1], [이미지2] ... 형태로 표기
- 총 지정된 사진수만큼 삽입
- 글 – 사진 – 글 – 사진의 일정한 리듬으로 배치
- 이미지 블록 앞뒤로는 빈 줄 1줄

📌 키워드 사용 규칙:
- 키워드는 전체 글에서 지정된 횟수만큼만 삽입
- 억지 삽입 없이 문맥 자연스럽게 유지

📌 클로징 (마무리 단락):
- 후기처럼 마무리
- 직접적인 구매 유도 문장(X), 대신 제안형 문장(O)
- 마지막 문장에는 [링크]라고 적어주고, 그 링크를 클릭하도록 심리를 자극할것.
  다음 리스트에서 무작위로 한 문장을 선택해 출력하세요.
  절대 2개 이상 출력하지 말고, 반드시 1개만 선택해 그대로 출력하세요.
  - 이번주까지만 특가 이벤트 진행한다고 해요. 링크 남겨드릴테니 한번 보시는것도 좋을것 같네요
  - 이제품 첼린지도 하고 난리났더라구요. 링크 남겨드릴테니 후기랑 상세페이지 잘 비교해보시고 구매하시길 바라요
  - 이번에 못 사면 다음 입고까지 한참 기다려야 한다고 하네요ㅠㅠ 제가 구매했던 링크 남겨드릴게요
  - 저는 30%할인할때 샀는데 지금도 하는지는 모르겠네요. 제가 구매했던 링크 남겨드릴게요
  - 단순 건강식품이 아니라 기능성 인정 받은 제품이라 믿고 구매했습니다. 궁금하신분들을 위해 링크 남겨드릴게요
  - 배송 지연 안내 보고 늦기 전에 결제했는데 다행히 다음날 바로 도착하더라구요. 필요하신분들은 아래링크 참고하세요!
  - 저도 여기 보고 알게 돼서 들어가봤는데 확실히 다르더라구요. 제가 구매했던 링크 남겨드릴게요
  - 요즘 카페랑 SNS에서 후기가 폭주 중이던데.. 링크남겨드릴테니 필요하신분은 확인해보세요

---

✍️ 톤 & 문체:
- 말하듯이, 진솔하고 공감하는 톤. 존댓말 유지(~요).
- 과장된 광고 문구 자제. 정보+체험 혼합 톤.
- 동일 의미의 군더더기 반복은 가볍게 정리하되, 전체 분량은 ±10% 이내로 유지 (요약 금지)

🚫 금지 표현:
- "치료", "완치", "효과 보장" 등 의료법 위반 표현
- "대박", "강추", "꼭 사세요" 등 직접적 구매 유도
- 이모지 사용 금지
- 불릿/번호 목록, 표, 코드블록 출력 금지

---

[형식 규칙]

1. 문장 단위 줄바꿈
   - 한 줄에는 반드시 한 문장만 둔다.
   - 목표 길이: (공백 포함) 22~40자/줄.
   - 허용 범위: 18~55자/줄. 55자를 넘기면 반드시 문장을 둘로 나눠라.

2. 문장 길이가 애매할 때의 분할 기준 (우선순위 순)
   A. 쉼표(,), '그리고/그래서/하지만/또한/다만/혹은' 같은 접속사 앞에서 끊는다.
   B. 조사 '고/며/면서/지만/라도' 앞에서 자연스럽게 끊는다.
   C. 수식어구가 길면 수식어와 핵심 서술(주어+서술어) 사이에서 끊는다.
   D. 의미 단위가 약하면 어미를 조정해 두 개의 완결 문장으로 만든다.
   E. 위 기준이 모두 어색하면, 핵심 주장/사실 우선 문장을 앞으로, 보충/예시는 다음 줄로 보낸다.
   - 끊은 뒤에는 각 줄이 '완결된 문장'이 되도록 종결어미(~요/~다)로 마무리한다.
   - '…'로 문장을 끝내며 끊지 않는다. (생략 부호는 문장 중간에서만 사용)

3. 문단 구성
   - 문단은 2~4줄(=2~4문장)로 묶는다.
   - 기본은 3줄, 흐름상 길거나 짧을 땐 2줄 또는 4줄을 허용한다.
   - 문단과 문단 사이에는 빈 줄 1줄만 둔다. (2줄 이상 금지)

4. 제목과 소제목
   - 제목은 첫 줄에 사용하되, Markdown H1(# 제목) 표기만 허용.
   - 제목 다음에는 빈 줄 1줄을 둔다. 소제목이 없으면 새로 만들지 않는다.

5. 이미지 플레이스홀더
   - [이미지N] 형식은 그대로 보존하고, 한 줄을 독립 문단처럼 배치한다.
   - 순서/번호를 임의 수정하거나 추가/삭제하지 않는다.
   - 이미지 블록 앞뒤로는 빈 줄 1줄을 둔다.

6. 기호/문장부호
   - 큰따옴표는 그대로 유지하되, 불필요한 인용부호 추가 금지.
   - 마침표 남발 금지. 그러나 각 줄은 완결 문장으로 끝낸다.
   - 이모지는 사용하지 않는다. 불릿/번호 목록, 표, 코드블록 출력 금지.

7. 금지 사항
   - 줄글(한 줄에 두 문장 이상) 금지.
   - 문단 사이 공백 2줄 이상 금지.
   - 제목 추가 생성/임의 소제목 생성 금지.
   - 임의 요약/삭제/추가 정보 삽입 금지.
   - 문장 중간 강제 개행(의미 단절) 금지.

[출력 예시 미니 샘플] ← 구조만 참고 (실제 내용 아님)

# 제목 예시

첫 문장은 문제 상황을 짧게 제시해요.
다음 문장은 맥락을 연결해요.
세 번째 문장은 독자의 공감을 이끌어요.

해결을 위해 시도했던 방법을 간단히 말해요.
핵심 실패 요인을 한 문장으로 정리해요.
전환 문장으로 다음 단락을 예고해요.

[이미지1]

제품을 사용한 계기를 한 문장으로 말해요.
사용 직후 느낀 포인트를 한 문장으로 말해요.
다음 문장으로 결과를 간결히 마무리해요."""

    user = f"""[시스템 자동 전달]
제목: {title}

[사용자 입력]
상위 노출 키워드: {keyword}
제품명: {product.get('name', '')}
제품 USP (차별 포인트): {product.get('usp', '')}
타겟층: {product.get('target', '')}
주요 성분: {product.get('ingredients', '')}
나만의 키워드: {product.get('brand_keyword', '')}
구매여정 단계: {stage}
사진 수: {photo_count}장
키워드 반복 수: {keyword_repeat}회

위 정보를 기반으로, 제목과 맥락이 맞는 후기형 블로그 본문을 작성해주세요."""

    return system, user

# 하위 호환: 기존 코드에서 _build_blog_prompts 호출하는 곳 대응
def _build_blog_prompts(keyword, stage, product, photo_count, keyword_repeat):
    """하위 호환용 — 제목+본문 통합 프롬프트 (레거시)"""
    return _build_blog_body_prompt(keyword, stage, product, photo_count, keyword_repeat, '')

def _save_notion(keyword_data):
    headers = {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '키워드': {'title': [{'text': {'content': keyword_data['keyword']}}]},
        '상태': {'select': {'name': '미사용'}},
    }
    if keyword_data.get('competition') and keyword_data['competition'] != '-':
        props['경쟁 강도'] = {'select': {'name': keyword_data['competition']}}
    if keyword_data.get('contact_point'):
        props['구매여정_단계'] = {'select': {'name': keyword_data['contact_point']}}

    payload = {'parent': {'database_id': KEYWORD_DB_ID}, 'properties': props}
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers, json=payload, timeout=10)
        return r.status_code == 200
    except Exception:
        return False

# ── 중간저장 ──
def _save_progress(results, remaining):
    try:
        with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
            json.dump({'results': results, 'remaining': remaining, 'ts': datetime.now().isoformat()}, f, ensure_ascii=False)
    except Exception as e:
        print(f"[_save_progress] 저장 실패: {e}")

def _load_progress():
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

# ───────────────────────────── API ENDPOINTS ─────────────────────────────

# ── 카페24 API ──

def _cafe24_load_token():
    if os.path.exists(CAFE24_TOKEN_FILE):
        with open(CAFE24_TOKEN_FILE, 'r') as f:
            return json.load(f)
    return {}

def _cafe24_save_token(token_data):
    with open(CAFE24_TOKEN_FILE, 'w') as f:
        json.dump(token_data, f, ensure_ascii=False)

def _cafe24_refresh_token():
    token = _cafe24_load_token()
    if not token.get('refresh_token'):
        return None
    auth = base64.b64encode(f'{CAFE24_CLIENT_ID}:{CAFE24_CLIENT_SECRET}'.encode()).decode()
    r = req.post(f'https://{CAFE24_MALL_ID}.cafe24api.com/api/v2/oauth/token',
                 headers={'Authorization': f'Basic {auth}', 'Content-Type': 'application/x-www-form-urlencoded'},
                 data={'grant_type': 'refresh_token', 'refresh_token': token['refresh_token']}, timeout=10)
    if r.status_code == 200:
        new_token = r.json()
        new_token['refresh_token'] = new_token.get('refresh_token', token['refresh_token'])
        _cafe24_save_token(new_token)
        return new_token
    return None

def _cafe24_api(endpoint):
    token = _cafe24_load_token()
    if not token.get('access_token'):
        token = _cafe24_refresh_token()
    if not token:
        return None
    headers = {'Authorization': f'Bearer {token["access_token"]}', 'Content-Type': 'application/json',
               'X-Cafe24-Api-Version': '2024-03-01'}
    r = req.get(f'https://{CAFE24_MALL_ID}.cafe24api.com/api/v2/{endpoint}', headers=headers, timeout=15)
    if r.status_code == 401:
        token = _cafe24_refresh_token()
        if token:
            headers['Authorization'] = f'Bearer {token["access_token"]}'
            r = req.get(f'https://{CAFE24_MALL_ID}.cafe24api.com/api/v2/{endpoint}', headers=headers, timeout=15)
    if r.status_code == 200:
        return r.json()
    return None

@app.get("/api/cafe24/auth-url")
async def cafe24_auth_url():
    """카페24 OAuth 인증 URL 생성"""
    scope = 'mall.read_salesreport,mall.read_order,mall.read_analytics'
    redirect_uri = os.environ.get('CAFE24_REDIRECT_URI', '')
    url = f'https://{CAFE24_MALL_ID}.cafe24api.com/api/v2/oauth/authorize?response_type=code&client_id={CAFE24_CLIENT_ID}&scope={scope}&redirect_uri={redirect_uri}'
    return {'url': url}

@app.post("/api/cafe24/auth-callback")
async def cafe24_auth_callback(request: Request):
    """인증코드로 토큰 발급"""
    body = await request.json()
    code = body.get('code', '')
    if not code:
        return JSONResponse({'error': 'code 필요'}, 400)
    auth = base64.b64encode(f'{CAFE24_CLIENT_ID}:{CAFE24_CLIENT_SECRET}'.encode()).decode()
    try:
        r = req.post(f'https://{CAFE24_MALL_ID}.cafe24api.com/api/v2/oauth/token',
                     headers={'Authorization': f'Basic {auth}', 'Content-Type': 'application/x-www-form-urlencoded'},
                     data={'grant_type': 'authorization_code', 'code': code,
                           'redirect_uri': os.environ.get('CAFE24_REDIRECT_URI', '')}, timeout=10)
        if r.status_code == 200:
            _cafe24_save_token(r.json())
            return {'ok': True}
        return JSONResponse({'ok': False, 'error': r.text[:300]}, 400)
    except Exception as e:
        return JSONResponse({'ok': False, 'error': str(e)}, 500)

@app.get("/api/cafe24/status")
async def cafe24_status():
    """카페24 연동 상태 확인"""
    token = _cafe24_load_token()
    return {'connected': bool(token.get('access_token')), 'mall_id': CAFE24_MALL_ID}

@app.get("/api/cafe24/sales")
async def cafe24_sales(start: str = '', end: str = ''):
    """매출/주문 데이터 조회"""
    if not start or not end:
        today = time.strftime('%Y-%m-%d')
        start = start or today
        end = end or today
    data = _cafe24_api(f'admin/orders/count?start_date={start}&end_date={end}')
    sales = _cafe24_api(f'admin/salesreport?start_date={start}&end_date={end}')
    return {'orders': data, 'sales': sales, 'period': f'{start} ~ {end}'}

@app.get("/api/cafe24/analytics")
async def cafe24_analytics(start: str = '', end: str = ''):
    """접속/유입 통계 조회"""
    if not start or not end:
        today = time.strftime('%Y-%m-%d')
        start = start or today
        end = end or today
    data = _cafe24_api(f'admin/analytics/dailyvisits?start_date={start}&end_date={end}')
    return {'analytics': data, 'period': f'{start} ~ {end}'}

# ── UTM 관리 ──
UTM_FILE = os.path.join(os.path.dirname(__file__), "utm_history.json")

@app.post("/api/utm/generate")
async def utm_generate(request: Request):
    """UTM 파라미터 생성"""
    body = await request.json()
    base_url = body.get('url', '')
    channel = body.get('channel', '')
    keyword = body.get('keyword', '')
    campaign = body.get('campaign', '')
    if not base_url:
        return JSONResponse({'error': 'URL 필요'}, 400)
    params = {
        'utm_source': channel or 'direct',
        'utm_medium': 'content',
        'utm_campaign': campaign or keyword or 'default',
        'utm_term': keyword,
        'utm_content': channel
    }
    sep = '&' if '?' in base_url else '?'
    utm_url = base_url + sep + '&'.join(f'{k}={quote(str(v))}' for k, v in params.items() if v)
    # 이력 저장
    history = []
    if os.path.exists(UTM_FILE):
        with open(UTM_FILE, 'r') as f:
            history = json.load(f)
    history.append({'url': utm_url, 'channel': channel, 'keyword': keyword, 'campaign': campaign, 'created': time.strftime('%Y-%m-%d %H:%M')})
    with open(UTM_FILE, 'w') as f:
        json.dump(history[-500:], f, ensure_ascii=False)  # 최근 500개 유지
    return {'utm_url': utm_url, 'params': params}

@app.get("/api/utm/history")
async def utm_history():
    """UTM 생성 이력 조회"""
    if os.path.exists(UTM_FILE):
        with open(UTM_FILE, 'r') as f:
            return {'history': json.load(f)}
    return {'history': []}

@app.get("/")
async def serve_dashboard():
    return FileResponse(os.path.join(os.path.dirname(__file__), "dashboard.html"))

@app.get("/사용안내서.html")
async def serve_manual():
    return FileResponse(os.path.join(os.path.dirname(__file__), "사용안내서.html"), media_type="text/html")

@app.patch("/api/status/assign-channel")
async def assign_channel(body: dict):
    page_id = body.get('page_id', '')
    channels = body.get('channels', [])
    if not page_id:
        return JSONResponse({'ok': False, 'error': 'page_id required'}, 400)
    headers = {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {'배정 채널': {'multi_select': [{'name': ch} for ch in channels]}}
    try:
        r = req.patch(f'https://api.notion.com/v1/pages/{page_id}',
                      headers=headers, json={'properties': props}, timeout=10)
        return {'ok': r.status_code == 200}
    except Exception:
        return JSONResponse({'ok': False, 'error': 'Notion API error'}, 500)

@app.patch("/api/status/record-work")
async def record_work(body: dict):
    """콘텐츠 작업 기록 (배포 후)"""
    page_id = body.get('page_id', '')
    if not page_id:
        return JSONResponse({'ok': False, 'error': 'page_id required'}, 400)
    headers = {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {'발행_상태': {'select': {'name': '발행완료'}}}
    if body.get('posted_url'):
        props['발행_URL'] = {'url': body['posted_url']}
    if body.get('deploy_date'):
        props['생성일'] = {'date': {'start': body['deploy_date']}}
    if body.get('work_account'):
        props['작업계정'] = {'rich_text': [{'text': {'content': body['work_account']}}]}
    if body.get('work_cafe'):
        props['작업카페'] = {'rich_text': [{'text': {'content': body['work_cafe']}}]}
    if body.get('agency'):
        props['실행사'] = {'rich_text': [{'text': {'content': body['agency']}}]}
    if body.get('work_cost') is not None and body['work_cost'] != '':
        try:
            props['작업비용'] = {'number': int(body['work_cost'])}
        except Exception:
            pass
    try:
        r = req.patch(f'https://api.notion.com/v1/pages/{page_id}',
                      headers=headers, json={'properties': props}, timeout=10)
        return {'ok': r.status_code == 200}
    except Exception:
        return JSONResponse({'ok': False, 'error': 'Notion API error'}, 500)

@app.post("/api/keywords/upload-excel")
async def upload_excel(file: UploadFile = File(...)):
    content = await file.read()
    keywords = []
    if file.filename.endswith('.csv'):
        text = content.decode('utf-8-sig')
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            if row and row[0].strip():
                keywords.append(row[0].strip())
    else:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True)
        ws = wb.active
        for row in ws.iter_rows(min_col=1, max_col=1, values_only=True):
            if row[0] and str(row[0]).strip():
                keywords.append(str(row[0]).strip())
    return {'keywords': keywords}

@app.post("/api/keywords/expand")
async def expand_keywords(request: Request):
    body = await request.json()
    seeds = body.get('keywords', [])
    mode = body.get('mode', 'excel')  # 'excel' (1회 확장) 또는 'repeat' (반복 확장)

    async def generate():
      try:
        all_kws = {}
        visited = set()

        for s in seeds:
            all_kws[s] = '시드'
            yield f"data: {json.dumps({'type':'keyword','keyword':s,'source':'시드'}, ensure_ascii=False)}\n\n"

        loop = asyncio.get_running_loop()

        yield f"data: {json.dumps({'type':'progress','msg':'브라우저 시작 중...','cur':0,'total':0}, ensure_ascii=False)}\n\n"
        await _selenium_semaphore.acquire()
        driver = None
        try:
            driver = await loop.run_in_executor(executor, _create_driver)

            if mode == 'excel':
                # ── 모드1: 엑셀 기반 1회 확장 (자완/연관/함께찾는만) ──
                total = len(seeds)
                yield f"data: {json.dumps({'type':'progress','msg':f'엑셀 기반 확장: {total}개 키워드의 자완/연관/함께찾는 수집','cur':0,'total':total}, ensure_ascii=False)}\n\n"

                for i, kw in enumerate(seeds):
                    if kw in visited:
                        continue
                    visited.add(kw)

                    yield f"data: {json.dumps({'type':'progress','msg':f'확장 중: {kw} ({i+1}/{total}) | 누적 {len(all_kws)}개','cur':i+1,'total':total}, ensure_ascii=False)}\n\n"

                    ac = await loop.run_in_executor(executor, _autocomplete, kw)
                    for ak in ac:
                        if _valid_kw(ak) and ak not in all_kws:
                            all_kws[ak] = '자동완성'
                            yield f"data: {json.dumps({'type':'keyword','keyword':ak,'source':'자동완성'}, ensure_ascii=False)}\n\n"

                    related = await loop.run_in_executor(executor, _expand_selenium, driver, kw)
                    for rk, src in related.items():
                        if rk not in all_kws:
                            all_kws[rk] = src
                            yield f"data: {json.dumps({'type':'keyword','keyword':rk,'source':src}, ensure_ascii=False)}\n\n"

                    # 10개 시드마다 중간 저장
                    if (i + 1) % 10 == 0:
                        remaining_seeds = seeds[i+1:]
                        expand_snapshot = [{'keyword': k, 'source': s} for k, s in all_kws.items()]
                        await loop.run_in_executor(executor, _save_progress, expand_snapshot, remaining_seeds)
                        yield f"data: {json.dumps({'type':'saved','count':len(all_kws),'msg':f'중간 저장 완료 ({len(all_kws)}개)'}, ensure_ascii=False)}\n\n"

                    await asyncio.sleep(1.5)

                yield f"data: {json.dumps({'type':'progress','msg':f'엑셀 기반 확장 완료. 총 {len(all_kws)}개','cur':total,'total':total}, ensure_ascii=False)}\n\n"

            else:
                # ── 모드2: 반복 확장 (기존 BFS) ──
                queue = list(seeds)
                round_num = 0

                while queue:
                    round_num += 1
                    current_batch = list(queue)
                    queue = []

                    yield f"data: {json.dumps({'type':'progress','msg':f'[라운드 {round_num}] {len(current_batch)}개 키워드 확장 시작 (누적 {len(all_kws)}개)','cur':0,'total':len(current_batch)}, ensure_ascii=False)}\n\n"

                    for i, kw in enumerate(current_batch):
                        if kw in visited:
                            continue
                        visited.add(kw)

                        yield f"data: {json.dumps({'type':'progress','msg':f'[라운드 {round_num}] 확장 중: {kw} ({i+1}/{len(current_batch)}) | 누적 {len(all_kws)}개','cur':i+1,'total':len(current_batch)}, ensure_ascii=False)}\n\n"
                        ac = await loop.run_in_executor(executor, _autocomplete, kw)
                        for ak in ac:
                            if _valid_kw(ak) and ak not in all_kws:
                                all_kws[ak] = '자동완성'
                                if ak not in visited:
                                    queue.append(ak)
                                yield f"data: {json.dumps({'type':'keyword','keyword':ak,'source':'자동완성'}, ensure_ascii=False)}\n\n"

                        related = await loop.run_in_executor(executor, _expand_selenium, driver, kw)
                        for rk, src in related.items():
                            if rk not in all_kws:
                                all_kws[rk] = src
                                if rk not in visited:
                                    queue.append(rk)
                                yield f"data: {json.dumps({'type':'keyword','keyword':rk,'source':src}, ensure_ascii=False)}\n\n"

                        await asyncio.sleep(1.5)

                    new_count = len(queue)
                    yield f"data: {json.dumps({'type':'progress','msg':f'[라운드 {round_num} 완료] 신규 {new_count}개 발견 → 누적 {len(all_kws)}개','cur':len(current_batch),'total':len(current_batch)}, ensure_ascii=False)}\n\n"

                    if not queue:
                        yield f"data: {json.dumps({'type':'progress','msg':f'더 이상 새로운 키워드 없음. 총 {round_num}라운드 완료.','cur':1,'total':1}, ensure_ascii=False)}\n\n"
        finally:
            if driver:
                await loop.run_in_executor(executor, driver.quit)
            _selenium_semaphore.release()

        # ── 검색량 조회 ──
        kw_list = [k for k in all_kws.keys() if _valid_kw(k)]
        yield f"data: {json.dumps({'type':'progress','msg':f'검색량 조회 중... ({len(kw_list)}개)','cur':0,'total':1}, ensure_ascii=False)}\n\n"
        vol = await loop.run_in_executor(executor, _search_volume, kw_list)

        result_list = []
        for k in kw_list:
            v = vol.get(k, {'pc': 0, 'mo': 0})
            pc = v['pc'] if isinstance(v['pc'], int) else 0
            mo = v['mo'] if isinstance(v['mo'], int) else 0
            result_list.append({
                'keyword': k,
                'source': all_kws[k],
                'pc': pc,
                'mo': mo,
                'total': pc + mo,
            })

        yield f"data: {json.dumps({'type':'complete','keywords':result_list,'total':len(result_list)}, ensure_ascii=False)}\n\n"
      except Exception as e:
        print(f"[expand_keywords] 에러: {e}")
        yield f"data: {json.dumps({'type':'error','message':f'키워드 확장 중 오류: {e}'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/keywords/search-volume")
async def search_volume_api(request: Request):
    """중지 후 검색량만 별도 조회"""
    body = await request.json()
    keywords = body.get('keywords', [])
    loop = asyncio.get_running_loop()
    vol = await loop.run_in_executor(executor, _search_volume, keywords)
    return vol

@app.post("/api/keywords/analyze")
async def analyze_keywords(request: Request):
    body = await request.json()
    keywords = body.get('keywords', [])
    resume = body.get('resume', False)

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        today = datetime.now()
        results = []
        start_idx = 0

        # 이전 진행분 복구
        if resume:
            prog = _load_progress()
            if prog:
                results = prog.get('results', [])
                remaining = prog.get('remaining', [])
                keywords_to_process = remaining
                start_idx = len(results)
                yield f"data: {json.dumps({'type':'resume','existing':results,'start':start_idx}, ensure_ascii=False)}\n\n"
            else:
                keywords_to_process = keywords
        else:
            keywords_to_process = keywords

        total = start_idx + len(keywords_to_process)

        # 검색량 배치 조회
        yield f"data: {json.dumps({'type':'progress','msg':'검색량 조회 중...','cur':0,'total':total}, ensure_ascii=False)}\n\n"
        vol = await loop.run_in_executor(executor, _search_volume, keywords_to_process)

        for i, kw in enumerate(keywords_to_process):
            idx = start_idx + i + 1
            yield f"data: {json.dumps({'type':'progress','msg':f'분석 중: {kw}','cur':idx,'total':total}, ensure_ascii=False)}\n\n"

            try:
                # SERP 분석
                serp = await loop.run_in_executor(executor, _analyze_serp, kw, today)
                v = vol.get(kw, {'pc':0,'mo':0})
            except Exception as kw_err:
                print(f"[analyze] 키워드 '{kw}' 분석 실패: {kw_err}")
                serp = {'competition': '-', 'content_tab_rank': '-', 'content_format': '-', 'articles': [], 'top6_tabs': '-'}
                v = {'pc': 0, 'mo': 0}

            row = {
                'keyword': kw,
                'pc': v['pc'], 'mo': v['mo'], 'total': v['pc'] + v['mo'],
                'competition': serp.get('competition', '-'),
                'content_tab_rank': serp.get('content_tab_rank', '-'),
                'content_format': serp.get('content_format', '-'),
                'contact_point': '',
                'articles': serp.get('articles', []),
                'top6_tabs': serp.get('top6_tabs', '-'),
            }
            results.append(row)

            yield f"data: {json.dumps({'type':'result','row':row,'cur':idx,'total':total}, ensure_ascii=False)}\n\n"

            # 30개마다 중간저장
            if idx % 30 == 0:
                remaining = keywords_to_process[i+1:]
                await loop.run_in_executor(executor, _save_progress, results, remaining)
                yield f"data: {json.dumps({'type':'saved','count':len(results)}, ensure_ascii=False)}\n\n"

            await asyncio.sleep(1.5)

        # 완료 시 진행파일 삭제
        if os.path.exists(PROGRESS_FILE):
            os.remove(PROGRESS_FILE)

        yield f"data: {json.dumps({'type':'complete','total':len(results)}, ensure_ascii=False)}\n\n"
      except Exception as e:
        print(f"[analyze_keywords] 에러: {e}")
        yield f"data: {json.dumps({'type':'error','message':f'키워드 분석 중 오류: {e}'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/keywords/save-notion")
async def save_notion(request: Request):
    body = await request.json()
    items = body.get('items', [])
    loop = asyncio.get_running_loop()
    success = 0
    for item in items:
        ok = await loop.run_in_executor(executor, _save_notion, item)
        if ok: success += 1
        await asyncio.sleep(0.3)
    return {'success': success, 'total': len(items)}

@app.get("/api/keywords/check-progress")
async def check_progress():
    prog = _load_progress()
    if prog:
        return {'has_progress': True, 'count': len(prog.get('results',[])), 'timestamp': prog.get('ts','')}
    return {'has_progress': False}

# ── Gemini Flash API (접촉지점 판별) ──
GEMINI_SYSTEM_PROMPT = '아래 키워드 각각을 검색하는 사람이 구매여정 6단계 중 어디에 있는지 판별해. 키워드 자체의 맥락을 보고 판단해. 0_무지: 문제 자체를 모름 1_인지: 문제 인식 (예: 전립선 증상) 2_호기심: 해결책 탐색 (예: 전립선에 좋은 음식) 3_정보습득: 해결책 찾는 중 (예: 전립선 영양제) 4_의심: 비교/검토 (예: 전립선 영양제 부작용) 5_구매직전: 구매 의사 명확 (예: 전립선 영양제 추천) JSON으로 응답해: {"키워드": "단계", ...}'

def _call_gemini_contact_point(keywords_batch):
    """Gemini Flash로 접촉지점 판별 (50개씩)"""
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=" + GEMINI_API_KEY
    user_msg = '\n'.join(keywords_batch)
    payload = {
        'system_instruction': {'parts': [{'text': GEMINI_SYSTEM_PROMPT}]},
        'contents': [{'parts': [{'text': user_msg}]}],
        'generationConfig': {'temperature': 0.1, 'responseMimeType': 'application/json'},
    }
    try:
        r = req.post(url, json=payload, timeout=30)
        if r.status_code == 200:
            data = r.json()
            text = data['candidates'][0]['content']['parts'][0]['text']
            return json.loads(text)
        else:
            print("Gemini API error %d: %s" % (r.status_code, r.text[:200]))
            return {}
    except Exception as e:
        print("Gemini API error: %s" % e)
        return {}

@app.post("/api/keywords/contact-point")
async def contact_point(request: Request):
    """키워드 접촉지점 판별 (Gemini Flash, 50개씩 배치)"""
    body = await request.json()
    keywords = body.get('keywords', [])
    loop = asyncio.get_running_loop()
    result = {}
    batch_size = 50
    for i in range(0, len(keywords), batch_size):
        batch = keywords[i:i+batch_size]
        batch_result = await loop.run_in_executor(executor, _call_gemini_contact_point, batch)
        result.update(batch_result)
    return result

# ───────────────────────────── BLOG ENDPOINTS ─────────────────────────────

@app.get("/api/blog/notion-keywords")
async def blog_notion_keywords():
    """노션 키워드 DB에서 블로그 배정 키워드 조회"""
    headers = {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    payload = {
        'filter': {
            'and': [
                {'property': '배정 채널', 'multi_select': {'contains': '블로그'}},
                {'property': '상태', 'select': {'equals': '미사용'}},
            ]
        },
        'page_size': 100,
    }
    try:
        r = req.post(f'https://api.notion.com/v1/databases/{KEYWORD_DB_ID}/query', headers=headers, json=payload, timeout=15)
        if r.status_code != 200:
            return {'keywords': [], 'error': r.text[:300]}
        data = r.json()
        keywords = []
        for page in data.get('results', []):
            props = page.get('properties', {})
            title_prop = props.get('키워드', {}).get('title', [])
            kw = title_prop[0]['text']['content'] if title_prop else ''
            stage_prop = props.get('구매여정_단계', {}).get('select')
            stage_name = stage_prop.get('name', '') if stage_prop else ''
            if kw:
                keywords.append({'keyword': kw, 'stage': stage_name, 'page_id': page['id']})
        return {'keywords': keywords}
    except Exception as e:
        return {'keywords': [], 'error': str(e)}

@app.post("/api/blog/check-forbidden")
async def blog_check_forbidden(request: Request):
    """셀프모아 금칙어 검사"""
    body = await request.json()
    text = body.get('text', '')
    if not text:
        return {'forbidden_words': [], 'count': 0, 'clean': True}
    try:
        r = req.post('https://www.selfmoa.com/filter/wordcheck.php',
                     data={'title': text},
                     headers={'Content-Type': 'application/x-www-form-urlencoded',
                              'User-Agent': 'Mozilla/5.0'},
                     timeout=15)
        r.encoding = 'utf-8'
        # 빨간색 하이라이트된 단어 추출
        found = re.findall(r'<font\s+color=["\']?red["\']?>(.*?)</font>', r.text, re.IGNORECASE)
        if not found:
            found = re.findall(r'color:\s*red[^>]*>(.*?)</', r.text, re.IGNORECASE)
        unique = list(dict.fromkeys(found))  # 중복 제거, 순서 유지
        return {'forbidden_words': unique, 'count': len(unique), 'clean': len(unique) == 0}
    except Exception as e:
        return JSONResponse({'forbidden_words': [], 'count': 0, 'clean': True, 'error': str(e)}, 500)

@app.post("/api/blog/fix-forbidden")
async def blog_fix_forbidden(request: Request):
    """금칙어를 대체어로 수정"""
    body = await request.json()
    text = body.get('text', '')
    forbidden_words = body.get('forbidden_words', [])
    if not text or not forbidden_words:
        return {'fixed_text': text, 'replacements': []}
    words_str = ', '.join(forbidden_words)
    sys_prompt = f"""너는 네이버 블로그 콘텐츠 편집 전문가야.
아래 원고에서 금칙어를 자연스러운 대체어로 바꿔줘.
금칙어 목록: [{words_str}]

규칙:
- 원고의 맥락과 흐름을 유지하면서 금칙어만 대체
- 대체어는 네이버 블로그에서 안전한 단어로 선택
- 원고 전체 분량과 구조는 그대로 유지
- 수정된 전체 원고만 출력 (설명 없이)"""
    loop = asyncio.get_running_loop()
    fixed = await loop.run_in_executor(executor, _call_claude, sys_prompt, text)
    replacements = [{'from': w, 'to': '(수정됨)'} for w in forbidden_words]
    return {'fixed_text': fixed, 'replacements': replacements}

@app.post("/api/blog/generate")
async def blog_generate(request: Request):
    """블로그 원고 생성 (SSE)"""
    body = await request.json()
    keywords = body.get('keywords', [])
    product = body.get('product', {})

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords)
        for i, kw_data in enumerate(keywords):
            kw = kw_data['keyword']
            stage = kw_data.get('stage', '3_정보습득')

            msg1 = '[%d/%d] %s — 상위글 분석 중...' % (i+1, total, kw)
            yield _sse({'type': 'progress', 'msg': msg1, 'cur': i, 'total': total})
            analysis = await loop.run_in_executor(executor, _analyze_top_for_blog, kw)

            pc = analysis['photo_count']
            kr = analysis['keyword_repeat']
            # STEP 1: 제목 생성
            msg2 = '[%d/%d] %s — 제목 생성 중...' % (i+1, total, kw)
            yield _sse({'type': 'progress', 'msg': msg2, 'cur': i, 'total': total})
            overrides = _prompt_load_overrides()
            title_sys = overrides.get('블로그_제목', None)
            if title_sys:
                title_usr = f"상위 노출 키워드: {kw}"
            else:
                title_sys, title_usr = _build_blog_title_prompt(kw, product)
            title_raw = await loop.run_in_executor(executor, _call_claude, title_sys, title_usr)
            title = title_raw.strip().replace('제목:', '').replace('제목 :', '').strip()
            if '\n' in title:
                title = title.split('\n')[0].strip()

            # STEP 2: 본문 생성 (제목을 변수로 전달)
            msg3 = '[%d/%d] %s — 본문 생성 중... (사진%d장, 키워드%d회)' % (i+1, total, kw, pc, kr)
            yield _sse({'type': 'progress', 'msg': msg3, 'cur': i, 'total': total})
            body_sys = overrides.get('블로그_본문', None)
            if body_sys:
                body_usr = f"[시스템 자동 전달]\n제목: {title}\n\n[사용자 입력]\n상위 노출 키워드: {kw}\n제품명: {product.get('name','')}\n제품 USP (차별 포인트): {product.get('usp','')}\n타겟층: {product.get('target','')}\n주요 성분: {product.get('ingredients','')}\n나만의 키워드: {product.get('brand_keyword','')}\n구매여정 단계: {stage}\n사진 수: {pc}장\n키워드 반복 수: {kr}회\n\n위 정보를 기반으로, 제목과 맥락이 맞는 후기형 블로그 본문을 작성해주세요."
            else:
                body_sys, body_usr = _build_blog_body_prompt(kw, stage, product, pc, kr, title)
            body_text = await loop.run_in_executor(executor, _call_claude, body_sys, body_usr)
            body_text = body_text.strip()

            actual_repeat = body_text.count(kw)
            char_count = len(body_text)

            result = {
                'keyword': kw, 'stage': stage, 'title': title, 'body': body_text,
                'photo_count': pc, 'keyword_repeat': kr,
                'actual_repeat': actual_repeat, 'char_count': char_count,
                'page_id': kw_data.get('page_id', ''),
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[blog_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'블로그 원고 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/blog/save-notion")
async def blog_save_notion(request: Request):
    """블로그 원고를 노션 콘텐츠 DB에 저장"""
    body = await request.json()
    headers = {
        'Authorization': f'Bearer {NOTION_TOKEN}',
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '블로그'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('body_summary'):
        props['본문'] = {'rich_text': [{'text': {'content': body['body_summary'][:2000]}}]}
    if body.get('photo_count') is not None:
        props['사진수'] = {'number': body['photo_count']}
    if body.get('keyword_repeat') is not None:
        props['키워드_반복수'] = {'number': body['keyword_repeat']}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}

    body_text = body.get('body', '')
    if body_text:
        paragraphs = [p.strip() for p in body_text.split('\n\n') if p.strip()]
        children = []
        for para in paragraphs[:100]:
            for k in range(0, len(para), 2000):
                children.append({
                    'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}
                })
        payload['children'] = children[:100]

    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ───────────────────────────── CAFE SEO ─────────────────────────────

def _search_cafe_top_article(keyword):
    """네이버 검색에서 카페 상위 1위 글 URL+제목 찾기"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get('https://search.naver.com/search.naver?query=%s&where=nexearch' % quote(keyword), headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        # 카페 링크 찾기
        for a in soup.find_all('a', href=re.compile(r'cafe\.naver\.com')):
            href = a.get('href', '')
            if '/ArticleRead' in href or re.search(r'cafe\.naver\.com/[^/]+/\d+', href):
                title_text = a.get_text(strip=True)
                return {'url': href, 'title': title_text}
        # fallback: 카페 검색 탭
        r2 = req.get('https://search.naver.com/search.naver?query=%s&where=article' % quote(keyword), headers=headers, timeout=10)
        soup2 = BeautifulSoup(r2.text, 'html.parser')
        for a in soup2.find_all('a', href=re.compile(r'cafe\.naver\.com')):
            href = a.get('href', '')
            if '/ArticleRead' in href or re.search(r'cafe\.naver\.com/[^/]+/\d+', href):
                title_text = a.get_text(strip=True)
                return {'url': href, 'title': title_text}
    except Exception as e:
        print("Cafe search error: %s" % e)
    return {'url': '', 'title': ''}

def _crawl_cafe_article(url):
    """네이버 카페 글 크롤링 (제목, 본문)"""
    try:
        mobile_url = url.replace('cafe.naver.com', 'm.cafe.naver.com')
        if 'm.m.cafe' in mobile_url:
            mobile_url = mobile_url.replace('m.m.cafe', 'm.cafe')
        headers = {'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15'}
        r = req.get(mobile_url, headers=headers, timeout=10, allow_redirects=True)
        soup = BeautifulSoup(r.text, 'html.parser')
        title = ''
        body_text = ''
        for sel in ['.tit_area .title', 'h3.title_text', '.se-title-text', '.article_header .title']:
            el = soup.select_one(sel)
            if el:
                title = el.get_text(strip=True)
                break
        for sel in ['.se-main-container', '.article_viewer', '.ContentRenderer', '#postContent']:
            el = soup.select_one(sel)
            if el:
                body_text = el.get_text('\n', strip=True)
                break
        return {'title': title, 'body': body_text[:5000]}
    except Exception as e:
        print("Cafe crawl error: %s" % e)
        return {'title': '', 'body': ''}

def _build_cafe_title_prompt(keyword, original_title):
    system = """📌 역할:
너는 네이버 카페 상위노출을 위한 제목 리라이팅 전문가야.

[목표]
기존에 상위노출된 카페 제목을 기반으로, 메인 키워드를 유지하면서 마지막 말(어미·단어)만 자연스럽게 바꿔주는 리라이팅 작업을 수행한다.
광고처럼 보이지 않고, 궁금형/공감형 톤으로 바꿔 클릭을 유도할 수 있게 한다.

[지침]
1. 제목의 전체 구조는 유지
2. 마지막 말(어미 or 단어)만 자연스럽게 바꿔줘
3. 클릭을 유도할 수 있는 궁금형 or 공감형 문장으로 바꾸기
4. 광고 느낌 X ("추천", "인생템", "꼭 사세요" 등 금지)
5. 이모티콘 사용 금지
6. 문장 길이 25~35자 이내
7. 정보 탐색형 문장으로 구성 — '궁금증'을 자극하면 클릭률(CTR)이 올라감

[제목 유형 참고]
- 고민형: "(30대 여자) 출산탈모 때문에 미치겠어요ㅠㅠ"
- 직관형: "40대 되면 출산탈모 심해지나요?"
- 궁금형: "다이어트 영양제 진짜 효과 있나요?"
- 결핍형: "다이어트 영양제 먹고 붓기 빠졌다는 후기 봤는데 진짜인가요?"

[잘못된 제목 vs 올바른 제목]
- ❌ 다이어트영양제 추천드려요 → ✅ 다이어트 영양제 진짜 효과 있나요?
- ❌ 체지방 줄이는 약 후기 → ✅ 체지방 감소 도움된다는 영양제, 효과 있을까요?
- ❌ 출산탈모 걱정이네요 → ✅ (30대 여자) 출산탈모 때문에 미치겠어요ㅠㅠ

[출력 규칙]
- 출력은 제목 1줄만
- 불필요한 부가 문구 없이 제공
- 예시: 입력 "다이어트 자전거 도움되나요?" → 출력 "다이어트 자전거 효과 있나요?\""""
    if original_title:
        user = "- 기존 제목: %s\n- 메인 키워드: %s → 제목에 그대로 유지" % (original_title, keyword)
    else:
        user = "- 기존 제목: (없음 - 새로 생성)\n- 메인 키워드: %s → 제목에 그대로 유지" % keyword
    return system, user

def _build_cafe_body_prompt(keyword, title, original_body, settings, product):
    repeat_count = settings.get('repeat_count', 5)
    sub_keywords = settings.get('sub_keywords', '')
    forbidden = product.get('forbidden', '')
    photo_count = settings.get('photo_count', 8)
    char_target = settings.get('char_target', 1500)
    ref_text = original_body if original_body else '(없음 - 새로 생성)'

    system = """📌 역할:
너는 네이버 카페 상위노출 기준에 맞춰, 결핍 고민과 증상을 얘기하고 질문형식의 글을 작성하는 카피라이터야.

광고처럼 보이면 안 되고, 누구나 메인 키워드에 맞는 증상을 공감할 수 있는 글을 써서 오래 글을 읽을 수 있게 써야 해.

[본문 구조]
1. 오프닝 (결핍) — 증상과 문제 얘기, 공감 유도. 개인 고민, 문제 인식.
   예: "요즘 살이 안 빠져서 너무 스트레스예요ㅠㅠ 식단도 해보고 운동도 해봤는데 그대로예요."
2. 시도 (실패 경험) — 기존 해결 방법과 한계.
   예: "한약, 디톡스 주스, 카페인 다이어트 등 여러 방법 다 해봤는데 금방 요요가 오더라구요."
3. 전환 (대안 등장) — 새로운 제품/방법 선택과 고민.
   예: "그래서 이번엔 다이어트 영양제 쪽으로 알아봤어요. 식욕 억제보다는 체지방 대사를 돕는 쪽으로요."
4. 결과 (변화 체감) — 구체적인 후기.
   예: "2주 정도 먹었는데 붓기가 빠지고 피로도 덜 느껴지더라구요."
5. 마무리 (질문형 문장) — 사용자 참여 유도.
   예: "혹시 다이어트 영양제 꾸준히 먹어보신 분 계신가요?"

[키워드 규칙]
- 메인 키워드: 지정된 반복횟수 +1회 본문에 삽입
- 서브 키워드: 지정된 횟수 +1회 삽입
- 키워드 누락 시 실패
- 금칙어 포함 시 실패

[사진 표시 규칙]
- 문단과 문단 사이 또는 문장 중간에 어색하지 않게 [어울릴 사진] 삽입이라고 표기
- 어떤 이미지를 넣으면 좋을지 짧은 설명도 함께 작성
- 사진 수: 지정된 수 +1개 삽입
- 사진 유형: 손 위 제품 컷, 섭취 모습, 체중계, 생활 배경 등 실사용자 느낌
- 제품명 확대 촬영, 홍보용 배너 이미지 금지
- 예시:
  - [어울릴 사진: 운동 후 다이어트 영양제와 물컵이 함께 있는 컷]
  - [어울릴 사진: 체중계 수치 변화 사진]
  - [어울릴 사진: 아침 공복에 영양제 섭취하는 장면]

[문체 규칙]
- 줄바꿈은 3~5줄 간격
- 자연스러운 후기·질문 톤, 진솔한 체험 공유 느낌
- 광고 어투 금지: "강추", "인생템", "꼭 사세요", "효과 최고!", "무조건 됩니다!" 등
- 후기형·질문형·체험형 톤 유지 (명령형 X)
- 본문 중간중간 질문형 문장을 넣어 체류시간 확보

[출력 규칙]
- 출력 시 제목 없이 바로 본문만
- 글자 수: 지정된 글자수 ±100자 (공백 제외)"""

    user = """- 제목: %s
- 참고 경쟁사글 (없다면 무시하고 새로 작성): %s
- 메인 키워드: %s
- 메인 키워드 반복횟수: %d → 총 %d에서 +1회 본문 삽입
- 서브 키워드: %s → 서브 키워드 +1회 삽입
- 글자 수: %d ±100자 (공백 제외)
- 금칙어 목록: %s → 포함 시 실패
- 사진 수: %d +1개 삽입""" % (title, ref_text, keyword, repeat_count, repeat_count, sub_keywords, char_target, forbidden, photo_count)

    return system, user

def _build_cafe_comments_prompt(keyword, body_text, brand_keyword, alternatives=''):
    alt_section = ""
    if alternatives:
        alt_section = "\n사용자가 입력한 차선책 정보:\n\n" + alternatives
    else:
        alt_section = """
아래 기본 차선책을 참고하여 댓글을 작성하세요.

자연 식이요법 (식단 관리 및 음식으로 개선)
장점: 건강한 식습관으로 항산화 성분 및 영양 섭취, 비교적 안전함
단점: 꾸준히 지키기 어려움, 식단만으로 체지방 감소나 노화 개선 효과에는 한계
메시지 포인트: "음식만으로는 체지방을 줄이고 활력을 되찾기엔 한계가 있더라."

기본 영양제 및 비타민 (멀티비타민, 콜라겐 등)
장점: 필요한 기본 영양 보충 가능, 간편함
단점: 특허 성분이나 체지방 감소 등 특화 기능성은 부족, 떨어진 컨디션을 끌어올리기에는 역부족
메시지 포인트: "일반 영양제만으론 예전 같지 않은 체력이나 피부 탄력까지 끌어올리기엔 부족하더라구요."

전통 다이어트 방법 (한약, 디톡스 주스 등)
장점: 몸에 좋다고 전해지는 방법으로 일부 디톡스 효과 기대, 심리적 위안
단점: 과학적·임상적 검증 부족, 고가의 한약은 경제적 부담, 디톡스 주스만으로는 영양 불균형 초래
메시지 포인트: "민간요법이라고 좋다 해서 해봤지만, 솔직히 큰 효과는 못 느꼈고 돈만 들었어요."

피부과 시술 및 휴식 (전문의 관리, 일시적 휴식)
장점: 전문 관리와 휴식을 통해 일시적으로 피부나 컨디션 개선에 도움
단점: 단기간 일시적 케어에 그침 – 근본적인 체질 개선이나 체중 관리는 어려움, 지속 시 비용 부담 큼
메시지 포인트: "클리닉 가서 시술도 받아봤지만 잠깐일 뿐이고, 근본적으로 달라지는 건 없더라구요."

운동 및 식단 다이어트 (자체 체중 감량 시도)
장점: 체중 감량과 건강 회복에 도움, 특별한 비용 없음
단점: 바쁜 일상 속 운동 시간 확보 어려움, 무리한 식단 조절은 건강을 해칠 수 있음, 꾸준한 실천 어려움
메시지 포인트: "운동이나 식이조절이 좋다는 건 알지만, 막상 바쁘게 지내다 보면 실천이 거의 불가능하더라구요."

일반 다이어트 보조제 (다이어트 약, 카페인 함유 제품 등)
장점: 빠른 체중 감량 효과 기대, 복용 간편
단점: 카페인 등 성분으로 부작용 위험 (두근거림, 불면 등), 건강에 무리가 갈 수 있음, 요요 등 지속적 관리 어려움
메시지 포인트: "효과 빠르다는 다이어트 약은 부작용도 걱정되고 솔직히 무서워서 시도도 못 하겠더라구요."
"""

    system = """역할:
너는 네이버 카페 글에 달릴 자연스러운 댓글 10개를 생성하는 작가야.
댓글은 실제 사용자의 다양한 반응처럼 보여야 하고, 제품 삽입이 자연스러워야 해.

[작성 정보]

제품의 차선책 (경쟁 대안)
{alt_section}

[지시사항]

상황: 바쁜 일상과 스트레스 탓인지 피로, 기억력 저하, 체중 증가와 피부 탄력 저하로 고민하는 사람이 쓴 게시글에 달린 댓글들을 작성합니다. 각 댓글은 서로 다른 사람이 작성한 것으로, 앞 사람의 의견에 공감하거나 자신의 경험을 추가하면서 자연스럽게 이어지는 흐름이어야 합니다. 모든 댓글은 일상적인 존댓말 체로 쓰되, 너무 딱딱하지 않게 편안하고 친근한 어조를 사용하세요. 필요한 경우 감정을 나타내는 이모티콘이나 표현(예: ㅠㅠ, ㅎㅎ, ;;, 👍 등)을 적절히 섞어 자연스러운 분위기를 만드세요.

1~2번째 댓글 (공감 및 차선책 언급): 문제 상황에 깊이 공감하며 비슷한 어려움을 겪어봤다는 경험담으로 시작하세요. 많은 사람들이 흔히 시도하는 대안들을 언급하고, 본인도 열심히 해봤지만 효과를 거의 못 봤거나 일시적이었다는 식으로 써주세요. (주의: 이 단계에서는 특정 제품 언급이나 광고 티가 나면 안 됩니다!)

3~4번째 댓글 (해결책 제시 – 제품명 언급 없이 장점 소개): 새로운 해결 방법을 제시하는 흐름으로 전환합니다. 구체적으로 특화된 건강기능식품이 요즘 나와 있다는 정보를 공유하되, 아직 제품명은 직접 언급하지 마세요. (여전히 제품명이나 구매 유도는 하지 않습니다.)

5~6번째 댓글 (제품명 공개 및 경험담으로 추천): 앞에서 언급한 보조제가 무엇인지 자연스럽게 제품명으로 연결합니다. 마치 이미 써본 사람들처럼, 제품 복용 후 구체적인 변화를 경험담 형태로 작성하세요. 한 댓글에서는 제품을 쓰고 효과를 체감한 내용을 담습니다. 또 다른 댓글에서는 제품의 강점(특허 성분 함유, 가성비 등)과 함께 주효과를 언급하세요. 각각 다양한 사용자의 입장에서 제품을 사용해 만족했다는 흐름을 보여주세요. (링크, 구매 강조 ❌ 단, 제품명이 노출됩니다.)

7번째 댓글 (효과 언급): 이어서 제품명을 한 번 더 등장시켜, 고민 해결 경험담을 써주세요. 제품 복용 후 효과를 본 사례를 구체적으로 적습니다. 안전하게 효과를 봤다는 점을 강조하세요. 추가로 개인 소감을 덧붙여줍니다.

8~10번째 댓글 (추가 공감, 조언 및 마무리): 마지막 댓글들에서는 앞선 대화를 자연스럽게 이어받으며 마무리합니다. 다른 사람들도 등장하여 영양제 관리에 공감하고 꾸준한 자기관리의 중요성을 강조하세요. "방치하면 더 악화된다"는 경고를 살짝 넣어 건강 관리의 필요성을 상기시키고, "꾸준히 하면 좋아진다"는 희망적인 조언을 덧붙이세요. 사회적 증거와 긍정적 분위기로 끝맺습니다. 독자들이 "나도 해봐야겠다"는 마음이 들도록 유도하며 대화를 마무리하세요.

각 댓글에 메인 키워드 1개 이상은 포함

광고 어투 X (강추, 무조건, 최고 등 과장 표현 및 구매 유도 금지)

[출력형식]
한줄에 한개의 댓글을 모두 작성하며, 각각의 댓글 사이의 줄바꿈은 1회만 할 것. 절대 줄바꿈을 2회하지 않을 것 (중요).

[형식]
맞아요... 저도 요즘 부쩍 피로도 심하고 몸이 예전같지 않아서 좋다는 음식이며 영양제까지 열심히 챙겨먹어봤는데도 피로 회복이나 체중감량엔 효과가 거의 없었어요 ㅠㅠ 역시 식단 조절만으로는 한계가 있나봐요...
저도 휴가도 내서 푹 쉬어보고 좋다는 영양제도 찾아 먹고, 한약이랑 디톡스 주스까지 열심히 따라해봤는데 바쁘게 지내다 보니 금세 지치고 기억력 떨어진 건 그대로이더라구요 ㅠㅠ 피부 푸석해진 게 신경 쓰여서 피부과 시술까지 받아봤지만 솔직히 잠깐 나아진 듯한 느낌만 들고 비용만 많이 들었어요;; 운동이나 다이어트도 해보려 했지만 일하고 집안일 하랴 꾸준히 하기가 제 마음처럼 잘 안 되더라구요. 정말 쉽지 않아요.
요즘은 집에서 꾸준히 챙겨먹을 수 있는 건강기능식품들이 잘 나와 있더라구요. 단순 종합비타민제가 아니라 떨어진 체력이나 기억력, 푸석해진 피부 회복부터 체중감량까지 도움을 주는 특허 성분의 영양제들이요.
저도 최근에 그런 다이어트 영양제를 하나 먹기 시작했는데, 확실히 전보다 덜 피곤하고 멍했던 정신도 맑아지는 느낌이에요 👍 제품 고를 땐 원료가 특허받았는지, 또 식약처 기능성 인정을 받은 건지 꼭 확인해야 된다더라구요!
맞아요, 항산화 다이어트 영양제 중에서는 후기가 제일 좋길래 저도 다이어트 시작하면서 함께 먹어봤거든요? 꾸준히 먹었더니 몸도 한결 가벼워지고 피부톤도 환해지는 것 같아요. 괜히 많이들 찾는 게 아니더라구요.
국내 유일 특허 조합이 들어간 다이어트 영양제인데 가격도 생각보다 괜찮아서 좋았어요ㅠㅠ 저는 사실 나이 들면서 늘어난 뱃살 때문에 걱정이 많았는데, 먹고 나서 주변에서 얼굴 좋아졌단 소리도 듣고 체지방률도 조금 내려가서 만족 중이에요 🙂
저는 살 찐 게 가장 고민이었는데 식약처에서 기능성 인정받은 다이어트 영양제가 있어서 시작했어요. 카페인 들어간 다이어트제가 아니라 국내 특허 받은 폴리페놀 성분으로 만들었다고 해서 그런지 카페인 민감한 저도 부담 없이 천천히 살이 빠지더라구요. 요요 없고 안전하게 빼고 있어서 너무 만족해요!
사실 다이어트 약이나 병원 치료 같은 데에 의존하는 건 부담도 크고 장기적으로도 어렵잖아요… 이렇게 집에서 꾸준히 영양제로 관리하는 게 확실히 답인 것 같아요.
몸 컨디션 떨어지고 체중도 늘고 피부 노화 오는 걸 그냥 방치하면 나중에 체중감량이나 피부 회복이 더 힘들어진다고 하더라구요. 그러니까 영양제라도 미리 챙겨두는 게 낫다고 해서 저도 열심히 챙겨보려구요.
많은 분들이 이렇게 꾸준히 영양제 챙겨드시면서 기운도 되찾고 체중감량에도 성공하셨다니까 저도 계속 챙겨보려구요. 병원 안 가고도 이렇게 관리할 수 있다는 게 진짜 다행이고 신기하네요ㅎㅎ""".format(alt_section=alt_section)

    alt_display = alternatives if alternatives else '(AI 자동 판단)'
    user = f"""- 실제 본문: {body_text[:2000]}
- 메인 키워드: {keyword}
- 제품명 or 간접 언급: {brand_keyword or keyword}
- 차선책 (경쟁 대안, 선택사항): {alt_display}"""

    return system, user

@app.get("/api/cafe/notion-keywords")
async def cafe_notion_keywords():
    """노션 키워드 DB에서 카페SEO 배정 키워드 조회"""
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    payload = {
        'filter': {
            'and': [
                {'property': '배정 채널', 'multi_select': {'contains': '카페SEO'}},
                {'property': '상태', 'select': {'equals': '미사용'}},
            ]
        },
        'page_size': 100,
    }
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % KEYWORD_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code != 200:
            return {'keywords': [], 'error': r.text[:300]}
        data = r.json()
        keywords = []
        for page in data.get('results', []):
            props = page.get('properties', {})
            title_prop = props.get('키워드', {}).get('title', [])
            kw = title_prop[0]['text']['content'] if title_prop else ''
            if kw:
                keywords.append({'keyword': kw, 'page_id': page['id']})
        return {'keywords': keywords}
    except Exception as e:
        return {'keywords': [], 'error': str(e)}

@app.post("/api/cafe/generate")
async def cafe_generate(request: Request):
    """카페SEO 원고 생성 (SSE): 제목→본문→댓글"""
    body = await request.json()
    keywords = body.get('keywords', [])
    urls = body.get('urls', [])
    product = body.get('product', {})
    settings = body.get('settings', {})

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords)
        url_list = [u.strip() for u in urls if u.strip()] if urls else []

        for i, kw_data in enumerate(keywords):
            kw = kw_data['keyword']
            url = url_list[i] if i < len(url_list) else ''
            original_title = ''
            original_body = ''

            # 상위글 분석 (사진수/키워드반복수/글자수 자동 계산)
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 상위글 분석 중 (사진수/반복수/글자수)...' % (i+1, total, kw), 'cur': i, 'total': total})
            cafe_analysis = await loop.run_in_executor(executor, _analyze_top_for_cafe, kw)
            kw_settings = {**settings}  # 키워드별 로컬 복사
            kw_settings['photo_count'] = cafe_analysis['photo_count']
            kw_settings['repeat_count'] = cafe_analysis['keyword_repeat']
            kw_settings['char_target'] = cafe_analysis['char_count']

            # 경쟁사 글 크롤링 (URL 직접 입력 or 자동 검색)
            if url:
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 경쟁사 글 크롤링 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                crawled = await loop.run_in_executor(executor, _crawl_cafe_article, url)
                original_title = crawled['title']
                original_body = crawled['body']
            else:
                # URL 없으면 네이버 검색으로 카페 상위글 자동 크롤링
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 카페 상위글 검색 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                found = await loop.run_in_executor(executor, _search_cafe_top_article, kw)
                if found['url']:
                    yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 상위글 크롤링 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                    crawled = await loop.run_in_executor(executor, _crawl_cafe_article, found['url'])
                    original_title = crawled['title'] or found['title']
                    original_body = crawled['body']

            # STEP 1: 제목
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 제목 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys1, usr1 = _build_cafe_title_prompt(kw, original_title)
            title = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
            title = title.strip().split('\n')[0].strip()

            # STEP 2: 본문
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 본문 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys2, usr2 = _build_cafe_body_prompt(kw, title, original_body, kw_settings, product)
            cafe_body = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
            cafe_body = cafe_body.strip()

            # STEP 3: 댓글 10개
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 댓글 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys3, usr3 = _build_cafe_comments_prompt(kw, cafe_body, product.get('brand_keyword', ''), product.get('alternatives', ''))
            comments = await loop.run_in_executor(executor, _call_claude, sys3, usr3)
            comments = comments.strip()

            # STEP 4: 타겟 맞춤 사진 수집 (샤오홍슈)
            target = product.get('target', '')
            photo_count = kw_settings.get('photo_count', 8)
            # 본문에서 [어울릴 사진] 개수 세기
            img_slots = len(re.findall(r'\[어울릴 사진[^\]]*\]|\[이미지\d+\]', cafe_body))
            need_count = max(img_slots, photo_count) if img_slots else photo_count
            image_paths = []
            if target:
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 사진 수집 중 (샤오홍슈)...' % (i+1, total, kw), 'cur': i, 'total': total})
                image_paths = await loop.run_in_executor(executor, _collect_xhs_for_cafe, target, need_count)
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 사진 %d장 수집+모자이크 완료' % (i+1, total, kw, len(image_paths)), 'cur': i, 'total': total})

            # 이미지 파일명 목록 (프론트에 전달)
            image_filenames = [os.path.basename(p) for p in image_paths]

            result = {
                'keyword': kw, 'title': title, 'body': cafe_body, 'comments': comments,
                'original_title': original_title, 'original_body': original_body,
                'page_id': kw_data.get('page_id', ''),
                'photo_count': kw_settings.get('photo_count', 8),
                'repeat_count': kw_settings.get('repeat_count', 5),
                'images': image_filenames,
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[cafe_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'카페SEO 원고 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/cafe/save-notion")
async def cafe_save_notion(request: Request):
    """카페SEO 원고를 노션 콘텐츠 DB에 저장"""
    body = await request.json()
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '카페'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('body_summary'):
        props['본문'] = {'rich_text': [{'text': {'content': body['body_summary'][:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}
    if body.get('photo_count') is not None:
        props['사진수'] = {'number': body['photo_count']}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}

    # 본문 + 댓글을 페이지 children으로
    children = []
    for text_block in [body.get('body', ''), '---\n댓글 10개:\n' + body.get('comments', '')]:
        paragraphs = [p.strip() for p in text_block.split('\n\n') if p.strip()]
        for para in paragraphs:
            for k in range(0, len(para), 2000):
                children.append({
                    'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}
                })
    payload['children'] = children[:100]

    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

@app.get("/api/cafe/temp-image/{filename}")
async def cafe_temp_image(filename: str):
    """카페 원고용 임시 이미지 서빙"""
    fpath = os.path.join(TEMP_PHOTO_DIR, filename)
    if os.path.exists(fpath):
        return FileResponse(fpath, media_type="image/jpeg")
    return Response(status_code=404)

@app.post("/api/cafe/docx")
async def cafe_docx(request: Request):
    """카페 원고를 docx 파일로 생성 + 다운로드"""
    body = await request.json()
    title = body.get('title', '')
    body_text = body.get('body', '')
    keyword = body.get('keyword', '')
    image_filenames = body.get('images', [])

    image_paths = [os.path.join(TEMP_PHOTO_DIR, f) for f in image_filenames]
    loop = asyncio.get_running_loop()
    fpath, fname = await loop.run_in_executor(executor, _create_cafe_docx, title, body_text, image_paths, keyword)

    safe_fname = quote(fname)
    return FileResponse(
        fpath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename*=UTF-8''" + safe_fname}
    )

# ───────────────────────────── CAFE VIRAL ─────────────────────────────

def _build_viral_stage1_prompt(category, target='', topic=''):
    system = """너는 네이버 카페에 활동하는 평범한 회원이다.
카페에 올릴 짧은 일상글을 1개 작성한다.

[역할]
- 카페 내부 실제 회원처럼 행동하는 사람
- 광고주, 마케터, 홍보 담당자가 아닌 "그냥 일상을 나누는 사람"

[글 주제 범위]
- 일상 피로, 체력 관리, 육아, 수면, 날씨, 음식, 운동, 생활 고민 등
- 비상업적 주제만 허용
- 제품, 브랜드, 성분, 쇼핑 관련 내용 절대 금지

[톤 & 문체 규칙]
- 존댓말 + 따뜻한 대화체
- 문장 끝: "~했어요", "~하더라구요", "~인 것 같아요" 등 구어체
- 감정 표현: ㅎㅎ, ㅠㅠ, ;;, 👍 중 1~2개 자연스럽게 사용
- 문장 단락: 2~3문장씩 끊기
- 전체 길이: 3~5줄 이내 (짧고 가볍게)

[구조]
- 제목: 일상적 감정/상황을 담은 짧은 문장 (1줄)
- 본문: 일상 상황 묘사 → 가벼운 감정 표현 → 공감 유도 질문 or 마무리

[절대금지]
- 제품명, 브랜드명, 성분명 언급
- 광고 어투: "추천드려요", "꼭 써보세요", "이거 진짜 좋아요"
- 링크, 해시태그, 쇼핑몰 언급
- 정보 전달형 문체 (블로그 톤, 리뷰 톤)
- 같은 문장 구조 반복"""
    user = "다음 조건으로 카페 일상글을 작성해줘.\n\n- 타겟층: %s\n- 글 주제: %s\n\n위 조건에 맞는 카페 일상글을 제목과 본문으로 작성해줘." % (
        target or category, topic or '자유 (일상, 피로, 건강, 운동 등)')
    return system, user

def _build_viral_stage2_prompt(category, target_concern, product_category=''):
    system = """너는 네이버 카페에서 활동하는 평범한 회원이다.
생활 속 고민을 나누는 카페 글을 1개 작성한다.

[역할]
- 특정 고민을 가진 실제 카페 회원
- 광고주가 아닌, 진짜 고민을 가진 사람처럼 글을 쓰는 사람
- 해결책을 "광고"하는 것이 아니라, 경험을 "나누는" 톤

[글 구조] — 반드시 아래 순서를 따른다
1. 제목: 감정형 또는 공감형 문장 (예: "요즘 피로가 너무 심해서요ㅠㅠ")
2. 오프닝: 일상적 상황 묘사 (예: "애 재우고 나면 체력이 바닥이에요…")
3. 중반: 감정 공유 + 기존에 시도한 방법 (예: "커피도 마셔보고 비타민도 챙겨봤는데 그대로네요ㅠ")
4. 후반: 간접적 해결 언급 — 카테고리 수준만 (예: "친구가 알려준 영양제 먹고 좀 나아졌어요ㅎㅎ")
5. 마무리: 질문 또는 공감 유도 (예: "다들 피로할 땐 어떻게 관리하세요?")

[카테고리 언급 규칙]
- 제품명, 브랜드명 절대 금지
- "카테고리" 수준까지만 허용
  예시: "천연 성분 영양제", "수면 도움 제품", "두피관리제", "저자극 클렌저"
- 카테고리 언급은 후반부에서 1회만

[톤 & 문체 규칙]
- 존댓말 + 따뜻한 대화체
- 문장 끝: "~했어요", "~하더라구요", "~인 것 같아요" 등 구어체
- 감정 표현: ㅎㅎ, ㅠㅠ, ;;, 👍 중 1~2개 자연스럽게 사용
- 문장 단락: 2~3문장씩 끊기
- 전체 길이: 5줄 이하

[대화 흐름]
질문 → 공감 → 경험공유 → 마무리

[절대금지]
- 제품명, 브랜드명, 성분명(구체적 성분) 언급
- 광고 어투: "추천드려요", "꼭 써보세요", "이거 진짜 좋아요"
- 링크, 해시태그, 쇼핑몰 언급
- 정보 전달형 문체 (블로그 톤, 리뷰 톤)
- CTA (구매 유도, 검색 유도)
- 같은 문장 구조 반복"""
    user = "다음 조건으로 카페 고민글을 작성해줘.\n\n- 타겟층: %s\n- 고민 키워드: %s\n- 제품 카테고리: %s\n\n위 조건에 맞게, 제목과 본문을 작성해줘.\n제품명은 절대 쓰지 말고 카테고리만 언급해." % (
        category, target_concern, product_category or '건강기능식품')
    return system, user

def _build_viral_stage3_prompt(category, target_concern, brand_keyword, product_name, usp, ingredients='', product_category=''):
    system = """너는 네이버 카페에서 활동하는 평범한 회원이다.
카페 침투글(본문 1개)과 댓글 세트(4개)를 함께 작성한다.

[역할]
- 실제 카페 회원처럼 행동하는 사람
- 제품을 "홍보"하는 것이 아니라, 생활 경험을 "나누는" 사람
- 본문 작성자와 댓글 작성자는 각각 다른 사람이다

[본문 구조] — 반드시 아래 순서를 따른다
1. 제목: 감정형 또는 공감형 문장
2. 오프닝: 일상적 상황 + 고민 묘사
3. 중반: 감정 공유 + 기존 시도 방법
4. 후반: 경험 공유형으로 제품/카테고리 자연스럽게 언급
5. 마무리: 질문 또는 공감 유도

[본문 내 제품 언급 규칙]
- 제품명 또는 나만의 키워드를 본문에서 1회만 언급 가능
- 반드시 경험 공유형 톤으로 ("~먹어보고 있는데 괜찮은 것 같아요")
- 홍보 톤 절대 금지 ("추천드려요", "꼭 써보세요")

[댓글 세트 규칙] — 4개 댓글을 아래 흐름으로 작성한다
- 댓글1: 공감 반응 ("저도 그래요ㅠㅠ", "완전 공감이에요")
- 댓글2: 경험 공유 (카테고리 수준 — "저는 천연 성분 들어간 거 먹고 좀 괜찮아졌어요ㅎㅎ")
- 댓글3: 관심/질문 ("그거 혹시 어디서 보셨어요?", "저도 궁금해요")
- 댓글4: 자연스러운 정보 제공 — 여기서만 제품명 또는 나만의 키워드 1회 언급 가능
  (예: "저도 요즘 {제품명} 먹고 있는데 꽤 괜찮아요ㅎㅎ")

[댓글 포인트]
- 모든 댓글은 "공감 → 경험 → 자연스러운 연결" 구조
- 각 댓글은 서로 다른 사람이 쓴 것처럼 말투와 표현이 달라야 한다
- 제품명은 댓글 세트 전체에서 1회만 (댓글4에서만)
- 링크 절대 금지

[톤 & 문체 규칙]
- 존댓말 + 따뜻한 대화체
- 문장 끝: "~했어요", "~하더라구요", "~인 것 같아요" 등 구어체
- 감정 표현: ㅎㅎ, ㅠㅠ, ;;, 👍 중 1~2개 자연스럽게 사용
- 문장 단락: 2~3문장씩 끊기
- 본문 전체 길이: 5줄 이하
- 댓글 각각: 1~2문장

[절대금지]
- 광고 어투: "추천드려요", "꼭 써보세요", "이거 진짜 좋아요"
- 링크, 해시태그, 쇼핑몰 언급
- 반복 문체 (동일 문장 구조 복붙)
- 제품명 과다 노출 (본문+댓글 통틀어 최대 2회)
- CTA (구매 유도, 검색 유도, "한번 찾아보세요")
- 본문 작성자와 댓글 작성자가 같은 말투를 쓰는 것"""
    user = """다음 조건으로 카페 침투글(본문 + 댓글 4개 세트)을 작성해줘.

- 타겟층: %s
- 고민 키워드: %s
- 제품명: %s
- 나만의 키워드: %s
- USP: %s
- 주요 성분: %s
- 제품 카테고리: %s

위 조건에 맞게 본문 1개와 댓글 4개를 작성해줘.
본문에서 제품 언급은 경험형으로 1회만, 댓글에서는 댓글4에서만 1회.""" % (
        category, target_concern, product_name, brand_keyword, usp, ingredients, product_category or '건강기능식품')
    return system, user

@app.post("/api/viral/generate")
async def viral_generate(request: Request):
    """카페바이럴 세트 생성 (SSE)"""
    body = await request.json()
    category = body.get('category', '')
    product = body.get('product', {})
    set_count = body.get('set_count', 3)
    target_concern = product.get('target_concern', '')
    brand_keyword = product.get('brand_keyword', '')
    product_name = product.get('name', '')
    usp = product.get('usp', '')
    ingredients = product.get('ingredients', '')
    product_category = product.get('product_category', '')

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total_steps = set_count * 3

        for s in range(set_count):
            step_base = s * 3

            # 1단계: 일상글
            yield _sse({'type': 'progress', 'msg': '[세트 %d/%d] 1단계 일상글 생성 중...' % (s+1, set_count), 'cur': step_base, 'total': total_steps})
            sys1, usr1 = _build_viral_stage1_prompt(category, product.get('target', ''), '')
            raw1 = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
            s1 = _parse_viral_output(raw1)

            # 2단계: 대화침투글
            yield _sse({'type': 'progress', 'msg': '[세트 %d/%d] 2단계 대화침투글 생성 중...' % (s+1, set_count), 'cur': step_base+1, 'total': total_steps})
            sys2, usr2 = _build_viral_stage2_prompt(category, target_concern, product_category)
            raw2 = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
            s2 = _parse_viral_output(raw2)

            # 3단계: 제품인지글 + 댓글
            yield _sse({'type': 'progress', 'msg': '[세트 %d/%d] 3단계 제품인지글+댓글 생성 중...' % (s+1, set_count), 'cur': step_base+2, 'total': total_steps})
            sys3, usr3 = _build_viral_stage3_prompt(category, target_concern, brand_keyword, product_name, usp, ingredients, product_category)
            raw3 = await loop.run_in_executor(executor, _call_claude, sys3, usr3)
            s3 = _parse_viral_stage3(raw3)

            result = {
                'set_num': s + 1,
                'stage1': s1,
                'stage2': s2,
                'stage3': s3,
            }
            yield _sse({'type': 'result', 'data': result, 'cur': step_base+3, 'total': total_steps})

        yield _sse({'type': 'complete', 'total': set_count})
      except Exception as e:
        print(f"[viral_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'카페바이럴 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

def _parse_viral_output(raw):
    """제목: / 본문: 형식 파싱"""
    title = ''
    body = ''
    lines = raw.strip().split('\n')
    body_lines = []
    in_body = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('제목:') or stripped.startswith('제목 :'):
            title = stripped.split(':', 1)[1].strip()
            in_body = False
        elif stripped.startswith('본문:') or stripped.startswith('본문 :'):
            body_start = stripped.split(':', 1)[1].strip()
            if body_start:
                body_lines.append(body_start)
            in_body = True
        elif in_body:
            body_lines.append(line)
        elif not title:
            title = stripped
            in_body = True
        else:
            body_lines.append(line)
    body = '\n'.join(body_lines).strip()
    if not body and not title:
        title = lines[0].strip() if lines else ''
        body = '\n'.join(lines[1:]).strip()
    return {'title': title, 'body': body}

def _parse_viral_stage3(raw):
    """3단계: 글(제목+본문) + 댓글 파싱"""
    result = {'title': '', 'body': '', 'comments': ''}
    parts = re.split(r'\[댓글\]|\[글\]', raw, flags=re.IGNORECASE)
    if len(parts) >= 2:
        post_part = parts[1] if len(parts) >= 3 else parts[0]
        comment_part = parts[-1] if len(parts) >= 2 else ''
        if len(parts) >= 3:
            post_part = parts[1]
            comment_part = parts[2]
        else:
            post_part = parts[0]
            comment_part = parts[1]
        parsed = _parse_viral_output(post_part)
        result['title'] = parsed['title']
        result['body'] = parsed['body']
        result['comments'] = comment_part.strip()
    else:
        parsed = _parse_viral_output(raw)
        result['title'] = parsed['title']
        result['body'] = parsed['body']
    return result

@app.post("/api/viral/save-notion")
async def viral_save_notion(request: Request):
    """카페바이럴 단계별 노션 저장"""
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '카페'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('body_summary'):
        props['본문'] = {'rich_text': [{'text': {'content': body['body_summary'][:2000]}}]}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}

    content_text = body.get('body', '')
    if body.get('comments'):
        content_text += '\n\n---\n댓글:\n' + body['comments']
    if content_text:
        children = []
        for para in [p.strip() for p in content_text.split('\n\n') if p.strip()][:100]:
            for k in range(0, len(para), 2000):
                children.append({
                    'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}
                })
        payload['children'] = children[:100]

    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ───────────────────────────── JISIKIN ─────────────────────────────

def _build_jisikin_title_prompt(keyword, product):
    system = """너는 클릭을 유도하는 지식인 질문 제목을 작성하는 마케팅 전문가다.

조건:
- 질문자 입장에서 실제로 궁금해서 작성한 것처럼 보여야 함
- 공백 포함 25자 이내
- 브랜드명은 절대 언급하지 마라
- 광고 티는 나면 안 됨
- 사용자들이 "어? 나도 이런데?" 라고 느낄 수 있는 제목을 작성하라
- 상위 노출 키워드는 제목에 무조건 1회 자연스럽게 들어가야 함

위 조건을 바탕으로, 사람들이 네이버 지식인에서 검색하고 싶은 제목을 한 줄로 만들어줘.
제목 1개만 출력하고, 다른 설명은 쓰지 마라.

예시:
- 가르시니아 먹어도 살이 안 빠져요 제발 급해요
- 식욕 억제제 부작용 너무 심한데 괜찮나요 ㅠㅠ
- 가르시니아 다이어트 보조제 효과 있나요 광고사절
- 다이어트 약 먹었더니 속 울렁거려요 미쳐요"""
    user = "[상위 노출 키워드]\n%s\n\n[제품 정보 요약]\n제품명: %s\nUSP: %s\n타겟층: %s\n주요 성분: %s" % (
        keyword, product.get('name',''), product.get('usp',''),
        product.get('target',''), product.get('ingredients',''))
    return system, user

def _build_jisikin_body_prompt(keyword, product):
    system = """너는 실제 소비자가 네이버 지식인에 질문을 올리는 것처럼 보이는 마케팅 카피를 작성하는 전문가다.

네이버 지식인에서 사람들이 자주 검색할 법한 진짜 사용자처럼 보이는 질문 본문을 1개만 작성하세요.

조건:
- 반드시 실제 고민처럼 보여야 한다 (광고 느낌 ❌)
- 질문자 입장에서 말하듯이 작성할 것
- 너무 매끄럽거나 인위적인 문장은 ❌
- 총 6~8문장으로 구성하되, 길게 느껴지지 않게
- 마지막 문장은 반드시 "~인가요?", "~있을까요?" 식의 질문형으로 끝날 것
- 상위 노출 키워드는 질문 본문 안에 정확히 한 번만 자연스럽게 포함될 것
- 브랜드명은 절대 사용하지 말 것
- 1인칭 시점, 구어체, 이모티콘 가끔 사용
- 다른 제품을 써본 경험 + 불만족 포인트 포함하면 설득력 ↑"""
    user = "[상위 노출 키워드]\n%s\n\n[제품 정보 요약]\n제품명: %s\nUSP: %s\n타겟층: %s\n주요 성분: %s" % (
        keyword, product.get('name',''), product.get('usp',''),
        product.get('target',''), product.get('ingredients',''))
    return system, user

def _build_jisikin_answers_prompt(keyword, question_title, question_body, product):
    brand_kw = product.get('brand_keyword', '')
    system = """너는 네이버 지식인에 올라온 질문에 대해
실제 사용자가 남긴 것처럼 보이는 두 개의 신뢰도 높은 답변을 작성하는 마케팅 카피 전문가다.

답변 목적:
- 광고처럼 보이지 않으면서도,
- 실제 사용 경험 기반으로 문제 공감 → 해결 과정 → 제품 유도(나만의 키워드)

작성 조건:

1. 답변은 반드시 두 개 작성
   → 답변 1, 답변 2는 완전히 독립된 사람의 글처럼 보여야 함
   → 내용 반복 ❌
   → 동일한 문체 반복 금지 (스팸 인식됨)

2. 각 답변에는 반드시 나만의 키워드가 자연스럽게 1회 포함되어야 함
   (설명: 광고 티가 안나게 하는 우회 키워드)

3. 각 답변의 길이는 최소 10줄 이상
   (너무 짧으면 신뢰도가 떨어짐)

4. 각 답변은 다음 흐름을 따라야 함:

---

공감
- 질문자의 고민에 "저도 그랬어요", "~~ 너무 고민이었어요" 등으로 진심으로 공감
- 광고 같은 말투 ❌, 진짜 써본 사람처럼 말하듯이 작성

경험 + 해결 과정 공유
- 본인이 어떤 방법을 시도했는지
- 기존 제품에서 실패한 경험
- 현재 사용하는 제품에서 느낀 효과
- 성분/기능을 언급하되, 전문가처럼 말하지 말 것

제품 유도 (자연스럽게)
- 브랜드명은 쓰지 말고,
- 나만의 키워드를 "요즘은 그냥 OOO 위주로 써요" 식으로 한 번만 써줄 것

---

✅ [답변 1 템플릿]
1. 질문자의 문제에 "저도 그랬어요", "한동안 저도 그게 심했어요" 식으로 시작
2. 그로 인해 겪은 스트레스나 불편함을 설명
3. 처음에는 어떤 걸 시도했는지 말함
4. 어떤 이유로 현재 제품/방법을 찾게 됐는지 말함
5. 사용 방식, 시간, 사용감, 변화된 점 등을 구체적으로 설명
6. 기대감과 불신 사이에서 느낀 점 언급 (예: "광고 같아서 반신반의했는데...")
7. 비교 경험 (예: "이전엔 이런 성분 썼는데...")
8. 결과적으로 증상이 어떻게 나아졌는지 언급
9. 아직 완벽하진 않지만 예전보단 낫다고 마무리
10. 마지막 줄에 나만의 키워드를 포함해 자연스럽게 마무리
    - 예: "요즘은 그냥 {나만의 키워드}로 정착했어요."

---

✅ [답변 2 템플릿]
1. 질문과 비슷한 경험이 있었다는 식으로 시작
2. 상황을 구체적으로 설명 (예: 계절/장소/스트레스 등 상황성 강조)
3. 처음 썼던 제품이나 방법이 잘 안 먹혔다는 경험
4. 제품을 선택하기까지의 고민이나 비교 과정 묘사
5. 자신이 중요하게 본 조건 (예: 성분, 사용감, 가격 등) 언급
6. 쓰고 나서의 변화 (즉각적인 효과/점진적 변화 등)
7. 개인적인 습관이나 팁 (예: 사용하는 시간대, 보조 루틴)
8. 예전엔 상상도 못 했던 변화를 간접적으로 언급
9. 마무리는 "요즘은 ○○만 씀" 식으로 경험 강조
10. 나만의 키워드를 사용해 자연스럽게 닫는다
    - 예: "그래서 지금은 {나만의 키워드}만 계속 쓰고 있어요."

---

톤과 스타일:
- 말하듯 자연스럽고, 실제 사용 경험을 공유하는 후기 느낌
- 전문가처럼 설명하지 않고, 생활 속 체험담을 강조
- 광고 문구, 과장된 표현, 브랜드명 절대 금지

출력 형식:
답변 1:
(내용)

답변 2:
(내용)"""
    user = "[상위 노출 키워드]\n%s\n\n[제품 정보 요약]\n제품명: %s\nUSP: %s\n타겟층: %s\n주요 성분: %s\n\n[나만의 키워드]\n%s" % (
        keyword, product.get('name',''), product.get('usp',''),
        product.get('target',''), product.get('ingredients',''), brand_kw)
    return system, user

def _build_jisikin_direct_answer_prompt(question_text, keyword, product):
    """직접 답변 모드: 실제 고민글에 대한 답변 1개 생성"""
    brand_kw = product.get('brand_keyword', '')
    system = """역할: 네이버 지식인에 올라온 실제 질문에 대해, 진짜 사용자가 남긴 것처럼 보이는 신뢰도 높은 답변 1개를 작성하는 전문가.

목표: 실제 질문자의 구체적인 상황에 맞춰 공감하고, 본인의 경험을 공유하듯 자연스럽게 답변한다.

핵심 원칙:

[1. 맞춤 대응]
- 질문자가 언급한 구체적인 상황(증상, 고민, 시도한 방법 등)에 직접 반응하라
- 일반적인 답변이 아니라 해당 질문에만 해당하는 맞춤 답변이어야 한다
- 질문에서 언급한 키워드/상황을 답변에서 다시 언급하여 실제 대화처럼 보이게 하라

[2. 솔루션 섞기]
- 우리 제품(나만의 키워드)만 단독 추천하지 마라
- 다른 방법(운동, 식단, 생활습관 등)이나 다른 일반적인 제품/성분도 함께 자연스럽게 섞어서 언급하라
- 나만의 키워드를 특별히 강조하지 말고 여러 솔루션 중 하나로 자연스럽게 녹여라
- 여러 가지를 시도해본 사람처럼 보여야 한다

[3. 답변 구조]
- 공감: 질문자의 고민에 진심으로 공감 ("저도 그랬어요" 등)
- 경험 + 해결 과정: 본인이 시도한 여러 방법, 성공/실패 경험
- 솔루션들: 다양한 방법을 언급하며 나만의 키워드를 그 중 하나로 자연스럽게 포함

[4. 톤앤매너]
- 말하듯 자연스럽고, 실제 사용 경험을 공유하는 후기 느낌
- 전문가처럼 설명하지 않고, 생활 속 체험담 강조
- 답변 최소 10줄 이상

[5. 나만의 키워드]
- 답변 안에 나만의 키워드를 자연스럽게 1회만 포함
- 브랜드명 직접 언급 금지
- 강조하거나 추천하는 톤 금지 — 그냥 "나는 이것도 써봤다" 수준

출력: 답변 1개만 출력 (부연 설명 없이)

금지사항:
- 광고 문구, 과장된 표현 금지
- 브랜드명 직접 언급 금지
- 나만의 키워드 2회 이상 사용 금지
- 링크 삽입 금지
- "꼭 써보세요", "추천드려요" 같은 직접 추천 금지"""
    user = "실제 고민글:\n%s\n\n상위 노출 키워드: %s\n제품 정보: %s, %s, %s, %s\n나만의 키워드: %s" % (
        question_text, keyword,
        product.get('name',''), product.get('ingredients',''),
        product.get('usp',''), product.get('target',''), brand_kw)
    return system, user

def _parse_jisikin_answers(raw):
    """답변 1, 답변 2 분리"""
    parts = re.split(r'✅\s*\[답변\s*[12]\]|답변\s*[12]\s*:', raw)
    answer1 = parts[1].strip() if len(parts) > 1 else ''
    answer2 = parts[2].strip() if len(parts) > 2 else ''
    if not answer1 and not answer2:
        half = len(raw) // 2
        answer1 = raw[:half].strip()
        answer2 = raw[half:].strip()
    return answer1, answer2

@app.get("/api/jisikin/notion-keywords")
async def jisikin_notion_keywords():
    """노션 키워드 DB에서 지식인 배정 키워드 조회"""
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    # 지식인SEO 또는 지식인바이럴
    results_all = []
    for channel_name in ['지식인', '지식인SEO', '지식인바이럴']:
        payload = {
            'filter': {
                'and': [
                    {'property': '배정 채널', 'multi_select': {'contains': channel_name}},
                    {'property': '상태', 'select': {'equals': '미사용'}},
                ]
            },
            'page_size': 100,
        }
        try:
            r = req.post('https://api.notion.com/v1/databases/%s/query' % KEYWORD_DB_ID, headers=headers, json=payload, timeout=15)
            if r.status_code == 200:
                for page in r.json().get('results', []):
                    props = page.get('properties', {})
                    title_prop = props.get('키워드', {}).get('title', [])
                    kw = title_prop[0]['text']['content'] if title_prop else ''
                    pid = page['id']
                    if kw and not any(x['page_id'] == pid for x in results_all):
                        results_all.append({'keyword': kw, 'page_id': pid})
        except Exception as e:
            print(f"[jisikin] notion query error: {e}")
    return {'keywords': results_all}

@app.post("/api/jisikin/generate")
async def jisikin_generate(request: Request):
    """지식인 질문+답변 생성 (SSE)"""
    body = await request.json()
    keywords = body.get('keywords', [])
    product = body.get('product', {})

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords)

        for i, kw_data in enumerate(keywords):
            kw = kw_data['keyword']

            # STEP 1: 질문 제목
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 질문 제목 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys1, usr1 = _build_jisikin_title_prompt(kw, product)
            q_title = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
            q_title = q_title.strip().split('\n')[0].strip().strip('"').strip()

            # STEP 2: 질문 본문
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 질문 본문 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys2, usr2 = _build_jisikin_body_prompt(kw, product)
            q_body = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
            q_body = q_body.strip()

            # STEP 3: 답변 2개
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 답변 2개 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys3, usr3 = _build_jisikin_answers_prompt(kw, q_title, q_body, product)
            raw_answers = await loop.run_in_executor(executor, _call_claude, sys3, usr3)
            answer1, answer2 = _parse_jisikin_answers(raw_answers)

            result = {
                'keyword': kw, 'q_title': q_title, 'q_body': q_body,
                'answer1': answer1, 'answer2': answer2,
                'page_id': kw_data.get('page_id', ''),
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[jisikin_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'지식인 콘텐츠 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/jisikin/generate-direct")
async def jisikin_generate_direct(request: Request):
    """지식인 직접 답변 생성 (SSE) — 실제 고민글에 답변"""
    body = await request.json()
    questions = body.get('questions', [])
    product = body.get('product', {})

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(questions)

        for i, q in enumerate(questions):
            question_text = q.get('text', '')
            keyword = q.get('keyword', '')

            yield _sse({'type': 'progress', 'msg': '[%d/%d] 답변 생성 중...' % (i+1, total), 'cur': i, 'total': total})
            sys_p, usr_p = _build_jisikin_direct_answer_prompt(question_text, keyword, product)
            answer = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
            answer = answer.strip()

            result = {
                'question_text': question_text,
                'keyword': keyword,
                'answer': answer,
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[jisikin_generate_direct] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'지식인 직접 답변 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/jisikin/save-notion")
async def jisikin_save_notion(request: Request):
    """지식인 콘텐츠 노션 저장"""
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('q_title', '')}}]},
        '채널': {'select': {'name': '지식인'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('q_body'):
        props['본문'] = {'rich_text': [{'text': {'content': body['q_body'][:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}

    content = body.get('q_body', '')
    answers_text = ''
    if body.get('answer1'):
        answers_text += '✅ [답변 1]\n' + body['answer1']
    if body.get('answer2'):
        answers_text += '\n\n✅ [답변 2]\n' + body['answer2']
    full_text = content + '\n\n---\n' + answers_text if answers_text else content

    children = []
    for para in [p.strip() for p in full_text.split('\n\n') if p.strip()][:100]:
        for k in range(0, len(para), 2000):
            children.append({
                'object': 'block', 'type': 'paragraph',
                'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}
            })
    payload['children'] = children[:100]

    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ───────────────────────────── YOUTUBE VIDEO SEARCH ─────────────────────────

@app.post("/api/youtube/search-videos")
async def youtube_search_videos(request: Request):
    """키워드로 YouTube 영상 검색 (yt-dlp)"""
    body = await request.json()
    keyword = body.get('keyword', '').strip()
    count = min(int(body.get('count', 50)), 100)
    if not keyword:
        return JSONResponse({'error': '키워드를 입력하세요'}, 400)

    def _search():
        try:
            import yt_dlp
            ydl_opts = {
                'quiet': True,
                'extract_flat': True,
                'force_generic_extractor': False,
            }
            videos = []
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                result = ydl.extract_info(f'ytsearch{count}:{keyword}', download=False)
                for entry in result.get('entries', []):
                    if not entry:
                        continue
                    vid = {
                        'id': entry.get('id', ''),
                        'title': entry.get('title', ''),
                        'url': entry.get('url', f"https://www.youtube.com/watch?v={entry.get('id','')}"),
                        'channel': entry.get('channel', entry.get('uploader', '')),
                        'duration': entry.get('duration'),
                        'view_count': entry.get('view_count'),
                    }
                    videos.append(vid)
            return {'videos': videos, 'keyword': keyword, 'total': len(videos)}
        except ImportError:
            return {'error': 'yt-dlp가 설치되지 않았습니다. pip install yt-dlp'}
        except Exception as e:
            return {'error': str(e)}

    result = await asyncio.get_running_loop().run_in_executor(executor, _search)
    if 'error' in result:
        return JSONResponse(result, 500)
    return result

@app.post("/api/youtube/fetch-video-details")
async def youtube_fetch_video_details(request: Request):
    """영상 URL 목록에서 제목/설명을 일괄 크롤링"""
    body = await request.json()
    urls = body.get('urls', [])
    if not urls:
        return JSONResponse({'error': 'URL 목록이 비어있습니다'}, 400)

    def _fetch_all():
        results = []
        headers = {'User-Agent': 'Mozilla/5.0', 'Accept-Language': 'ko-KR,ko;q=0.9'}
        for url in urls[:50]:
            vid = {'url': url, 'title': '', 'description': ''}
            try:
                r = req.get(url, headers=headers, timeout=10)
                title_match = re.search(r'<title>(.*?)</title>', r.text)
                if title_match:
                    vid['title'] = title_match.group(1).replace(' - YouTube', '').strip()
                desc_match = re.search(r'"shortDescription":"(.*?)"', r.text)
                if desc_match:
                    vid['description'] = desc_match.group(1).replace('\\n', '\n').strip()[:2000]
            except Exception:
                pass
            results.append(vid)
        return results

    results = await asyncio.get_running_loop().run_in_executor(executor, _fetch_all)
    return {'videos': results}

# ───────────────────────────── YOUTUBE COMMENTS ─────────────────────────────

def _build_youtube_summary_prompt(video_title, script, description):
    """1단계: 영상 요약 (temperature 0.3 — 정확하게)"""
    context = script if script else description if description else video_title
    system = """너는 유튜브 영상의 핵심 내용을 요약하는 AI야.
반드시 한국어로 답변하며, 불필요한 내용은 제외해야 해."""
    user = """영상에 대한 정보는 아래 두 가지를 참고해:
1) 영상 제목: %s
2) 영상 소개(스크립트 또는 더보기): %s

위의 두 가지 정보를 함께 분석하여 영상의 주요 내용을 핵심만 500자 내외로 요약해줘!""" % (video_title, context[:3000])
    return system, user


def _build_youtube_comment_prompt(video_title, description, brand_keyword):
    """유튜브 바이럴 댓글 3단 시나리오 생성 (temperature 0.9 — 창의적으로)"""
    system = """너는 유튜브 영상에 달리는 바이럴 댓글을 작성하는 마케팅 전문가다.

주어진 영상 제목과 더보기 내용을 참고하여,
실제 시청자가 남긴 것처럼 보이는 바이럴 댓글 3단 시나리오를 만들어라.

[3단 댓글 시나리오 구조]

1단계 — 밑밥 (공감형 질문):
- 영상 내용에 공감하면서, 나만의 키워드와 관련된 질문을 던지는 댓글
- 진짜 시청자가 궁금해서 남긴 것처럼 자연스럽게

2단계 — 해결사 (자연스러운 키워드 삽입):
- 1단계 댓글에 대한 답변 형태로, 나만의 키워드를 자연스럽게 추천하는 댓글
- 경험 기반으로 "나는 이거 써봤는데~" 식의 톤

3단계 — 쐐기 (구매 행동 유도):
- 2단계를 보고 "나도 해봐야겠다" 식으로 반응하는 댓글
- 구매 의향을 간접적으로 드러내는 톤

[작성 규칙]
- 광고 티 절대 금지
- 브랜드명 직접 언급 금지
- 실제 유튜브 댓글처럼 짧고 자연스러운 구어체
- 각 단계별 댓글 1개씩, 총 3개 출력

[출력 형식]
댓글1 (밑밥):
(내용)

댓글2 (해결사):
(내용)

댓글3 (쐐기):
(내용)"""
    user = """[영상 제목]
%s

[더보기 내용]
%s

[나만의 키워드]
%s""" % (video_title, description, brand_keyword)
    return system, user

@app.post("/api/youtube/fetch-info")
async def youtube_fetch_info(request: Request):
    """YouTube URL에서 제목/설명/자막 크롤링"""
    body = await request.json()
    url = body.get('url', '')
    if not url:
        return JSONResponse({'error': 'URL 필요'}, 400)
    result = {'title': '', 'description': '', 'transcript': ''}
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        r = req.get(url, headers=headers, timeout=10)
        # 제목 추출
        title_match = re.search(r'<title>(.*?)</title>', r.text)
        if title_match:
            result['title'] = title_match.group(1).replace(' - YouTube', '').strip()
        # 설명 추출
        desc_match = re.search(r'"shortDescription":"(.*?)"', r.text)
        if desc_match:
            result['description'] = desc_match.group(1).replace('\\n', '\n').strip()[:2000]
    except Exception:
        pass
    return result

@app.post("/api/youtube/generate")
async def youtube_generate(request: Request):
    """유튜브 댓글 생성 (SSE)"""
    body = await request.json()
    videos = body.get('videos', [])
    product_name = body.get('product_name', '')
    brand_keyword = body.get('brand_keyword', product_name)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(videos)
        for i, v in enumerate(videos):
            title = v.get('title', '')
            description = v.get('description', '')
            # 1단계: 영상 요약 (temperature 0.3 — 참고용)
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 영상 요약 중...' % (i+1, total, title[:30]), 'cur': i, 'total': total})
            sum_sys, sum_usr = _build_youtube_summary_prompt(title, v.get('script', ''), description)
            summary = await loop.run_in_executor(executor, _call_claude, sum_sys, sum_usr, 0.3, 1024)
            summary = summary.strip()
            # 2단계: 3단 시나리오 댓글 생성 (temperature 0.9 — 창의적으로)
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 댓글 3단 시나리오 생성 중...' % (i+1, total, title[:30]), 'cur': i, 'total': total})
            cmt_sys, cmt_usr = _build_youtube_comment_prompt(title, description, brand_keyword)
            comment = await loop.run_in_executor(executor, _call_claude, cmt_sys, cmt_usr, 0.9, 1024)
            comment = comment.strip()
            result = {
                'title': title, 'link': v.get('link', ''),
                'script': v.get('script', ''), 'description': description,
                'summary': summary, 'comment': comment,
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})
        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[youtube_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'유튜브 댓글 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/youtube/save-notion")
async def youtube_save_notion(request: Request):
    """유튜브 댓글 노션 저장"""
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '유튜브'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('comment'):
        props['본문'] = {'rich_text': [{'text': {'content': body['comment'][:2000]}}]}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    children = []
    if body.get('comment'):
        children.append({'object': 'block', 'type': 'paragraph',
            'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': body['comment'][:2000]}}]}})
    if children:
        payload['children'] = children

    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ──────────────────── YOUTUBE AUTO-POST (CommentBoost 통합) ──────────────────

# 자동게시 상태
_yt_autopost_state = {
    "running": False,
    "current_task": None,
    "progress": 0,
    "total": 0,
    "logs": [],
    "results": {"success": 0, "fail": 0, "skip": 0},
}
_yt_autopost_lock = threading.Lock()

# 모듈 임포트 (src/)
try:
    from src.youtube_bot import YouTubeBot
    from src.fingerprint import FingerprintManager
    from src.safety_rules import SafetyRules
    from src.smm_client import SMMClient
    from src.comment_tracker import CommentTracker
    _yt_modules_available = True
except ImportError as _ie:
    print(f"[CommentBoost] 모듈 임포트 실패: {_ie}")
    _yt_modules_available = False

_yt_safety_rules = SafetyRules() if _yt_modules_available else None
_yt_comment_tracker = CommentTracker() if _yt_modules_available else None
_yt_fingerprint_mgr = FingerprintManager() if _yt_modules_available else None
_yt_smm_client = SMMClient() if _yt_modules_available else None

# YouTube 계정 관리
_YT_ACCOUNTS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config', 'yt_accounts.json')
os.makedirs(os.path.dirname(_YT_ACCOUNTS_FILE), exist_ok=True)

def _load_yt_accounts():
    if os.path.exists(_YT_ACCOUNTS_FILE):
        try:
            return json.loads(open(_YT_ACCOUNTS_FILE, encoding='utf-8').read())
        except (json.JSONDecodeError, OSError) as e:
            print(f"[yt_accounts] 로드 오류: {e}")
    return []

def _save_yt_accounts(accounts):
    tmp = _YT_ACCOUNTS_FILE + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(accounts, f, ensure_ascii=False, indent=2)
    os.replace(tmp, _YT_ACCOUNTS_FILE)

def _yt_add_log(msg, level="info"):
    import datetime as _dt
    with _yt_autopost_lock:
        _yt_autopost_state["logs"].append({
            "time": _dt.datetime.now().strftime("%H:%M:%S"),
            "msg": msg, "level": level,
        })
        if len(_yt_autopost_state["logs"]) > 200:
            _yt_autopost_state["logs"] = _yt_autopost_state["logs"][-100:]

def _yt_update_state(**kwargs):
    """thread-safe 상태 업데이트"""
    with _yt_autopost_lock:
        _yt_autopost_state.update(kwargs)

def _yt_get_state(*keys):
    """thread-safe 상태 읽기"""
    with _yt_autopost_lock:
        if len(keys) == 1:
            return _yt_autopost_state.get(keys[0])
        return {k: _yt_autopost_state.get(k) for k in keys}

@app.get("/api/youtube/accounts")
async def yt_get_accounts():
    """YouTube 계정 목록 조회"""
    accounts = _load_yt_accounts()
    result = []
    for acc in accounts:
        label = acc.get("label", acc.get("email", "unknown"))
        status = _yt_safety_rules.get_account_status(label) if _yt_safety_rules else {}
        result.append({**acc, "password": "***", **status})
    return {"accounts": result}

@app.post("/api/youtube/accounts")
async def yt_add_account(request: Request):
    """YouTube 계정 추가"""
    body = await request.json()
    email = body.get("email", "").strip()
    label = body.get("label", email)
    if not email:
        return JSONResponse({"error": "이메일 필요"}, 400)

    accounts = _load_yt_accounts()
    for acc in accounts:
        if acc.get("email") == email:
            return JSONResponse({"error": "이미 등록된 계정"}, 400)

    accounts.append({
        "email": email, "label": label,
        "password": body.get("password", ""),
        "proxy": body.get("proxy", ""),
        "adspower_profile_id": body.get("adspower_profile_id", ""),
        "active": True,
        "status": "활성",  # 활성/휴식/정지/폐기
        "aging_status": body.get("aging_status", "완료"),  # 진행중/완료
        "created_at": datetime.now().strftime("%Y-%m-%d"),
        "last_used_at": None,
        "total_comments": 0,
        "notes": body.get("notes", ""),
    })
    _save_yt_accounts(accounts)
    return {"success": True}

@app.delete("/api/youtube/accounts/{email}")
async def yt_delete_account(email: str):
    """YouTube 계정 삭제"""
    accounts = _load_yt_accounts()
    accounts = [a for a in accounts if a.get("email") != email]
    _save_yt_accounts(accounts)
    return {"success": True}

@app.patch("/api/youtube/accounts/{email}")
async def yt_update_account(email: str, request: Request):
    """YouTube 계정 정보 수정 (상태, 메모, 프록시 등)"""
    body = await request.json()
    accounts = _load_yt_accounts()
    for acc in accounts:
        if acc.get("email") == email:
            for key in ["status", "aging_status", "notes", "proxy", "adspower_profile_id", "active"]:
                if key in body:
                    acc[key] = body[key]
            _save_yt_accounts(accounts)
            return {"success": True}
    return JSONResponse({"error": "계정을 찾을 수 없습니다."}, 404)

@app.post("/api/youtube/test-login")
async def yt_test_login(request: Request):
    """YouTube 계정 로그인 테스트"""
    if not _yt_modules_available:
        return JSONResponse({"error": "playwright 모듈 미설치. pip install playwright && playwright install chromium"}, 500)

    body = await request.json()
    email = body.get("email", "")
    accounts = _load_yt_accounts()
    account = next((a for a in accounts if a.get("email") == email), None)
    if not account:
        return JSONResponse({"error": "계정을 찾을 수 없습니다."}, 404)

    label = account.get("label", email)

    def _test():
        bot = YouTubeBot(account_label=label, fingerprint_manager=_yt_fingerprint_mgr)
        bot.headless = True
        try:
            bot.start_browser(account)
            logged_in = bot.login_youtube(account)
            return {"logged_in": logged_in, "label": label}
        except Exception as e:
            return {"logged_in": False, "error": str(e)}
        finally:
            bot.close()

    result = await asyncio.get_running_loop().run_in_executor(executor, _test)
    return result

@app.post("/api/youtube/manual-login")
async def yt_manual_login(request: Request):
    """YouTube 수동 로그인 (브라우저 열기)"""
    if not _yt_modules_available:
        return JSONResponse({"error": "playwright 모듈 미설치"}, 500)

    body = await request.json()
    email = body.get("email", "")
    accounts = _load_yt_accounts()
    account = next((a for a in accounts if a.get("email") == email), None)
    if not account:
        return JSONResponse({"error": "계정을 찾을 수 없습니다."}, 404)

    label = account.get("label", email)

    def _login():
        bot = YouTubeBot(account_label=label, fingerprint_manager=_yt_fingerprint_mgr)
        bot.headless = False  # 화면 보이기 필수
        try:
            bot.start_browser(account)
            success = bot.manual_login(timeout_sec=120)
            return {"success": success, "label": label}
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            bot.close()

    result = await asyncio.get_running_loop().run_in_executor(executor, _login)
    return result

@app.post("/api/youtube/auto-post")
async def yt_auto_post(request: Request):
    """YouTube 댓글 자동 게시 시작"""
    if not _yt_modules_available:
        return JSONResponse({"error": "playwright 모듈 미설치"}, 500)

    if _yt_get_state("running"):
        return JSONResponse({"error": "이미 실행 중입니다."}, 400)

    body = await request.json()
    tasks = body.get("tasks", [])  # [{youtube_url, comment_text, page_id?}, ...]
    headless = body.get("headless", True)

    if not tasks:
        return JSONResponse({"error": "작업 목록이 비어있습니다."}, 400)

    accounts = _load_yt_accounts()
    active_accounts = [a for a in accounts if a.get("active", True)]
    if not active_accounts:
        return JSONResponse({"error": "활성 계정이 없습니다."}, 400)

    def _run():
        with _yt_autopost_lock:
            _yt_autopost_state.update({
                "running": True, "progress": 0, "total": len(tasks),
                "logs": [], "results": {"success": 0, "fail": 0, "skip": 0},
                "current_task": None,
            })

        account_idx = 0
        bot = None
        current_label = ""

        try:
            for i, task in enumerate(tasks):
                if not _yt_get_state("running"):
                    _yt_add_log("사용자에 의해 중지됨", "warning")
                    break

                youtube_url = task.get("youtube_url", "")
                comment_text = task.get("comment_text", "")
                url_short = youtube_url[:50] + "..." if len(youtube_url) > 50 else youtube_url

                _yt_update_state(progress=i, current_task=f"[{i+1}/{len(tasks)}] {url_short}")

                # 계정 라운드로빈
                account = active_accounts[account_idx % len(active_accounts)]
                label = account.get("label", account.get("email", "unknown"))

                # 계정 변경 시 브라우저 재시작
                if label != current_label:
                    if bot:
                        bot.close()
                    current_label = label
                    antidetect = body.get("antidetect_mode", "stealth")
                    bot = YouTubeBot(
                        account_label=label,
                        fingerprint_manager=_yt_fingerprint_mgr,
                        antidetect_mode=antidetect,
                    )
                    bot.headless = headless
                    try:
                        bot.start_browser(account)
                        logged_in = bot.login_youtube(account)
                    except Exception as bot_err:
                        _yt_add_log(f"[{label}] 브라우저 시작 실패: {bot_err}", "error")
                        try: bot.close()
                        except Exception: pass
                        bot = None
                        with _yt_autopost_lock:
                            _yt_autopost_state["results"]["fail"] += 1
                        account_idx += 1
                        continue
                    if not logged_in:
                        _yt_add_log(f"[{label}] 로그인 실패 — 건너뜀", "error")
                        try: bot.close()
                        except Exception: pass
                        bot = None
                        with _yt_autopost_lock:
                            _yt_autopost_state["results"]["fail"] += 1
                        account_idx += 1
                        continue

                # 안전 규칙 검사
                passed, reason = _yt_safety_rules.check_all_rules(
                    current_label, youtube_url, comment_text, skip_interval=False
                )
                if not passed:
                    _yt_add_log(f"[건너뜀] {reason}", "warning")
                    with _yt_autopost_lock:
                        _yt_autopost_state["results"]["skip"] += 1
                    continue

                # 인간형 딜레이
                if i > 0:
                    delay_info = _yt_safety_rules.get_human_delay("comment")
                    _yt_add_log(f"🧑 {delay_info['description']}", "info")
                    time.sleep(delay_info["delay_sec"])

                # 댓글 작성
                _yt_add_log(f"[댓글 작성 중] {url_short}", "info")
                comment_url = bot.post_comment(youtube_url, comment_text)

                if comment_url:
                    _yt_autopost_state["results"]["success"] += 1
                    _yt_safety_rules.record_comment(current_label, youtube_url, comment_text)
                    _yt_add_log(f"[성공] {comment_url}", "success")

                    # 계정 사용 기록 업데이트
                    try:
                        _accs = _load_yt_accounts()
                        for _a in _accs:
                            if _a.get("label") == current_label or _a.get("email") == account.get("email"):
                                _a["last_used_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                                _a["total_comments"] = _a.get("total_comments", 0) + 1
                                break
                        _save_yt_accounts(_accs)
                    except Exception:
                        pass

                    # 트래킹 등록
                    if _yt_comment_tracker:
                        _yt_comment_tracker.register_comment(
                            comment_url=comment_url,
                            video_url=youtube_url,
                            comment_text=comment_text,
                            account_label=current_label,
                        )

                    # SMM 좋아요 주문 대기 등록 (auto_like 모드)
                    if body.get("auto_like") and _yt_smm_client and _yt_smm_client.enabled and comment_url:
                        like_qty = body.get("like_quantity", _yt_smm_client.default_quantity)
                        with _yt_autopost_lock:
                            _yt_autopost_state.setdefault("pending_likes", []).append({
                                "comment_url": comment_url,
                                "quantity": like_qty,
                                "video_title": url_short,
                                "status": "pending_approval",
                            })
                        _yt_add_log(f"👍 좋아요 주문 대기: {like_qty}개 ({url_short})", "info")

                    # Notion 업데이트 (page_id가 있으면)
                    page_id = task.get("page_id")
                    if page_id and NOTION_TOKEN:
                        try:
                            _headers_n = {
                                'Authorization': f'Bearer {NOTION_TOKEN}',
                                'Content-Type': 'application/json',
                                'Notion-Version': '2022-06-28',
                            }
                            _payload = {"properties": {
                                "댓글 url": {"url": comment_url},
                                "상태": {"select": {"name": "댓글완료"}},
                            }}
                            req.patch(f'https://api.notion.com/v1/pages/{page_id}',
                                     headers=_headers_n, json=_payload, timeout=15)
                        except Exception:
                            pass
                else:
                    with _yt_autopost_lock:
                        _yt_autopost_state["results"]["fail"] += 1
                    _yt_add_log(f"[실패] 댓글 작성 실패", "error")

                with _yt_autopost_lock:
                    _yt_autopost_state["progress"] = i + 1

        except Exception as e:
            _yt_add_log(f"[에러] {str(e)}", "error")
        finally:
            if bot:
                bot.close()
            _yt_update_state(running=False, current_task=None)
            _yt_add_log("자동 게시 완료", "info")

    threading.Thread(target=_run, daemon=True).start()
    return {"success": True, "total": len(tasks)}

@app.get("/api/youtube/auto-post/status")
async def yt_auto_post_status():
    """자동 게시 진행 상태"""
    with _yt_autopost_lock:
        return {
            "running": _yt_autopost_state["running"],
            "progress": _yt_autopost_state["progress"],
            "total": _yt_autopost_state["total"],
            "current_task": _yt_autopost_state["current_task"],
            "results": dict(_yt_autopost_state["results"]),
            "logs": list(_yt_autopost_state["logs"][-50:]),
        }

@app.post("/api/youtube/auto-post/stop")
async def yt_auto_post_stop():
    """자동 게시 중지"""
    _yt_update_state(running=False)
    return {"success": True}

@app.get("/api/youtube/tracking/summary")
async def yt_tracking_summary():
    """댓글 트래킹 요약"""
    if not _yt_comment_tracker:
        return {"total": 0, "active": 0, "hidden": 0, "deleted": 0, "total_likes": 0}
    return _yt_comment_tracker.get_summary()

@app.get("/api/youtube/safety/status")
async def yt_safety_status():
    """안전 규칙 상태 (전체 계정)"""
    if not _yt_safety_rules:
        return {"today_total": 0, "accounts": []}

    accounts = _load_yt_accounts()
    statuses = []
    for acc in accounts:
        label = acc.get("label", acc.get("email", "unknown"))
        statuses.append(_yt_safety_rules.get_account_status(label))

    return {
        "today_total": _yt_safety_rules.get_today_total_success(),
        "accounts": statuses,
    }

@app.post("/api/youtube/safety/allow-video")
async def yt_safety_allow_video(request: Request):
    """동일 영상 차단 수동 해제 (재작업 허용)"""
    if not _yt_safety_rules:
        return JSONResponse({"error": "안전 규칙 모듈 미로드"}, 500)
    body = await request.json()
    url = body.get("youtube_url", "")
    if not url:
        return JSONResponse({"error": "YouTube URL 필요"}, 400)
    ok = _yt_safety_rules.allow_video(url)
    return {"success": ok, "msg": "해당 영상 재작업이 허용되었습니다." if ok else "영상 ID 추출 실패"}

@app.get("/api/youtube/safety/posted-videos")
async def yt_safety_posted_videos():
    """댓글이 작성된 영상 목록"""
    if not _yt_safety_rules:
        return {"videos": []}
    return {"videos": _yt_safety_rules.get_posted_videos()}

# ───────────────────────────── SMM (좋아요 구매) ─────────────────────────────

@app.get("/api/youtube/smm/status")
async def yt_smm_status():
    """SMM 패널 상태 (활성화 여부 + 잔액)"""
    if not _yt_smm_client:
        return {"enabled": False, "error": "SMM 모듈 미로드"}
    if not _yt_smm_client.enabled:
        return {"enabled": False, "balance": None, "msg": "SMM_ENABLED=false"}
    balance = await asyncio.get_running_loop().run_in_executor(executor, _yt_smm_client.get_balance)
    return {"enabled": True, "balance": balance}

@app.get("/api/youtube/smm/services")
async def yt_smm_services():
    """SMM 서비스 목록 조회"""
    if not _yt_smm_client or not _yt_smm_client.enabled:
        return {"services": [], "error": "SMM 비활성화"}
    services = await asyncio.get_running_loop().run_in_executor(executor, _yt_smm_client.get_services)
    # YouTube 관련만 필터
    yt_services = [s for s in services if 'youtube' in str(s.get('name','')).lower() or 'yt' in str(s.get('name','')).lower()]
    return {"services": yt_services, "all_count": len(services)}

@app.post("/api/youtube/smm/order")
async def yt_smm_order(request: Request):
    """좋아요 수동 주문 (사용자 승인 후 호출)"""
    if not _yt_smm_client or not _yt_smm_client.enabled:
        return JSONResponse({"error": "SMM 비활성화"}, 400)
    body = await request.json()
    comment_url = body.get("comment_url", "")
    quantity = body.get("quantity", _yt_smm_client.default_quantity)
    service_id = body.get("service_id", None)
    if not comment_url:
        return JSONResponse({"error": "댓글 URL 필요"}, 400)
    result = await asyncio.get_running_loop().run_in_executor(
        executor, lambda: _yt_smm_client.order_likes(comment_url, quantity, service_id)
    )
    return result

@app.get("/api/youtube/smm/pending-likes")
async def yt_smm_pending_likes():
    """승인 대기 중인 좋아요 주문 목록"""
    return {"pending": _yt_autopost_state.get("pending_likes", [])}

@app.post("/api/youtube/smm/approve-likes")
async def yt_smm_approve_likes(request: Request):
    """사용자가 승인한 좋아요 주문 일괄 실행"""
    if not _yt_smm_client or not _yt_smm_client.enabled:
        return JSONResponse({"error": "SMM 비활성화"}, 400)

    pending = _yt_autopost_state.get("pending_likes", [])
    approved = [p for p in pending if p.get("status") == "pending_approval"]
    if not approved:
        return {"success": 0, "msg": "승인 대기 중인 주문 없음"}

    results = []
    for item in approved:
        result = await asyncio.get_running_loop().run_in_executor(
            executor,
            lambda url=item["comment_url"], qty=item["quantity"]: _yt_smm_client.order_likes(url, qty)
        )
        item["status"] = "ordered" if "order" in result else "error"
        item["order_id"] = result.get("order")
        item["error"] = result.get("error")
        results.append(result)

    success = sum(1 for r in results if "order" in r)
    return {"success": success, "total": len(approved), "results": results}

@app.post("/api/youtube/smm/clear-pending")
async def yt_smm_clear_pending():
    """대기 목록 초기화"""
    with _yt_autopost_lock:
        _yt_autopost_state["pending_likes"] = []
    return {"success": True}

@app.post("/api/youtube/smm/check-orders")
async def yt_smm_check_orders(request: Request):
    """주문 상태 확인"""
    if not _yt_smm_client:
        return JSONResponse({"error": "SMM 모듈 미로드"}, 400)
    body = await request.json()
    order_ids = body.get("order_ids", [])
    if not order_ids:
        return {"orders": {}}
    result = await asyncio.get_running_loop().run_in_executor(
        executor, lambda: _yt_smm_client.check_orders(order_ids)
    )
    return {"orders": result}

# ───────────────────────────── IP 변경 (아이폰 테더링) ─────────────────────────────

@app.post("/api/youtube/ip-change")
async def yt_ip_change():
    """Wi-Fi 토글로 아이폰 테더링 IP 변경 (macOS)"""
    import subprocess
    import platform
    if platform.system() != "Darwin":
        return JSONResponse({"error": "macOS에서만 지원"}, 400)

    def _toggle_wifi():
        try:
            # Wi-Fi 끄기
            subprocess.run(["networksetup", "-setairportpower", "en0", "off"], check=True, timeout=10)
            time.sleep(3)
            # Wi-Fi 켜기
            subprocess.run(["networksetup", "-setairportpower", "en0", "on"], check=True, timeout=10)
            time.sleep(5)
            # 새 IP 확인
            result = subprocess.run(["curl", "-s", "https://api.ipify.org"], capture_output=True, text=True, timeout=15)
            new_ip = result.stdout.strip()
            return {"success": True, "new_ip": new_ip}
        except Exception as e:
            return {"success": False, "error": str(e)}

    result = await asyncio.get_running_loop().run_in_executor(executor, _toggle_wifi)
    return result

@app.get("/api/youtube/ip-check")
async def yt_ip_check():
    """현재 공인 IP 확인"""
    import subprocess
    try:
        result = await asyncio.get_running_loop().run_in_executor(
            executor,
            lambda: subprocess.run(["curl", "-s", "https://api.ipify.org"], capture_output=True, text=True, timeout=10)
        )
        return {"ip": result.stdout.strip()}
    except Exception as e:
        return {"ip": None, "error": str(e)}

# ───────────────────────────── 모듈 상태 ─────────────────────────────

@app.get("/api/youtube/autopost/modules-status")
async def yt_modules_status_v2():
    """CommentBoost 모듈 + Stealth 로드 상태"""
    stealth_available = False
    try:
        from src.youtube_bot import has_stealth
        stealth_available = has_stealth()
    except Exception:
        pass
    return {
        "available": _yt_modules_available,
        "stealth": stealth_available,
        "smm_enabled": _yt_smm_client.enabled if _yt_smm_client else False,
    }

# ───────────────────────────── KEYWORD STATUS ─────────────────────────────

def _notion_query_all(db_id, filter_obj=None):
    """노션 DB 전체 페이지 조회 (페이지네이션 포함)"""
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    all_results = []
    has_more = True
    start_cursor = None
    while has_more:
        payload = {'page_size': 100}
        if filter_obj:
            payload['filter'] = filter_obj
        if start_cursor:
            payload['start_cursor'] = start_cursor
        try:
            r = req.post('https://api.notion.com/v1/databases/%s/query' % db_id, headers=headers, json=payload, timeout=30)
            if r.status_code != 200:
                print("Notion query error: %s" % r.text[:200])
                break
            data = r.json()
            all_results.extend(data.get('results', []))
            has_more = data.get('has_more', False)
            start_cursor = data.get('next_cursor')
        except Exception as e:
            print("Notion query error: %s" % e)
            break
    return all_results

def _extract_prop(props, name, prop_type):
    """노션 property에서 값 추출"""
    p = props.get(name, {})
    if prop_type == 'title':
        t = p.get('title', [])
        return t[0]['text']['content'] if t else ''
    elif prop_type == 'rich_text':
        t = p.get('rich_text', [])
        return t[0]['text']['content'] if t else ''
    elif prop_type == 'select':
        s = p.get('select')
        return s.get('name', '') if s else ''
    elif prop_type == 'multi_select':
        ms = p.get('multi_select', [])
        return [x.get('name', '') for x in ms]
    elif prop_type == 'number':
        return p.get('number') or 0
    elif prop_type == 'relation':
        return [x.get('id', '') for x in p.get('relation', [])]
    elif prop_type == 'date':
        d = p.get('date')
        return d.get('start', '') if d else ''
    elif prop_type == 'url':
        return p.get('url', '') or ''
    return ''

@app.get("/api/status/sync")
async def status_sync():
    """키워드 현황 데이터 동기화 (키워드 DB + 콘텐츠 DB 조인)"""
    loop = asyncio.get_running_loop()

    # 키워드 DB 조회
    kw_pages = await loop.run_in_executor(executor, _notion_query_all, KEYWORD_DB_ID, None)
    # 콘텐츠 DB 조회
    ct_pages = await loop.run_in_executor(executor, _notion_query_all, CONTENT_DB_ID, None)

    # 콘텐츠 → 키워드 릴레이션 매핑
    content_by_kw = {}  # kw_page_id → [content_info, ...]
    for page in ct_pages:
        props = page.get('properties', {})
        kw_rels = _extract_prop(props, '키워드', 'relation')
        ct_info = {
            'title': _extract_prop(props, '제목', 'title'),
            'channel': _extract_prop(props, '채널', 'select'),
            'prod_status': _extract_prop(props, '생산 상태', 'select'),
            'deploy_status': _extract_prop(props, '발행_상태', 'select'),
            'deploy_date': _extract_prop(props, '생성일', 'date'),
            'deploy_url': _extract_prop(props, '발행_URL', 'url'),
            'body_summary': _extract_prop(props, '본문', 'rich_text'),
        }
        for kw_id in kw_rels:
            if kw_id not in content_by_kw:
                content_by_kw[kw_id] = []
            content_by_kw[kw_id].append(ct_info)

    # 키워드 → 최종 행 만들기
    rows = []
    for page in kw_pages:
        pid = page['id']
        props = page.get('properties', {})
        keyword = _extract_prop(props, '키워드', 'title')
        if not keyword:
            continue
        channels = _extract_prop(props, '배정 채널', 'multi_select')
        search_vol = _extract_prop(props, '검색량', 'number')
        competition = _extract_prop(props, '경쟁 강도', 'select')
        contact = _extract_prop(props, '구매여정_단계', 'select')
        status = _extract_prop(props, '상태', 'select')

        contents = content_by_kw.get(pid, [])
        work_status = '생성완료' if contents else '미작업'
        deploy_status = ''
        deploy_date = ''
        deploy_url = ''
        content_title = ''
        body_preview = ''
        if contents:
            ct = contents[0]
            deploy_status = ct.get('deploy_status', '')
            deploy_date = ct.get('deploy_date', '')
            deploy_url = ct.get('deploy_url', '')
            content_title = ct.get('title', '')
            body_preview = ct.get('body_summary', '')[:200]

        rows.append({
            'page_id': pid, 'keyword': keyword,
            'channels': ','.join(channels) if channels else '',
            'search_vol': search_vol, 'competition': competition,
            'contact_point': contact, 'kw_status': status,
            'work_status': work_status, 'deploy_status': deploy_status,
            'deploy_date': deploy_date, 'deploy_url': deploy_url,
            'content_title': content_title, 'body_preview': body_preview,
            'exposure': '-', 'rank': '', 'views': '',
        })

    # 집계
    total = len(rows)
    created = sum(1 for r in rows if r['work_status'] == '생성완료')
    deployed = sum(1 for r in rows if r['deploy_status'] == '발행완료')
    unworked = total - created
    summary = {'total': total, 'created': created, 'deployed': deployed, 'exposed': 0, 'unworked': unworked}

    return {'rows': rows, 'summary': summary}

def _check_exposure_one(keyword, deploy_url):
    """네이버 검색에서 해당 URL 노출 여부 확인"""
    if not deploy_url:
        return {'exposure': '-', 'rank': ''}
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get('https://search.naver.com/search.naver?query=%s&where=nexearch' % quote(keyword), headers=headers, timeout=10)
        html = r.text
        # URL 존재 확인
        if deploy_url in html:
            # 순위 추정: 링크 등장 위치
            soup = BeautifulSoup(html, 'html.parser')
            rank = 0
            for i, a in enumerate(soup.find_all('a', href=True)):
                if deploy_url in a.get('href', ''):
                    rank = i + 1
                    break
            return {'exposure': '노출중', 'rank': str(rank) if rank else '확인됨'}
        return {'exposure': '미노출', 'rank': ''}
    except Exception:
        return {'exposure': '-', 'rank': ''}

@app.post("/api/status/check-exposure")
async def check_exposure(request: Request):
    """노출 체크 (SSE)"""
    body = await request.json()
    items = body.get('items', [])  # [{keyword, deploy_url}, ...]

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(items)
        results = []
        for i, item in enumerate(items):
            kw = item.get('keyword', '')
            url = item.get('deploy_url', '')
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s 노출 확인 중...' % (i+1, total, kw), 'cur': i+1, 'total': total})
            result = await loop.run_in_executor(executor, _check_exposure_one, kw, url)
            result['keyword'] = kw
            results.append(result)
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})
            await asyncio.sleep(1.5)
        yield _sse({'type': 'complete', 'total': total, 'results': results})
      except Exception as e:
        print(f"[check_exposure] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'노출 체크 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

# ───────────────────────────── TIKTOK ─────────────────────────────

def _build_tiktok_prompt(keyword, appeal, buying_one, product, forbidden):
    brand_kw = product.get('brand_keyword', '')
    system = """역할: 틱톡 숏폼 영상의 스크립트를 작성하는 바이럴 콘텐츠 전문가.

목표: 제품의 소구점과 타겟 고민을 기반으로, 틱톡에서 자연스럽게 시청되고 검색 유도 효과를 만드는 숏폼 영상 스크립트 1개를 생성한다. UGC(User Generated Content) 스타일 — 실제 사용자가 찍은 듯한 자연스러운 톤이 핵심.

스크립트 구조 (15~30초 분량):

[1. 후킹] (0~3초)
- 시청자의 스크롤을 멈추게 하는 첫 문장
- 부정 편향 or 궁금증 자극 or 충격적 사실
- 소구점의 구매원씽에서 파생된 결핍 언어 사용

[2. 문제 공감] (3~10초)
- 타겟이 겪는 구체적 상황/증상 묘사
- 혼잣말 톤: "나도 진짜 이거 때문에 미치는 줄…"
- 기존에 시도했던 방법 + 실패 경험 간략히

[3. 전환점] (10~20초)
- 해결 계기를 자연스럽게 소개
- 나만의 키워드를 1회 자연스럽게 삽입
- 구체적 변화 언급: 기간 + 체감 효과

[4. 마무리] (20~30초)
- 간결한 결과 요약
- 검색 유도: "궁금하면 {나만의 키워드} 검색해봐"
- 또는 댓글 유도: "나만 이런 거 아니지…?"

작성 규칙:
[톤과 말투]
- 자연스러운 구어체. 친구한테 말하듯이.
- 혼잣말 톤 중심: "~~했음", "~~뒤집어짐ㅠㅠ"
- 딱딱하거나 격식 있는 표현 금지
- 이모티콘/감정 표현 자연스럽게 사용

[길이] 공백 제외 200~400자. 짧고 임팩트 있게.

[제품 노출]
- 나만의 키워드를 스크립트 안에 1회만 자연스럽게 포함
- 제품명 직접 언급 금지 (나만의 키워드만 사용)
- 해시태그·구매링크 금지

[영상 연출 가이드]
- 스크립트 옆에 [연출: ...] 메모를 간단히 표기

출력 형식:
[후킹] (0~3초)
(대사)
[연출: ...]

[문제 공감] (3~10초)
(대사)
[연출: ...]

[전환점] (10~20초)
(대사)
[연출: ...]

[마무리] (20~30초)
(대사)
[연출: ...]

금지사항:
- "강추", "인생템", "꼭 사세요" 등 광고 어투 금지
- 구매 링크·쇼핑몰 언급 금지
- 나만의 키워드 2회 이상 사용 금지
- 제품명(브랜드명) 직접 언급 금지
- 스튜디오 촬영 느낌 금지 — UGC 느낌 유지
"""

    user = "메인 키워드: %s\n소구점: %s\n구매원씽: %s\n제품명: %s\n주요 성분: %s\n핵심 특징: %s\n타겟층: %s\n나만의 키워드: %s" % (
        keyword, appeal, buying_one,
        product.get('name',''), product.get('ingredients',''),
        product.get('usp',''), product.get('target',''), brand_kw)
    return system, user

@app.get("/api/tiktok/notion-keywords")
async def tiktok_notion_keywords():
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    payload = {
        'filter': {'and': [
            {'property': '배정 채널', 'multi_select': {'contains': '틱톡'}},
            {'property': '상태', 'select': {'equals': '미사용'}},
        ]},
        'page_size': 100,
    }
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % KEYWORD_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code != 200:
            return {'keywords': []}
        keywords = []
        for page in r.json().get('results', []):
            props = page.get('properties', {})
            t = props.get('키워드', {}).get('title', [])
            kw = t[0]['text']['content'] if t else ''
            if kw:
                keywords.append({'keyword': kw, 'page_id': page['id']})
        return {'keywords': keywords}
    except Exception:
        return {'keywords': []}

@app.post("/api/tiktok/generate")
async def tiktok_generate(request: Request):
    body = await request.json()
    keywords = body.get('keywords', [])
    product = body.get('product', {})
    appeal = body.get('appeal', '')
    buying_one = body.get('buying_one', '')
    forbidden = body.get('forbidden', '')
    count = body.get('count', 1)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords) * count
        idx = 0
        for kw_data in keywords:
            kw = kw_data['keyword']
            for c in range(count):
                idx += 1
                label = '%s' % kw if count == 1 else '%s (#%d)' % (kw, c+1)
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 스크립트 생성 중...' % (idx, total, label), 'cur': idx-1, 'total': total})
                sys_p, usr_p = _build_tiktok_prompt(kw, appeal, buying_one, product, forbidden)
                script = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
                script = script.strip()
                result = {
                    'keyword': kw, 'script': script,
                    'page_id': kw_data.get('page_id', ''),
                    'num': c + 1,
                }
                yield _sse({'type': 'result', 'data': result, 'cur': idx, 'total': total})
        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[tiktok_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'틱톡 스크립트 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/tiktok/save-notion")
async def tiktok_save_notion(request: Request):
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    kw = body.get('keyword', '')
    props = {
        '제목': {'title': [{'text': {'content': '%s 틱톡 스크립트' % kw}}]},
        '채널': {'select': {'name': '틱톡'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    script = body.get('script', '')
    if script:
        props['본문'] = {'rich_text': [{'text': {'content': script[:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    if script:
        children = []
        for para in [p.strip() for p in script.split('\n\n') if p.strip()][:100]:
            for k in range(0, len(para), 2000):
                children.append({'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}})
        payload['children'] = children[:100]
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ───────────────────────────── SHORTS (숏츠 제작) ─────────────────────────────

def _build_shorts_topics_prompt(material, content_type):
    """숏츠 주제 5개 제안 프롬프트"""
    style_guide = ""
    if content_type == "썰형":
        style_guide = """'썰형' 스타일:
1. 강력한 훅 (갈등/좌절): 이혼, 이별, 자존감 바닥, 관계 파탄 등 인생 최악의 순간을 제시.
2. 감정적 고통 (상황 묘사): 상대방의 경멸, 무시, 혹은 자신의 비참함을 구체적으로 묘사.
3. 탐색 (절박함): '이대로 살 수 없다'는 절박함으로 방법을 찾기 시작.
4. 발견 (사회적 증거): 리뷰, 후기 등 신뢰할 만한 근거를 보고 제품을 알게 됨.
5. 극적인 반전 (결과): 상상 이상으로 과장되고 극적인 결과.
6. 행동 유도 (CTA): 명확한 검색/터치 지시."""
    else:
        style_guide = """'정보형' 스타일:
- 잘못된 상식 지적, 전문가 비밀 폭로, 자가 진단 체크리스트 등
- 정보가치가 높아 자연스럽게 시청하게 되는 구조
- 권위/데이터를 활용한 신뢰 확보"""

    system = f"""너는 숏츠 영상 대본을 전문적으로 작성하는 마케팅 어시스턴트다.

[기본 재료]
- 제품명: {material.get('product', '')}
- 타겟 고객: {material.get('target', '')}
- 타겟의 핵심 문제: {material.get('problem', '')}
- 문제가 유발하는 감정: {material.get('emotion', '')}
- 신뢰 근거: {material.get('trust', '')}
- CTA: {material.get('cta', '')}

[유형: {content_type}]
{style_guide}

임무: 위 기본 재료를 활용하여 '{content_type}' 유형에 맞는 창의적이고 매력적인 주제(앵글) 5가지를 새롭게 창작하여 제안하라.
각 주제가 왜 타겟에게 매력적일지 1줄 요약 이유를 포함해야 한다.

출력 형식 (정확히 따를 것):
1. [주제 제목] — [매력 포인트 1줄]
2. [주제 제목] — [매력 포인트 1줄]
3. [주제 제목] — [매력 포인트 1줄]
4. [주제 제목] — [매력 포인트 1줄]
5. [주제 제목] — [매력 포인트 1줄]"""

    return system, f"'{content_type}' 유형으로 주제 5가지를 제안해줘."


def _build_shorts_script_prompt(material, content_type, topic, length):
    """숏츠 대본 생성 프롬프트"""
    style_guide = ""
    if content_type == "썰형":
        style_guide = """[핵심 공식]
1. 강력한 훅 (갈등/좌절): 이혼, 이별, 자존감 바닥, 관계 파탄 등 인생 최악의 순간을 제시.
2. 감정적 고통 (상황 묘사): 상대방의 경멸, 무시, 혹은 자신의 비참함을 구체적으로 묘사. (예: 눈도 안 마주침, 도망침, 우울증)
3. 탐색 (절박함): '이대로 살 수 없다', '너무 창피하다'는 절박함으로 방법을 찾기 시작.
4. 발견 (사회적 증거): 리뷰, 틱톡 후기 등 신뢰할 만한 근거를 보고 제품을 알게 됨.
5. 극적인 반전 (결과): 상상 이상으로 과장되고 극적인 결과. (예: 구체적 숫자 곁들이기)
6. 행동 유도 (CTA): 명확한 지시.
7. 말투는 ~네요. ~다. ~고요 등으로 화자가 직접 이야기를 전달하는 형태. ('' 대사표현 금지)"""
    else:
        style_guide = """[핵심 공식]
초반 3초 - 결핍 강조 또는 잘못된 상식 지적
정보 제공 → 한계 제시
제품 소개 + 권위 부여 + 효과 제시 (구체적 숫자 곁들이기)
CTA"""

    system = f"""너는 숏츠 영상 대본을 전문적으로 작성하는 마케팅 어시스턴트다.

[기본 재료]
- 제품명: {material.get('product', '')}
- 타겟 고객: {material.get('target', '')}
- 타겟의 핵심 문제: {material.get('problem', '')}
- 문제가 유발하는 감정: {material.get('emotion', '')}
- 신뢰 근거: {material.get('trust', '')}
- CTA: {material.get('cta', '')}

[유형: {content_type}]
{style_guide}

작성 규칙:
- 약 {length}자 분량의 숏츠 대본을 작성한다.
- 이 대본은 그대로 TTS 음성으로 읽히므로, 읽었을 때 자연스러워야 한다.
- 이모지, 해시태그, 특수기호 사용 금지 (TTS가 읽을 수 없음)
- [연출], [자막], [장면] 같은 메타 표기 금지 — 순수 나레이션 텍스트만 출력
- 문장은 짧게 끊어서. 줄바꿈으로 문장을 구분.

주의사항: 대본 출력시 오로지 대본 내용만 출력. 다른 설명 없이."""

    return system, f"다음 주제로 대본을 작성해줘: {topic}"


def _build_shorts_hooks_prompt():
    """썸네일 훅(제목) 생성 프롬프트"""
    return """당신은 칩 히스의 '스틱(Stick!)' 원칙과 '자청' 스타일의 욕망/결핍 기반 카피라이팅 원칙을 모두 마스터한, 숏폼 비디오 전문 바이럴 마케터입니다.
당신의 임무는 초반 2초 이탈률을 0%에 가깝게 만드는 '스크롤 스토퍼(Scroll-stopper)' 훅(제목)을 생성하는 것입니다.

# 작업 프로세스
1. 1단계 (분석): 스크립트를 읽고, [핵심 메시지], [핵심 타겟 고객], [타겟의 결핍 또는 욕망]을 내부적으로 정의합니다.
2. 2단계 (생성): 아래 7가지 원칙을 창의적으로 조합하여 훅(제목) 카피 10개를 생성합니다.

# 생성 원칙 (스틱! + 자청 결합)
1. 단순성 (Simple): 핵심을 하나의 강력하고 짧은 문장으로 압축.
2. 의외성 (Unexpected): 타겟의 통념이나 예상을 깨뜨림. (예: "OOO, 사기였습니다.")
3. 구체성 (Concrete): 감각적이고 구체적인 단어, 숫자, 고유명사 사용.
4. 권위/신뢰 (Authority): 전문가, 데이터, 연구 결과를 암시.
5. 스토리 (Story): 극적인 변화나 경험을 암시.
6. 금지/위협 (Prohibition): 손실 회피 심리 자극. "하지 마라", "이거 모르면 손해".
7. 자아 흠집 (Ego Scratch): 자존심, 우월감, 불안감을 건드림.

# 산출물
- 2~3초 이내 인지 가능한 매우 짧고 간결한 훅 카피 10개
- 각 카피 뒤에 사용한 핵심 원칙을 괄호 안에 명시 (예: (의외성, 구체성))
- 번호를 매겨 출력"""


def _elevenlabs_tts_with_timestamps(text, voice_id, model_id="eleven_multilingual_v2"):
    """ElevenLabs TTS 호출 — 음성 + 캐릭터별 타임스탬프 반환"""
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}/with-timestamps"
    headers = {
        "xi-api-key": ELEVENLABS_API_KEY,
        "Content-Type": "application/json",
    }
    payload = {
        "text": text,
        "model_id": model_id,
        "voice_settings": {"stability": 0.5, "similarity_boost": 0.75},
    }
    r = req.post(url, headers=headers, json=payload, timeout=120)
    if r.status_code != 200:
        raise Exception(f"ElevenLabs API 에러 ({r.status_code}): {r.text[:300]}")
    return r.json()


def _generate_srt_from_alignment(text, alignment, words_per_segment=3):
    """캐릭터 타임스탬프로부터 SRT 자막 파일 생성"""
    chars = alignment.get("characters", [])
    starts = alignment.get("character_start_times_seconds", [])
    ends = alignment.get("character_end_times_seconds", [])

    words = []
    current_word = ""
    word_start = None
    word_end = None
    for i, ch in enumerate(chars):
        if ch.strip() == "" or ch in ("\n", "\r"):
            if current_word:
                words.append({"text": current_word, "start": word_start, "end": word_end})
                current_word = ""
                word_start = None
        else:
            if word_start is None:
                word_start = starts[i] if i < len(starts) else 0
            word_end = ends[i] if i < len(ends) else word_start
            current_word += ch
    if current_word:
        words.append({"text": current_word, "start": word_start, "end": word_end})

    segments = []
    for i in range(0, len(words), words_per_segment):
        group = words[i:i + words_per_segment]
        seg_text = " ".join(w["text"] for w in group)
        seg_start = group[0]["start"]
        seg_end = group[-1]["end"]
        segments.append({"start": seg_start, "end": seg_end, "text": seg_text})

    def _fmt_time(sec):
        h = int(sec // 3600)
        m = int((sec % 3600) // 60)
        s = int(sec % 60)
        ms = int((sec - int(sec)) * 1000)
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    srt_lines = []
    for idx, seg in enumerate(segments, 1):
        srt_lines.append(str(idx))
        srt_lines.append(f"{_fmt_time(seg['start'])} --> {_fmt_time(seg['end'])}")
        srt_lines.append(seg["text"])
        srt_lines.append("")

    return "\n".join(srt_lines)


@app.get("/api/shorts/voices")
async def shorts_voices():
    """ElevenLabs 사용 가능한 음성 목록"""
    if not ELEVENLABS_API_KEY:
        return JSONResponse({"error": "ELEVENLABS_API_KEY가 설정되지 않았습니다"}, 400)
    try:
        r = req.get("https://api.elevenlabs.io/v1/voices", headers={"xi-api-key": ELEVENLABS_API_KEY}, timeout=15)
        if r.status_code != 200:
            return JSONResponse({"error": f"API 에러: {r.status_code}"}, 500)
        voices = []
        for v in r.json().get("voices", []):
            voices.append({
                "voice_id": v["voice_id"],
                "name": v["name"],
                "category": v.get("category", ""),
                "labels": v.get("labels", {}),
                "preview_url": v.get("preview_url", ""),
            })
        return {"voices": voices}
    except Exception as e:
        return JSONResponse({"error": str(e)}, 500)


@app.post("/api/shorts/topics")
async def shorts_topics(request: Request):
    """숏츠 주제 5개 제안"""
    body = await request.json()
    material = body.get("material", {})
    content_type = body.get("type", "썰형")

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        yield _sse({"type": "progress", "msg": "주제 생성 중..."})
        sys_p, usr_p = _build_shorts_topics_prompt(material, content_type)
        result = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
        result = result.strip()
        if result.startswith("[ERROR]"):
            yield _sse({"type": "error", "message": result})
            return
        yield _sse({"type": "topics", "text": result})
        yield _sse({"type": "complete"})
      except Exception as e:
        print(f"[shorts_topics] 에러: {e}")
        yield _sse({"type": "error", "message": f"주제 생성 중 오류: {e}"})

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/shorts/script")
async def shorts_script(request: Request):
    """숏츠 대본 생성"""
    body = await request.json()
    material = body.get("material", {})
    content_type = body.get("type", "썰형")
    topic = body.get("topic", "")
    length = body.get("length", 600)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        yield _sse({"type": "progress", "msg": "대본 생성 중..."})
        sys_p, usr_p = _build_shorts_script_prompt(material, content_type, topic, length)
        result = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
        result = result.strip()
        if result.startswith("[ERROR]"):
            yield _sse({"type": "error", "message": result})
            return
        yield _sse({"type": "script", "text": result})
        yield _sse({"type": "complete"})
      except Exception as e:
        print(f"[shorts_script] 에러: {e}")
        yield _sse({"type": "error", "message": f"대본 생성 중 오류: {e}"})

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/shorts/hooks")
async def shorts_hooks(request: Request):
    """썸네일 훅(제목) 10개 생성"""
    body = await request.json()
    script = body.get("script", "")

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        yield _sse({"type": "progress", "msg": "썸네일 훅 생성 중..."})
        sys_p = _build_shorts_hooks_prompt()
        usr_p = f"[스크립트 전문]\n{script}"
        result = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
        result = result.strip()
        if result.startswith("[ERROR]"):
            yield _sse({"type": "error", "message": result})
            return
        yield _sse({"type": "hooks", "text": result})
        yield _sse({"type": "complete"})
      except Exception as e:
        print(f"[shorts_hooks] 에러: {e}")
        yield _sse({"type": "error", "message": f"훅 생성 중 오류: {e}"})

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/api/shorts/tts")
async def shorts_tts(request: Request):
    """TTS 음성 + SRT 자막 생성"""
    body = await request.json()
    script = body.get("script", "")
    voice_id = body.get("voice_id", "")
    words_per_seg = body.get("words_per_segment", 3)

    if not script:
        return JSONResponse({"error": "대본이 없습니다"}, 400)
    if not voice_id:
        return JSONResponse({"error": "음성을 선택하세요"}, 400)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()

        yield _sse({"type": "progress", "msg": "음성 생성 중 (ElevenLabs)..."})
        tts_result = await loop.run_in_executor(
            executor, _elevenlabs_tts_with_timestamps, script, voice_id
        )

        audio_b64 = tts_result.get("audio_base64", "")
        if not audio_b64:
            yield _sse({"type": "error", "message": "음성 생성 실패: 오디오 데이터 없음"})
            return
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        audio_filename = f"shorts_{ts}.mp3"
        audio_path = os.path.join(SHORTS_DIR, audio_filename)
        with open(audio_path, "wb") as f:
            f.write(base64.b64decode(audio_b64))

        yield _sse({"type": "progress", "msg": "자막(SRT) 생성 중..."})
        alignment = tts_result.get("alignment", {})
        srt_content = _generate_srt_from_alignment(script, alignment, words_per_seg)
        srt_filename = f"shorts_{ts}.srt"
        srt_path = os.path.join(SHORTS_DIR, srt_filename)
        with open(srt_path, "w", encoding="utf-8") as f:
            f.write(srt_content)

        txt_filename = f"shorts_{ts}.txt"
        txt_path = os.path.join(SHORTS_DIR, txt_filename)
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(script)

        yield _sse({
            "type": "complete",
            "audio_url": f"/api/shorts/download/{audio_filename}",
            "srt_url": f"/api/shorts/download/{srt_filename}",
            "txt_url": f"/api/shorts/download/{txt_filename}",
            "srt_preview": srt_content[:500],
        })

      except Exception as e:
        print(f"[shorts_tts] 에러: {e}")
        yield _sse({"type": "error", "message": f"TTS 생성 중 오류: {e}"})

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/api/shorts/download/{filename}")
async def shorts_download(filename: str):
    """생성된 숏츠 파일 다운로드"""
    safe_name = os.path.basename(filename)
    if not re.match(r'^[\w\-.]+$', safe_name):
        return JSONResponse({"error": "invalid filename"}, status_code=400)
    fpath = os.path.join(SHORTS_DIR, safe_name)
    if not os.path.exists(fpath):
        return JSONResponse({"error": "파일을 찾을 수 없습니다"}, 404)
    media_types = {".mp3": "audio/mpeg", ".srt": "text/srt", ".txt": "text/plain"}
    ext = os.path.splitext(safe_name)[1]
    return FileResponse(fpath, filename=safe_name, media_type=media_types.get(ext, "application/octet-stream"))


# ───────────────────────────── COMMUNITY ─────────────────────────────

COMMUNITY_TONES = {
    '뽐뿌': '"이거 써봤는데" 후기형 + 사진 3장+ 언급. 광고 감시 강함. 중간 길이(500~1000자).',
    '디시인사이드': '짧고 솔직한 디시 특유 말투, 유동닉 스타일. 광고 적대적. 짧게(200~500자).',
    '에펨코리아': '디시와 유사하나 약간 더 정리된 톤. 과한 홍보 금지. 짧게(200~500자).',
    '클리앙': '분석적, IT/가전 특화 톤. 분석 없는 단순 추천은 무시당함. 중간(500~1000자).',
    '맘스홀릭': '따뜻한 존댓말, 엄마 입장 공감. 구매력 높지만 광고 감지력도 높음. 중간(500~1000자).',
    '더쿠': '연예·뷰티 톤, "~인데" "~이더라" 식. 광고 티 나면 집단 공격. 짧게(200~500자).',
    '인스티즈': '연예·뷰티 톤, 가벼운 대화체. 짧게(200~500자).',
    '화해': '성분 리뷰 특화, 분석적 톤. 성분 근거 없으면 신뢰 안 됨. 중간(500~1000자).',
    '블라인드': '직장인 톤, 간결하고 실용적. 광고 적대적. 짧게(200~500자).',
}

STRATEGY_NAME = {
    '1': '고민 공감 + 체험 후기형',
    '2': '추천 요청형 (낚시)',
    '3': '비교 리뷰형',
    '4': '자기 제품 까기 (역발상)',
}

def _build_community_post_prompt(community, strategy, keyword, appeal, buying_one, product, forbidden):
    brand_kw = product.get('brand_keyword', '')
    strategy_label = STRATEGY_NAME.get(strategy, strategy)
    system = """너는 온라인 커뮤니티에 자연스럽게 침투하는 체험 공유 글을 작성하는 마케팅 전문가다.

광고 글이 아닌, '진짜 사용자'처럼 보이는 체험 공유 글을 작성해야 한다.
"우리 제품 좋아요 구매하세요" 같은 직접 광고는 즉시 삭제 + 차단 + 네거티브 역효과로 이어진다.

---

[글 유형 — 지정된 유형에 맞춰 작성]

● 유형1: 고민 공감 + 체험 후기형
- 고민 공감으로 시작 → 공감 유도
- 여러 제품 써본 경험 → 객관성 확보
- 나만의 키워드 1회만 자연스럽게 삽입
- 마무리를 '질문'으로 → 댓글 유도 → 글 활성화
- 예시 제목: "탈모 때문에 이것저것 써봤는데 요즘 쓰는 거 공유합니다"

● 유형2: 추천 요청형 (낚시)
- 추천 요청 형태의 자연스러운 질문
- 나만의 키워드를 '들어본 적 있다' 수준으로 언급
- 댓글에서 다른 계정이 "나 써봤는데 괜찮았어" 답변할 수 있는 구조
- 예시 제목: "다이어트 보조제 추천 좀…"

● 유형3: 비교 리뷰형
- 3개 제품 비교 형식 → 정보 가치가 높아 삭제 안 됨
- 내 제품을 3개 중 1개로 자연스럽게 포함
- "개인 의견" 면책 문구 포함
- 예시 제목: "[미니 리뷰] 모공앰플 3종 비교 (A / B / C)"

● 유형4: 자기 제품 까기 (역발상)
- 단점을 먼저 말하면 "광고면 단점을 말하겠어?" → 진짜 사용자로 인식
- 단점은 사소하게(향, 가격), 장점은 구체적으로
- "세일 때 사라" 식으로 구매 유도하면서 광고 안 티남

---

[커뮤니티별 톤 매칭 — 반드시 해당 커뮤니티 말투에 맞출 것]

● 뽐뿌: "이거 써봤는데" 후기형, 30~50대 남녀, 가성비 강조, 사진 3장+ 언급
● 디시인사이드: 짧고 솔직한 톤, 반말, "~임/~함/~ㅋㅋ" 체, 질문형이 안전
● 에펨코리아: IT/일상 톤, 반말, 디시보다 약간 부드러움
● 클리앙: 분석적/논리적 톤, 존댓말, 성분/스펙 비교 선호
● 맘스홀릭: 육아맘 톤, 존댓말, 공감 중심, 가족 상황 언급
● 더쿠/인스티즈: 여성 커뮤 톤, 반말/존댓말 혼용, 가벼운 리뷰
● 화해: 성분 중심 리뷰, 분석적 톤, 존댓말
● 블라인드: 직장인 톤, 존댓말, 현실적 고민 공유

---

[작성 규칙]
- 광고 티 절대 금지 (조금이라도 나면 '홍보충' 낙인)
- 브랜드명 직접 언급 금지
- 나만의 키워드는 글당 1회만 자연스럽게 삽입
- 직접 링크 삽입 금지 → 키워드 검색 유도
- 제목과 본문 모두 출력
- 해당 커뮤니티에서 실제로 올라올 법한 글처럼 작성

[출력 형식]
제목: (내용)

본문:
(내용)"""

    forbidden_line = ("\n금칙어: %s" % forbidden) if forbidden else ''
    user = """[커뮤니티]
%s

[글 유형]
%s

[제품 정보]
제품명: %s
USP: %s
타겟층: %s
주요 성분: %s
소구점: %s
구매원씽: %s%s

[나만의 키워드]
%s""" % (community, strategy_label, product.get('name',''), product.get('usp',''),
         product.get('target',''), product.get('ingredients',''),
         appeal, buying_one, forbidden_line, brand_kw)
    return system, user

def _build_community_comments_prompt(community, post_body, brand_kw):
    tone = COMMUNITY_TONES.get(community, '')
    system = """역할: 커뮤니티 침투글에 달릴 자연스러운 자작 댓글을 작성하는 작가.

목표: 침투글에 대한 자작 댓글 1~2개를 생성한다. 다른 계정이 쓴 것처럼 보여야 한다.

커뮤니티 톤: %s

작성 규칙:
1. 댓글 1~2개 (과하면 의심)
2. 원글 작성자와 완전히 다른 사람의 톤
3. 나만의 키워드는 댓글 전체에서 1회만 (원글에서 이미 언급된 경우 안 써도 됨)
4. 공감형 or "나도 써봤는데" 경험 공유형

출력 형식:
[댓글 1] ...
[댓글 2] ...

금지사항:
- 원글과 같은 톤/말투 사용 금지
- 링크·광고 어투·과도한 칭찬 금지""" % tone
    user = "침투글 본문: %s\n나만의 키워드: %s" % (post_body[:2000], brand_kw)
    return system, user

@app.post("/api/community/generate")
async def community_generate(request: Request):
    body = await request.json()
    keywords = body.get('keywords', [])
    community = body.get('community', '')
    strategy = body.get('strategy', '1')
    product = body.get('product', {})
    appeal = body.get('appeal', '')
    buying_one = body.get('buying_one', '')
    forbidden = body.get('forbidden', '')
    include_comments = body.get('include_comments', True)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords)
        for i, kw in enumerate(keywords):
            keyword = kw if isinstance(kw, str) else kw.get('keyword', '')
            page_id = '' if isinstance(kw, str) else kw.get('page_id', '')

            # STEP 1: 침투글
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 침투글 생성 중...' % (i+1, total, keyword), 'cur': i, 'total': total})
            sys1, usr1 = _build_community_post_prompt(community, strategy, keyword, appeal, buying_one, product, forbidden)
            raw = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
            parsed = _parse_viral_output(raw)  # reuse title/body parser

            # STEP 2: 자작 댓글
            comments = ''
            if include_comments:
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 자작 댓글 생성 중...' % (i+1, total, keyword), 'cur': i, 'total': total})
                sys2, usr2 = _build_community_comments_prompt(community, parsed['body'], product.get('brand_keyword', ''))
                comments = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
                comments = comments.strip()

            result = {
                'keyword': keyword, 'community': community, 'strategy': strategy,
                'title': parsed['title'], 'body': parsed['body'], 'comments': comments,
                'page_id': page_id,
            }
            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})
        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[community_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'커뮤니티 침투글 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/community/save-notion")
async def community_save_notion(request: Request):
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '커뮤니티'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('body'):
        props['본문'] = {'rich_text': [{'text': {'content': body['body'][:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    full = body.get('body', '')
    if body.get('comments'):
        full += '\n\n---\n자작 댓글:\n' + body['comments']
    if full:
        children = []
        for para in [p.strip() for p in full.split('\n\n') if p.strip()][:100]:
            for k in range(0, len(para), 2000):
                children.append({'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}})
        payload['children'] = children[:100]
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ───────────────────────────── KEYWORD ENDPOINTS (continued) ─────────────────────────────

@app.post("/api/keywords/download-excel")
async def download_excel(request: Request):
    """서버 사이드 엑셀 생성 + 다운로드"""
    from fastapi.responses import Response
    body = await request.json()
    items = body.get('items', [])
    filename = body.get('filename', '키워드')
    sheet_name = body.get('sheet', 'Sheet1')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name

    if items:
        # 헤더
        headers = list(items[0].keys())
        ws.append(headers)
        # 데이터
        for item in items:
            ws.append([item.get(h, '') for h in headers])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    safe_filename = quote(f"{filename}.xlsx")
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}",
        }
    )

# ═══════════════════════════ PHOTO COLLECTION ═══════════════════════════
PHOTO_DIR = os.path.join(os.path.dirname(__file__), "photos")
PHOTO_LIB_FILE = os.path.join(os.path.dirname(__file__), "photo_library.json")
os.makedirs(PHOTO_DIR, exist_ok=True)

def _translate_ko_to_zh(text):
    """한국어 → 중국어 번역 (Gemini Flash)"""
    try:
        url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=" + GEMINI_API_KEY
        payload = {
            "contents": [{"parts": [{"text": "다음 한국어를 중국어(간체)로 번역해줘. 번역 결과만 출력하고 다른 설명은 하지 마.\n\n" + text}]}]
        }
        r = req.post(url, json=payload, timeout=10)
        data = r.json()
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except Exception:
        return text

def _crawl_baidu_images(driver, query_zh, count):
    """바이두 이미지 검색 + 고해상도 이미지 URL 수집"""
    results = []
    try:
        url = f"https://image.baidu.com/search/index?tn=baiduimage&word={quote(query_zh)}"
        driver.get(url)
        time.sleep(3)

        # 스크롤해서 더 많은 이미지 로드
        scroll_count = max(1, count // 20)
        for _ in range(scroll_count):
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            time.sleep(1.5)

        # 이미지 추출 (objURL이 보통 고해상도 원본)
        import re
        html = driver.page_source
        # objURL": "http://..." 형태 정규식
        urls = re.findall(r'"objURL":"(http://[^"]+|https://[^"]+)"', html)
        if not urls:
            # 보조 수단: 일반 img 태그의 src / data-src
            imgs = driver.find_elements(By.CSS_SELECTOR, "img")
            for img in imgs:
                u = img.get_attribute("data-imgurl") or img.get_attribute("data-src") or img.get_attribute("src")
                if u and u.startswith("http") and "avatar" not in u and "logo" not in u:
                    if u not in urls:
                        urls.append(u)

        for u in urls[:count]:
            if u.startswith("http"):
                results.append(u)
    except Exception as e:
        print(f"[baidu] error: {e}")
    return results

XHS_PATH = os.environ.get('XHS_PATH', shutil.which('xhs') or '/Users/iconlms/Library/Python/3.11/bin/xhs')

def _crawl_xhs_images(query_zh, count):
    """샤오홍슈 이미지 수집 (xiaohongshu-cli 사용)"""
    results = []
    try:
        result = subprocess.run(
            [XHS_PATH, "search", query_zh, "--sort", "popular", "--json"],
            capture_output=True, text=True, timeout=30
        )
        data = json.loads(result.stdout)
        items = data.get("data", {}).get("items", [])

        for item in items:
            if len(results) >= count:
                break
            nc = item.get("note_card", {})
            # 이미지 타입 게시물만 (영상 제외)
            if nc.get("type") == "video":
                continue
            for img in nc.get("image_list", []):
                if len(results) >= count:
                    break
                # WB_DFT = 고해상도
                url = ""
                for info in img.get("info_list", []):
                    if info.get("image_scene") == "WB_DFT":
                        url = info.get("url", "")
                        break
                if not url:
                    url = img.get("info_list", [{}])[0].get("url", "") if img.get("info_list") else ""
                if url and url.startswith("http"):
                    results.append(url)
    except Exception as e:
        print(f"[xhs-cli] search error: {e}")
    return results

def _download_image(url, filename, category):
    """이미지 다운로드 + 카테고리별 폴더 저장"""
    try:
        cat_dir = os.path.join(PHOTO_DIR, category)
        os.makedirs(cat_dir, exist_ok=True)
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Referer": "https://image.baidu.com/"
        }
        r = req.get(url, headers=headers, timeout=15, stream=True)
        if r.status_code == 200:
            filepath = os.path.join(cat_dir, filename)
            with open(filepath, 'wb') as f:
                for chunk in r.iter_content(8192):
                    f.write(chunk)
            return True
    except Exception as e:
        print(f"[download] error {url}: {e}")
    return False

import cv2
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMP_PHOTO_DIR = os.path.join(os.path.dirname(__file__), "temp_photos")
OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(TEMP_PHOTO_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)

def _mosaic_faces(image_path):
    """OpenCV로 얼굴 감지 후 모자이크 처리"""
    try:
        img = cv2.imread(image_path)
        if img is None:
            return False
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        faces = cascade.detectMultiScale(gray, 1.1, 4, minSize=(30, 30))
        for (x, y, w, h) in faces:
            roi = img[y:y+h, x:x+w]
            small = cv2.resize(roi, (max(1, w//10), max(1, h//10)))
            mosaic = cv2.resize(small, (w, h), interpolation=cv2.INTER_NEAREST)
            img[y:y+h, x:x+w] = mosaic
        cv2.imwrite(image_path, img)
        return len(faces) > 0
    except Exception as e:
        print(f"[mosaic] error: {e}")
        return False

def _collect_xhs_for_cafe(target, count):
    """타겟층 기반으로 샤오홍슈 이미지 수집 + 모자이크"""
    # 타겟층 → 중국어 검색 키워드 변환
    zh_query = _translate_ko_to_zh(target + " 일상")
    print(f"[xhs-cafe] target='{target}' → query='{zh_query}'")

    urls = _crawl_xhs_images(zh_query, count + 3)
    if not urls:
        print("[xhs-cafe] no images found, trying simpler query")
        zh_query = _translate_ko_to_zh(target)
        urls = _crawl_xhs_images(zh_query, count + 3)

    downloaded = []
    headers = {"User-Agent": "Mozilla/5.0", "Referer": "https://www.xiaohongshu.com/"}
    for i, url in enumerate(urls):
        if len(downloaded) >= count:
            break
        try:
            r = req.get(url, headers=headers, timeout=15)
            if r.status_code == 200 and len(r.content) > 2000:
                fname = f"xhs_cafe_{int(time.time()*1000)}_{i}.jpg"
                fpath = os.path.join(TEMP_PHOTO_DIR, fname)
                with open(fpath, 'wb') as f:
                    f.write(r.content)
                _mosaic_faces(fpath)
                downloaded.append(fpath)
        except Exception:
            pass
        time.sleep(2)  # 차단 방지
    return downloaded

def _create_cafe_docx(title, body_text, image_paths, keyword):
    """카페 원고 docx 생성"""
    doc = DocxDocument()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    img_idx = 0
    for line in body_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if '[어울릴 사진' in line or '[이미지' in line:
            if img_idx < len(image_paths) and os.path.exists(image_paths[img_idx]):
                try:
                    doc.add_picture(image_paths[img_idx], width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(line)
                img_idx += 1
            else:
                doc.add_paragraph(line)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_kw = re.sub(r'[\\/*?:"<>|]', '', keyword)
    fname = f"{safe_kw}_카페원고_{ts}.docx"
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return fpath, fname

@app.get("/api/photo/translate")
async def photo_translate(text: str):
    loop = asyncio.get_running_loop()
    zh = await loop.run_in_executor(executor, _translate_ko_to_zh, text)
    return {"zh": zh}

@app.post("/api/collect-images")
@app.post("/api/photo/crawl")
async def photo_crawl(request: Request):
    body = await request.json()
    query_zh = body.get("query_zh", "")
    count = body.get("count", 30)
    sources = body.get("sources", ["baidu"])
    category = body.get("category", "제품사진")

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        yield f"data: {json.dumps({'type':'progress','msg':'브라우저 시작 중...','cur':0,'total':0}, ensure_ascii=False)}\n\n"
        await _selenium_semaphore.acquire()
        driver = None
        try:
            driver = await loop.run_in_executor(executor, _create_driver)

            all_urls = []
            if "baidu" in sources:
                yield f"data: {json.dumps({'type':'progress','msg':f'바이두 이미지 검색: {query_zh}','cur':0,'total':count}, ensure_ascii=False)}\n\n"
                baidu_urls = await loop.run_in_executor(executor, _crawl_baidu_images, driver, query_zh, count)
                all_urls.extend([("baidu", u) for u in baidu_urls])

            if "xhs" in sources:
                yield f"data: {json.dumps({'type':'progress','msg':f'샤오홍슈 검색: {query_zh}','cur':0,'total':count}, ensure_ascii=False)}\n\n"
                xhs_urls = await loop.run_in_executor(executor, _crawl_xhs_images, query_zh, count)
                if not xhs_urls:
                    yield f"data: {json.dumps({'type':'progress','msg':'⚠️ 샤오홍슈 수집 실패 (로그인/캡챠 차단 확인)','cur':0,'total':count}, ensure_ascii=False)}\n\n"
                all_urls.extend([("xhs", u) for u in xhs_urls])
        finally:
            if driver:
                await loop.run_in_executor(executor, driver.quit)
            _selenium_semaphore.release()

        # 이미지 다운로드
        total = min(len(all_urls), count)
        downloaded = 0
        for i, (src, url) in enumerate(all_urls[:count]):
            ts = int(time.time() * 1000)
            ext = "jpg"
            if ".png" in url.lower():
                ext = "png"
            elif ".webp" in url.lower():
                ext = "webp"
            filename = f"{src}_{ts}_{i}.{ext}"

            yield f"data: {json.dumps({'type':'progress','msg':f'다운로드 중 ({i+1}/{total})','cur':i+1,'total':total}, ensure_ascii=False)}\n\n"

            ok = await loop.run_in_executor(executor, _download_image, url, filename, category)
            if ok:
                downloaded += 1
                # 프론트엔드 호환성을 위해 filename은 상대경로처럼 전달 (업데이트 필요)
                rel_filename = f"{category}/{filename}"
                yield f"data: {json.dumps({'type':'image','filename':rel_filename,'category':category}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.3)

        yield f"data: {json.dumps({'type':'complete','total':downloaded}, ensure_ascii=False)}\n\n"
      except Exception as e:
        print(f"[photo_crawl] 에러: {e}")
        yield f"data: {json.dumps({'type':'error','message':f'이미지 수집 중 오류: {e}'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.get("/api/photo/thumb/{filename:path}")
async def photo_thumb(filename: str):
    filepath = os.path.realpath(os.path.join(PHOTO_DIR, filename))
    if not filepath.startswith(os.path.realpath(PHOTO_DIR)):
        return Response(status_code=403)
    if not os.path.exists(filepath):
        return {"error": "not found"}
    return FileResponse(filepath, media_type="image/jpeg")

@app.get("/api/photo/image/{filename:path}")
async def photo_image(filename: str):
    filepath = os.path.realpath(os.path.join(PHOTO_DIR, filename))
    if not filepath.startswith(os.path.realpath(PHOTO_DIR)):
        return Response(status_code=403)
    if not os.path.exists(filepath):
        return {"error": "not found"}
    return FileResponse(filepath)

@app.post("/api/photo/mosaic")
async def photo_mosaic(request: Request):
    """간단한 모자이크 처리 (PIL 없이 — 파일명에 _mosaic 태그 추가)"""
    body = await request.json()
    filenames = body.get("filenames", [])
    processed = []
    for fn in filenames:
        filepath = os.path.join(PHOTO_DIR, fn)
        if os.path.exists(filepath):
            # PIL이 없으므로 파일 복사 + 이름에 _mosaic 추가
            name, ext = os.path.splitext(fn)
            new_fn = f"{name}_mosaic{ext}"
            new_path = os.path.join(PHOTO_DIR, new_fn)
            import shutil
            shutil.copy2(filepath, new_path)
            processed.append(new_fn)
    return {"processed": processed, "count": len(processed)}

@app.post("/api/photo/delete")
async def photo_delete(request: Request):
    body = await request.json()
    filenames = body.get("filenames", [])
    deleted = 0
    for fn in filenames:
        filepath = os.path.join(PHOTO_DIR, fn)
        if os.path.exists(filepath):
            os.remove(filepath)
            deleted += 1
    return {"deleted": deleted}

@app.post("/api/photo/save-library")
async def photo_save_library(request: Request):
    body = await request.json()
    items = body.get("items", [])
    keyword = body.get("keyword", "")

    # 기존 라이브러리 로드
    lib = []
    if os.path.exists(PHOTO_LIB_FILE):
        with open(PHOTO_LIB_FILE, 'r', encoding='utf-8') as f:
            lib = json.load(f)

    # 새 항목 추가
    for item in items:
        lib.append({
            "filename": item["filename"],
            "category": item.get("category", "기타"),
            "keyword": keyword,
            "saved_at": datetime.now().isoformat()
        })

    with open(PHOTO_LIB_FILE, 'w', encoding='utf-8') as f:
        json.dump(lib, f, ensure_ascii=False, indent=2)

    return {"saved": len(items), "total": len(lib)}

@app.get("/api/photo/library")
async def photo_library():
    if not os.path.exists(PHOTO_LIB_FILE):
        return {"items": []}
    with open(PHOTO_LIB_FILE, 'r', encoding='utf-8') as f:
        lib = json.load(f)
    # 실제 파일 존재 여부 확인
    lib = [item for item in lib if os.path.exists(os.path.join(PHOTO_DIR, item["filename"]))]
    return {"items": lib}

# ═══════════════════════════ AD CREATIVE AUTOMATION ═══════════════════════════

AD_REFS_DIR = os.path.join(os.path.dirname(__file__), "ad_refs")
AD_OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "ad_outputs")
os.makedirs(AD_REFS_DIR, exist_ok=True)
os.makedirs(AD_OUTPUTS_DIR, exist_ok=True)

def _crawl_meta_ads(keyword, country="KR", count=30, advertiser=""):
    """메타 광고 라이브러리 크롤링 (Selenium) — 개선된 버전"""
    results = []
    driver = _create_driver()
    try:
        q = advertiser if advertiser else keyword
        url = "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=%s&media_type=image&q=%s" % (country, quote(q))
        driver.get(url)
        time.sleep(8)  # 초기 로딩 대기

        # 쿠키/로그인 팝업 닫기
        try:
            for label in ['닫기', 'Close', 'Decline optional cookies', '선택 쿠키 거부']:
                btns = driver.find_elements(By.XPATH, "//button[contains(text(), '%s') or @aria-label='%s']" % (label, label))
                for b in btns:
                    try: b.click(); time.sleep(1)
                    except Exception: pass
        except Exception:
            pass

        # 충분히 스크롤해서 광고 로드
        last_height = 0
        for scroll_i in range(max(8, count // 3)):
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            time.sleep(3)
            new_height = driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height and scroll_i > 3:
                break
            last_height = new_height

        # ── 전략: "게재 시작" 또는 "Started running" 텍스트를 기준으로 광고 카드 찾기 ──
        # 이 텍스트는 모든 광고 카드에 반드시 존재
        date_markers = driver.find_elements(By.XPATH,
            "//span[contains(text(),'Started running') or contains(text(),'게재 시작') or contains(text(),'시작일')]")

        print("[meta-ads] Found %d date markers" % len(date_markers))

        seen_texts = set()
        for marker in date_markers:
            if len(results) >= count:
                break
            try:
                # 날짜 마커에서 상위로 올라가면서 광고 카드 컨테이너 찾기
                container = marker
                for _ in range(15):
                    container = container.find_element(By.XPATH, "..")
                    # 광고 카드는 보통 높이가 200px 이상이고, 이미지를 포함
                    try:
                        h = container.size.get('height', 0)
                        w = container.size.get('width', 0)
                    except Exception:
                        h = w = 0
                    if h > 200 and w > 300:
                        imgs = container.find_elements(By.TAG_NAME, "img")
                        if len(imgs) >= 1:
                            break

                # 중복 체크
                card_text = container.text[:200]
                if card_text in seen_texts or len(card_text) < 20:
                    continue
                seen_texts.add(card_text)

                ad = {'headline': '', 'body': '', 'cta': '', 'advertiser': '', 'period': '', 'image_url': '', 'image_file': ''}

                # ── 이미지 추출 ──
                imgs = container.find_elements(By.TAG_NAME, "img")
                for img in imgs:
                    src = img.get_attribute("src") or ""
                    # 광고 이미지는 scontent/fbcdn URL + 실제 크기가 큰 것
                    try:
                        nw = driver.execute_script("return arguments[0].naturalWidth || 0", img)
                    except Exception:
                        nw = 0
                    if src and ("scontent" in src or "fbcdn" in src) and nw > 80:
                        ad['image_url'] = src
                        break

                # ── 텍스트 추출 (줄 단위로 파싱) ──
                full_text = container.text
                lines = [l.strip() for l in full_text.split('\n') if l.strip() and len(l.strip()) > 2]
                # 페이지 UI 텍스트 제거
                skip_words = ['광고 라이브러리', 'Ad Library', 'API', '브랜디드 콘텐츠', 'See summary', '요약 보기',
                              'About this ad', '이 광고 정보', 'See ad details', '광고 세부정보', '모든 광고']
                lines = [l for l in lines if not any(sw in l for sw in skip_words)]

                if lines:
                    ad['advertiser'] = lines[0]

                # 게재 기간
                for l in lines:
                    if '시작' in l or 'Started' in l or 'running' in l:
                        ad['period'] = l
                        break

                # 본문 = 가장 긴 텍스트 (광고주, 기간 제외)
                content_lines = [l for l in lines if l != ad['advertiser'] and l != ad['period'] and len(l) > 5]
                if content_lines:
                    content_lines.sort(key=len, reverse=True)
                    ad['body'] = content_lines[0]
                    # 헤드라인 = 두번째로 긴 텍스트 또는 짧은 텍스트
                    headlines = [l for l in content_lines if l != ad['body'] and 5 < len(l) < 80]
                    if headlines:
                        ad['headline'] = headlines[0]

                # CTA
                cta_btns = container.find_elements(By.CSS_SELECTOR, "a[role='button'], div[role='button']")
                for btn in cta_btns:
                    bt = btn.text.strip()
                    if bt and 2 < len(bt) < 25 and bt not in ('더 보기', 'See more', '좋아요', 'Like'):
                        ad['cta'] = bt
                        break

                if ad['body'] or ad['headline']:
                    # 이미지 다운로드
                    if ad['image_url']:
                        fname = "ad_%d_%d.jpg" % (int(time.time()*1000), len(results))
                        fpath = os.path.join(AD_REFS_DIR, fname)
                        try:
                            r = req.get(ad['image_url'], timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
                            if r.status_code == 200 and len(r.content) > 2000:
                                with open(fpath, 'wb') as f:
                                    f.write(r.content)
                                ad['image_file'] = fname
                        except Exception:
                            pass
                    results.append(ad)
            except Exception as e2:
                print("[meta-ads] card parse error: %s" % e2)
                continue

        # ── 폴백: date_markers가 없으면 일반적인 방식 시도 ──
        if not results:
            print("[meta-ads] Fallback: trying generic approach")
            # 페이지 소스에서 이미지 URL 직접 추출
            page_src = driver.page_source
            img_urls = re.findall(r'https://scontent[^"\']+', page_src)
            img_urls = list(dict.fromkeys(img_urls))[:count]  # 중복 제거
            for img_url in img_urls:
                if 'emoji' in img_url or 'avatar' in img_url or len(img_url) > 500:
                    continue
                ad = {'headline': keyword, 'body': '', 'cta': '', 'advertiser': '', 'period': '', 'image_url': img_url, 'image_file': ''}
                fname = "ad_%d_%d.jpg" % (int(time.time()*1000), len(results))
                fpath = os.path.join(AD_REFS_DIR, fname)
                try:
                    r = req.get(img_url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
                    if r.status_code == 200 and len(r.content) > 5000:
                        with open(fpath, 'wb') as f:
                            f.write(r.content)
                        ad['image_file'] = fname
                        results.append(ad)
                except Exception:
                    pass
                if len(results) >= count:
                    break

    except Exception as e:
        print("[meta-ads] crawl error: %s" % e)
    finally:
        driver.quit()
    return results

def _analyze_ad_refs(refs):
    """레퍼런스 패턴 분석 (AI)"""
    ref_text = ""
    for i, r in enumerate(refs):
        ref_text += "[%d] 헤드라인: %s | 본문: %s | CTA: %s\n" % (i+1, r.get('headline',''), r.get('body','')[:100], r.get('cta',''))

    system = "당신은 광고 소재 분석 전문가입니다."
    user = """아래 수집된 광고 레퍼런스 %d개를 분석해주세요.

%s

분석 항목:
1. 가장 많이 쓰이는 소구점 Top 3
2. 후킹 방식 분류 (공포/호기심/공감/결과제시/긴급성)
3. CTA 패턴 분포
4. 평균 헤드라인 길이
5. 자주 쓰이는 표현/키워드
6. 추천 광고 전략 요약

JSON으로 응답해.""" % (len(refs), ref_text)

    return _call_claude(system, user)

def _generate_ad_creatives(product, appeal, refs_analysis, ref_examples, platform, ad_type, count, forbidden):
    """광고 소재 생성 (AI)"""
    system = "당신은 메타/틱톡 DA 광고 소재를 기획하는 퍼포먼스 마케터입니다."
    user = """아래 정보를 기반으로 클릭률 높은 광고 소재를 %d개 기획해주세요.

## 광고 플랫폼: %s
## 광고 유형: %s

## 제품 정보
- 제품명: %s
- 주요 성분: %s
- 핵심 특징(USP): %s
- 타겟층: %s
- 나만의 키워드: %s

## 소구점: %s

## 레퍼런스 분석 결과
%s

## 별점 높은 레퍼런스 예시
%s

%s

## 출력 형식 (JSON 배열)
각 소재를 아래 형식으로:
[
  {
    "headline": "헤드라인 (15자 이내)",
    "body": "본문 텍스트 (100자 이내)",
    "cta": "CTA 버튼 텍스트",
    "overlay_text": "이미지 위 텍스트 오버레이 (2줄)",
    "ref_index": 참고한 레퍼런스 번호
  }
]""" % (
        count, platform, ad_type,
        product.get('name',''), product.get('ingredients',''), product.get('usp',''),
        product.get('target',''), product.get('brand_keyword',''),
        appeal, refs_analysis, ref_examples,
        ("## 광고 금지 문구\n다음 표현 사용 금지: " + forbidden) if forbidden else ""
    )

    return _call_claude(system, user)

def _create_ad_image(bg_path, overlay_text, output_path):
    """배경 이미지에 광고 텍스트 오버레이"""
    try:
        img = Image.open(bg_path).convert("RGBA")
        img = img.resize((1080, 1080), Image.LANCZOS)

        overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)

        try:
            font = ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", 56)
        except Exception:
            font = ImageFont.load_default()

        lines = overlay_text.split('\n')
        y_pos = img.height - 200
        draw.rectangle([0, y_pos - 30, img.width, img.height], fill=(0, 0, 0, 160))

        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            x = (img.width - tw) // 2
            draw.text((x, y_pos), line, fill="white", font=font)
            y_pos += 80

        result = Image.alpha_composite(img, overlay).convert("RGB")
        result.save(output_path, quality=95)
        return True
    except Exception as e:
        print("[ad-image] error: %s" % e)
        return False

@app.post("/api/ad/crawl-refs")
async def ad_crawl_refs(request: Request):
    """메타 광고 라이브러리에서 레퍼런스 수집 (SSE)"""
    body = await request.json()
    keyword = body.get('keyword', '')
    country = body.get('country', 'KR')
    count = body.get('count', 30)
    advertiser = body.get('advertiser', '')

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        yield _sse({'type': 'progress', 'msg': '메타 광고 라이브러리 크롤링 중...'})
        await _selenium_semaphore.acquire()
        try:
            refs = await loop.run_in_executor(executor, _crawl_meta_ads, keyword, country, count, advertiser)
        finally:
            _selenium_semaphore.release()
        yield _sse({'type': 'complete', 'refs': refs, 'total': len(refs)})
      except Exception as e:
        print(f"[ad_crawl_refs] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'광고 레퍼런스 수집 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.get("/api/ad/ref-image/{filename}")
async def ad_ref_image(filename: str):
    safe_name = os.path.basename(filename)
    fpath = os.path.join(AD_REFS_DIR, safe_name)
    if os.path.exists(fpath):
        return FileResponse(fpath, media_type="image/jpeg")
    return Response(status_code=404)

@app.post("/api/ad/analyze")
async def ad_analyze(request: Request):
    """레퍼런스 패턴 분석"""
    body = await request.json()
    refs = body.get('refs', [])
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(executor, _analyze_ad_refs, refs)
    return {'analysis': result}

@app.post("/api/ad/generate")
async def ad_generate(request: Request):
    """광고 소재 생성"""
    body = await request.json()
    product = body.get('product', {})
    appeal = body.get('appeal', '')
    analysis = body.get('analysis', '')
    ref_examples = body.get('ref_examples', '')
    platform = body.get('platform', '메타')
    ad_type = body.get('ad_type', '이미지')
    count = body.get('count', 5)
    forbidden = body.get('forbidden', '')
    ref_images = body.get('ref_images', {})  # {index: filename}

    loop = asyncio.get_running_loop()
    raw = await loop.run_in_executor(executor, _generate_ad_creatives,
        product, appeal, analysis, ref_examples, platform, ad_type, count, forbidden)

    # JSON 파싱
    creatives = []
    try:
        m = re.search(r'\[[\s\S]*\]', raw)
        if m:
            creatives = json.loads(m.group())
    except Exception:
        creatives = [{'headline': raw[:50], 'body': raw, 'cta': '자세히 보기', 'overlay_text': '', 'ref_index': 0}]

    # 각 소재에 이미지 합성
    for i, cr in enumerate(creatives):
        cr['image_file'] = ''
        overlay = cr.get('overlay_text', '')
        ref_idx = cr.get('ref_index', 0)
        bg_file = ref_images.get(str(ref_idx), '') or ref_images.get(str(ref_idx - 1), '')
        if bg_file and overlay:
            bg_path = os.path.join(AD_REFS_DIR, bg_file)
            if os.path.exists(bg_path):
                out_fname = "creative_%s_%d.jpg" % (int(time.time()*1000), i)
                out_path = os.path.join(AD_OUTPUTS_DIR, out_fname)
                ok = await loop.run_in_executor(executor, _create_ad_image, bg_path, overlay, out_path)
                if ok:
                    cr['image_file'] = out_fname

    return {'creatives': creatives}

@app.get("/api/ad/output-image/{filename}")
async def ad_output_image(filename: str):
    safe_name = os.path.basename(filename)
    fpath = os.path.join(AD_OUTPUTS_DIR, safe_name)
    if os.path.exists(fpath):
        return FileResponse(fpath, media_type="image/jpeg")
    return Response(status_code=404)

@app.post("/api/ad/save-notion")
async def ad_save_notion(request: Request):
    body = await request.json()
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('headline', '')}}]},
        '채널': {'select': {'name': '메타광고'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    summary = body.get('body', '') + ' | CTA: ' + body.get('cta', '')
    props['본문'] = {'rich_text': [{'text': {'content': summary[:2000]}}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

@app.post("/api/ad/upload-product-image")
async def ad_upload_product_image(file: UploadFile = File(...)):
    """제품 이미지 업로드"""
    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp', '.gif'}
    ext = os.path.splitext(file.filename or '')[1].lower()
    if ext not in allowed_ext:
        return JSONResponse({'error': f'허용되지 않는 파일 형식: {ext}'}, status_code=400)
    fname = "product_%s_%s" % (int(time.time()*1000), file.filename)
    fpath = os.path.join(AD_REFS_DIR, fname)
    with open(fpath, 'wb') as f:
        content = await file.read()
        f.write(content)
    return {'filename': fname}

# ═══════════════════════════ POWER CONTENT ═══════════════════════════

def _crawl_power_content(url):
    """파워컨텐츠 랜딩페이지 크롤링"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'}
        r = req.get(url, headers=headers, timeout=15, allow_redirects=True)
        soup = BeautifulSoup(r.text, 'html.parser')
        # 본문 추출 시도
        for sel in ['.se-main-container', '.content_view', '#content', '.article_body', '.post_ct', 'article', 'main']:
            el = soup.select_one(sel)
            if el and len(el.get_text(strip=True)) > 200:
                return el.get_text('\n', strip=True)[:10000]
        # 폴백: body 전체
        body = soup.find('body')
        if body:
            for tag in body.find_all(['script','style','nav','header','footer']):
                tag.decompose()
            return body.get_text('\n', strip=True)[:10000]
        return ''
    except Exception as e:
        print("[power-content] crawl error: %s" % e)
        return ''

def _build_pc_analysis_prompt(ref_text):
    system = "당신은 파워컨텐츠(설득형 랜딩페이지) 구조 분석 전문가입니다."
    user = """아래 파워컨텐츠 본문을 분석해.

%s

분석 항목:
1. total_chars: 전체 글자수 (공백 포함)
2. sections: 오프닝(1막)/솔루션(2막)/클로징(3막) 각각의 글자수와 비율
3. opening_method: 오프닝 전개 방식
4. solution_method: 솔루션 전개 방식 + 소거 개수
5. closing_method: 클로징 전개 방식
6. image_count: 이미지 삽입 위치와 개수
7. cta_count: CTA 위치와 개수
8. hooking_type: 후킹 방식 (부정편향/비밀폭로/비교검증)
9. keyword_repeats: 키워드 반복 횟수 추정
10. persuasion_flow: 전체 설득 흐름 (한 줄 요약)

JSON으로 응답해.""" % ref_text
    return system, user

def _build_pc_ad_prompt(keyword, appeal, buying_thing, product, hooking_type, forbidden):
    system = "당신은 네이버 파워컨텐츠 CTR 최적화 전문가입니다."
    user = """키워드와 소구점을 기반으로 네이버 파워컨텐츠 광고 소재를 생성해.

## 입력
- 메인 키워드: %s
- 소구점: %s
- 구매원씽: %s
- 제품 정보: 제품명=%s, 성분=%s, USP=%s
- 레퍼런스 후킹 방식: %s

## 출력 규격
1. 제목: 최대 28자 (공백 포함), 메인 키워드 포함
2. 설명: 최대 110자 (공백 포함), 궁금증 유발 + 혜택/근거/가치

제목 유형 (레퍼런스와 같은 유형):
- 부정편향: "절대 사면 안 되는 OOO", "효과 없었던 진짜 이유"
- 비밀폭로: "업계가 숨기는 OOO", "전문가는 안 알려주는"
- 비교검증: "3개 써보고 솔직 비교", "성분 분석해봤더니"

출력 형식:
제목: (제목)
설명: (설명)

%s
금지: 허위·과장, 근거 없는 효능, 낚시""" % (
        keyword, appeal, buying_thing,
        product.get('name',''), product.get('ingredients',''), product.get('usp',''),
        hooking_type, forbidden
    )
    return system, user

def _build_pc_body_prompt(keyword, stage, appeal, buying_thing, deficit_level, product, ad_title, ad_desc, analysis_json):
    system = "당신은 네이버 파워컨텐츠 랜딩 본문 전문 카피라이터입니다. 멘토 3단 공식 + BA 설득 기법으로 작성합니다."

    if stage in ('0_무지','1_인지','2_호기심','3_정보습득'):
        template = """[템플릿 A — 구매여정 0~3]

1막 오프닝: 문제를 '심각한 질병'으로 격상
- BA 카모플라주: 건강정보/전문가 분석 포맷으로 진입
- BA 점진화: 독자가 동의하는 증상에서 출발, 동의 3~5회 축적
- 결핍 마케팅: 방치 시 미래를 구체적으로 묘사
- 권위: "해외 논문", "전문가 발견" 등

2막 솔루션: 기존 상식 파괴 + 새 길 제시
- BA 집중화: 좌절 서사로 경쟁 경로 차단
- 소거법: 기존 대안 소거 → 우리 제품만 남김
- 성분 비교표 등으로 새 기준 증명

3막 클로징: 첫 번째 선택지 되기
- BA 메커니제이션: 작동 원리를 판매 언어로
- BA 재정의 가격축소: "하루 OOO원"
- CTA: "현명한 첫걸음을 내딛으세요" + [CTA: %s]""" % product.get('url','')
    else:
        template = """[템플릿 B — 구매여정 4~5]

1막 오프닝: 잘못된 선택의 위험성 경고
- BA 카모플라주: "직접 써보고 비교한 분석가" 포맷
- BA 점진화: 효과 없는 제품에 돈/시간 낭비가 진짜 문제

2막 솔루션: 경쟁자 압살
- BA 집중화: 경쟁사별 구체적 약점
- 소거법: A사 핵심 성분 없음, B사 함량 미달 → 우리만 충족

3막 클로징: 즉시 구매 유도
- FOMO + 가성비 증명
- CTA: "계속 돈을 낭비하시겠습니까?" + [CTA: %s]""" % product.get('url','')

    user = """아래 조건으로 파워컨텐츠 랜딩 본문을 작성해.

## 입력
- 광고 소재: 제목=%s / 설명=%s
- 메인 키워드: %s
- 구매여정: %s
- 소구점: %s / 구매원씽: %s / 결핍강도 목표: %s
- 제품: 이름=%s, 성분=%s, USP=%s, 타겟=%s, URL=%s

## 레퍼런스 구조 (이 구조를 그대로 따르되 내용만 교체)
%s

## 템플릿
%s

## 충돌방지 3규칙
1. 결핍강도 -3~-5, 한 문단 -2 이상 급락 금지
2. BA 재정의 우선, 공략집 프레이밍은 참고만
3. 섹션 순서 = 멘토 3단 공식 고정

## 공통 규칙
- 레퍼런스 전체 글자수와 비슷하게 (최소 2,500자, ±200자)
- 각 막 비율도 레퍼런스와 유사
- 메인 키워드 5~8회 자연스럽게 삽입
- 이미지/CTA 위치·개수도 레퍼런스와 동일
- [이미지N - 설명], [CTA: URL], [공정위 문구 삽입 위치] 표기

## 출력: 본문만 출력 (부연 설명 없이)""" % (
        ad_title, ad_desc, keyword, stage, appeal, buying_thing, deficit_level,
        product.get('name',''), product.get('ingredients',''), product.get('usp',''),
        product.get('target',''), product.get('url',''),
        analysis_json, template
    )
    return system, user

@app.post("/api/powercontent/analyze")
async def pc_analyze(request: Request):
    """레퍼런스 크롤링 + 구조 분석"""
    body = await request.json()
    url = body.get('url', '')
    raw_text = body.get('raw_text', '')
    loop = asyncio.get_running_loop()

    if url and not raw_text:
        raw_text = await loop.run_in_executor(executor, _crawl_power_content, url)
    if not raw_text:
        return {'error': '본문을 가져올 수 없습니다', 'analysis': ''}

    sys_p, usr_p = _build_pc_analysis_prompt(raw_text)
    analysis = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p)
    return {'raw_text': raw_text[:5000], 'analysis': analysis}

@app.post("/api/powercontent/generate")
async def pc_generate(request: Request):
    """파워컨텐츠 생성 (SSE): 구조분석 완료 후 → 광고소재 → 본문"""
    body = await request.json()
    keyword = body.get('keyword', '')
    stage = body.get('stage', '3_정보습득')
    appeal = body.get('appeal', '')
    buying_thing = body.get('buying_thing', '')
    deficit_level = body.get('deficit_level', '-4')
    product = body.get('product', {})
    analysis = body.get('analysis', '')
    forbidden = product.get('forbidden', '')

    # 구조분석에서 hooking_type 추출
    hooking_type = '부정편향'
    try:
        m = re.search(r'"hooking_type"\s*:\s*"([^"]+)"', analysis)
        if m: hooking_type = m.group(1)
    except (AttributeError, ValueError): pass

    # 구조분석에서 total_chars 추출
    target_chars = 3000
    try:
        m = re.search(r'"total_chars"\s*:\s*(\d+)', analysis)
        if m: target_chars = max(2500, int(m.group(1)))
    except (AttributeError, ValueError): pass

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()

        # STEP 2: 광고 소재 생성
        yield _sse({'type':'progress','msg':'광고 소재 생성 중 (제목 28자 + 설명 110자)...'})
        sys2, usr2 = _build_pc_ad_prompt(keyword, appeal, buying_thing, product, hooking_type, forbidden)
        ad_raw = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
        ad_title = ad_desc = ''
        for line in ad_raw.split('\n'):
            line = line.strip()
            if line.startswith('제목:') or line.startswith('제목 :'):
                ad_title = line.split(':', 1)[1].strip()
            elif line.startswith('설명:') or line.startswith('설명 :'):
                ad_desc = line.split(':', 1)[1].strip()
        if not ad_title:
            ad_title = ad_raw.split('\n')[0].strip()[:28]

        yield _sse({'type':'ad','title':ad_title,'desc':ad_desc})

        # STEP 3: 랜딩 본문 생성 (글자수 검증 포함, 최대 3회 재시도)
        yield _sse({'type':'progress','msg':'랜딩 본문 생성 중 (목표 %d자)...' % target_chars})
        sys3, usr3 = _build_pc_body_prompt(keyword, stage, appeal, buying_thing, deficit_level, product, ad_title, ad_desc, analysis)
        body_text = ''
        for attempt in range(3):
            body_text = await loop.run_in_executor(executor, _call_claude, sys3, usr3)
            char_count = len(body_text)
            if char_count >= 2500 and abs(char_count - target_chars) <= 400:
                break
            # 재시도
            yield _sse({'type':'progress','msg':'글자수 부족 (%d자). 재생성 중 (%d/3)...' % (char_count, attempt+2)})
            usr3 = "글자수가 부족합니다. 현재 %d자입니다. %d자 이상으로 다시 작성해주세요. 이전 내용의 구조를 유지하면서 각 단락을 더 풍부하게 확장해주세요.\n\n" % (char_count, target_chars) + usr3

        char_count = len(body_text)
        yield _sse({'type':'result','ad_title':ad_title,'ad_desc':ad_desc,'body':body_text,'char_count':char_count,'target_chars':target_chars})
        yield _sse({'type':'complete'})
      except Exception as e:
        print(f"[powercontent_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'파워컨텐츠 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/powercontent/docx")
async def pc_docx(request: Request):
    """파워컨텐츠 docx 생성"""
    body = await request.json()
    ad_title = body.get('ad_title', '')
    ad_desc = body.get('ad_desc', '')
    body_text = body.get('body', '')
    keyword = body.get('keyword', '')

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 광고 소재
    h = doc.add_heading('파워컨텐츠 광고 소재', level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run('제목: ' + ad_title)
    r1.bold = True
    r1.font.size = Pt(14)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('설명: ' + ad_desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph('─' * 40)
    doc.add_heading('랜딩 본문', level=1)

    for line in body_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('[이미지') or line.startswith('[CTA') or line.startswith('[공정위'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 120, 200)
                run.italic = True
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_kw = re.sub(r'[\\/*?:"<>|]', '', keyword)
    fname = "%s_파워컨텐츠_%s.docx" % (safe_kw, ts)
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return FileResponse(fpath, filename=fname, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/api/powercontent/save-notion")
async def pc_save_notion(request: Request):
    body = await request.json()
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    props = {
        '제목': {'title': [{'text': {'content': body.get('ad_title', '')}}]},
        '채널': {'select': {'name': '파워컨텐츠'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    summary = body.get('body', '')[:300]
    props['본문'] = {'rich_text': [{'text': {'content': summary}}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    # 본문을 페이지 children으로
    body_text = body.get('body', '')
    if body_text:
        paras = [p.strip() for p in body_text.split('\n\n') if p.strip()]
        children = []
        for para in paras[:100]:
            children.append({'object':'block','type':'paragraph','paragraph':{'rich_text':[{'type':'text','text':{'content':para[:2000]}}]}})
        payload['children'] = children
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ═══════════════════════════ DEPLOY SCHEDULE ═══════════════════════════

def _notion_query_by_date(date_str, channel=None):
    """노션 콘텐츠 DB에서 특정 날짜의 콘텐츠 집계"""
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    filters = [{'property': '생성일', 'date': {'equals': date_str}}]
    if channel:
        filters.append({'property': '채널', 'select': {'equals': channel}})
    payload = {'filter': {'and': filters}, 'page_size': 100}
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % CONTENT_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code == 200:
            return r.json().get('results', [])
        return []
    except Exception:
        return []

def _count_by_channel(results):
    """노션 결과를 채널별 카운트"""
    counts = {}
    for page in results:
        props = page.get('properties', {})
        ch = props.get('채널', {}).get('select', {})
        ch_name = ch.get('name', '기타') if ch else '기타'
        counts[ch_name] = counts.get(ch_name, 0) + 1
    return counts

@app.get("/api/schedule/today")
async def schedule_today():
    """오늘 생산 현황"""
    today = datetime.now().strftime('%Y-%m-%d')
    loop = asyncio.get_running_loop()
    results = await loop.run_in_executor(executor, _notion_query_by_date, today)
    counts = _count_by_channel(results)
    return {'date': today, 'counts': counts, 'total': len(results)}

@app.get("/api/schedule/week")
async def schedule_week():
    """이번 주 월~금 일별 생산 현황"""
    today = datetime.now().date()
    # 이번 주 월요일 찾기
    monday = today - timedelta(days=today.weekday())
    loop = asyncio.get_running_loop()
    week_data = []
    for i in range(7):
        d = monday + timedelta(days=i)
        ds = d.isoformat()
        results = await loop.run_in_executor(executor, _notion_query_by_date, ds)
        counts = _count_by_channel(results)
        week_data.append({'date': ds, 'weekday': ['월','화','수','목','금','토','일'][i], 'counts': counts, 'total': len(results)})
    return {'week': week_data, 'today': today.isoformat()}

# ═══════════════════════════ WEEKLY REPORT ═══════════════════════════

def _notion_query_range(start_date, end_date, db_id=None):
    """노션 DB에서 기간 내 콘텐츠 조회"""
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    payload = {
        'filter': {'and': [
            {'property': '생성일', 'date': {'on_or_after': start_date}},
            {'property': '생성일', 'date': {'on_or_before': end_date}},
        ]},
        'page_size': 100
    }
    target_db = db_id or CONTENT_DB_ID
    all_results = []
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % target_db, headers=headers, json=payload, timeout=20)
        if r.status_code == 200:
            data = r.json()
            all_results.extend(data.get('results', []))
            while data.get('has_more'):
                payload['start_cursor'] = data['next_cursor']
                r = req.post('https://api.notion.com/v1/databases/%s/query' % target_db, headers=headers, json=payload, timeout=20)
                if r.status_code != 200:
                    break
                data = r.json()
                all_results.extend(data.get('results', []))
    except Exception as e:
        print("[report] query error: %s" % e)
    return all_results

def _build_report_data(start, end, results):
    """리포트 데이터 구축"""
    channels = ['블로그','카페','지식인','카페바이럴','커뮤니티','파워컨텐츠','유튜브','틱톡','쓰레드','메타광고']
    production = {}
    published = {}
    for page in results:
        props = page.get('properties', {})
        ch_sel = props.get('채널', {}).get('select')
        ch = ch_sel.get('name', '기타') if ch_sel else '기타'
        production[ch] = production.get(ch, 0) + 1
        deploy_sel = props.get('발행_상태', {}).get('select')
        deploy = deploy_sel.get('name', '') if deploy_sel else ''
        if deploy == '발행완료' or deploy == '발행완료':
            published[ch] = published.get(ch, 0) + 1
    rows = []
    total_prod = total_pub = 0
    for ch in channels:
        p = production.get(ch, 0)
        pub = published.get(ch, 0)
        unpub = p - pub
        rate = round(pub / p * 100) if p > 0 else 0
        if p > 0:
            rows.append({'channel': ch, 'produced': p, 'published': pub, 'unpublished': unpub, 'rate': rate})
            total_prod += p
            total_pub += pub
    total_rate = round(total_pub / total_prod * 100) if total_prod > 0 else 0
    return {
        'period': '%s ~ %s' % (start, end),
        'rows': rows,
        'total': {'produced': total_prod, 'published': total_pub, 'unpublished': total_prod - total_pub, 'rate': total_rate}
    }

@app.post("/api/report/generate")
async def report_generate(request: Request):
    """주간 리포트 생성"""
    body = await request.json()
    start = body.get('start', '')
    end = body.get('end', '')
    loop = asyncio.get_running_loop()

    # 현재 기간 데이터
    results = await loop.run_in_executor(executor, _notion_query_range, start, end)
    report = _build_report_data(start, end, results)

    # 이전 기간 데이터 (같은 길이만큼 이전)
    from datetime import date as dt_date
    d_start = dt_date.fromisoformat(start)
    d_end = dt_date.fromisoformat(end)
    period_days = (d_end - d_start).days + 1
    prev_end = d_start - timedelta(days=1)
    prev_start = prev_end - timedelta(days=period_days - 1)
    prev_results = await loop.run_in_executor(executor, _notion_query_range, prev_start.isoformat(), prev_end.isoformat())
    prev_report = _build_report_data(prev_start.isoformat(), prev_end.isoformat(), prev_results)

    # 증감 계산
    for row in report['rows']:
        prev_row = next((r for r in prev_report['rows'] if r['channel'] == row['channel']), None)
        if prev_row:
            row['prev_produced'] = prev_row['produced']
            row['prev_published'] = prev_row['published']
            row['diff_produced'] = row['produced'] - prev_row['produced']
            row['diff_published'] = row['published'] - prev_row['published']
        else:
            row['prev_produced'] = 0
            row['prev_published'] = 0
            row['diff_produced'] = row['produced']
            row['diff_published'] = row['published']
    report['prev_total'] = prev_report['total']

    return report

@app.post("/api/report/ai-actions")
async def report_ai_actions(request: Request):
    """AI 추천 액션 생성"""
    body = await request.json()
    report_data = body.get('report', {})
    viral_data = body.get('viral', [])

    summary = "채널별 성과:\n"
    for row in report_data.get('rows', []):
        summary += "- %s: 생산 %d, 발행 %d, 달성률 %d%%\n" % (row['channel'], row['produced'], row['published'], row['rate'])
    summary += "\n총 생산: %d, 발행: %d, 달성률: %d%%\n" % (
        report_data.get('total', {}).get('produced', 0),
        report_data.get('total', {}).get('published', 0),
        report_data.get('total', {}).get('rate', 0)
    )
    if viral_data:
        summary += "\n카페바이럴 현황:\n"
        for v in viral_data:
            summary += "- %s: 1단계 %s, 완료=%s\n" % (v.get('cafe',''), v.get('date1',''), v.get('done3', False))

    system = "당신은 마케팅 성과 분석가입니다."
    user = """아래 주간 성과 데이터를 분석해서 다음 주 액션 3~5개를 제안해줘.

%s

제안 기준:
1. 노출률/달성률 낮은 채널 → 원인 분석 + 개선 방안
2. 미발행 콘텐츠 → 빠른 발행 독촉
3. 성과 좋은 채널 → 확장 제안
4. 바이럴 예정일 경과 → 발행 알림
5. 전체 생산성 개선 아이디어

구체적이고 실행 가능한 액션으로. 막연한 조언 금지.""" % summary

    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(executor, _call_claude, system, user)
    return {'actions': result}

@app.post("/api/report/docx")
async def report_docx(request: Request):
    """주간 리포트 docx 생성"""
    body = await request.json()
    report = body.get('report', {})
    actions = body.get('actions', '')

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading('주간 성과 리포트', level=0)
    doc.add_paragraph('기간: ' + report.get('period', ''))

    # 생산 현황 테이블
    doc.add_heading('생산 현황', level=1)
    rows_data = report.get('rows', [])
    if rows_data:
        table = doc.add_table(rows=len(rows_data)+2, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(['채널','생산','발행','미발행','달성률']):
            hdr[i].text = h
        for i, row in enumerate(rows_data):
            cells = table.rows[i+1].cells
            cells[0].text = row['channel']
            cells[1].text = str(row['produced'])
            cells[2].text = str(row['published'])
            cells[3].text = str(row['unpublished'])
            cells[4].text = '%d%%' % row['rate']
        # 합계
        t = report.get('total', {})
        last = table.rows[-1].cells
        last[0].text = '합계'
        last[1].text = str(t.get('produced', 0))
        last[2].text = str(t.get('published', 0))
        last[3].text = str(t.get('unpublished', 0))
        last[4].text = '%d%%' % t.get('rate', 0)

    # AI 추천 액션
    if actions:
        doc.add_heading('추천 액션', level=1)
        doc.add_paragraph(actions)

    period = report.get('period', '').replace(' ~ ', '_').replace('/', '-')
    fname = '주간리포트_%s.docx' % period
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return FileResponse(fpath, filename=fname, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ═══════════════════════════ CAFE AUTO COMMENT (카페 댓글 자동 등록) ═══════════════════════════

@app.post("/api/cafe/auto-comment")
async def cafe_auto_comment(request: Request):
    """카페 댓글 자동 등록 (SSE 스트리밍)"""
    body = await request.json()
    post_url = body.get('post_url', '')
    comments = body.get('comments', [])
    account_ids = body.get('account_ids', [])

    if not post_url or not comments:
        return JSONResponse({'error': '게시글 URL과 댓글 필요'}, 400)

    naver_accounts = _naver_load_accounts()
    accounts = []
    for acc_id in account_ids:
        acc = next((a for a in naver_accounts if a.get('id') == acc_id), None)
        if acc:
            accounts.append(acc)

    if len(accounts) < len(comments):
        return JSONResponse({'error': f'계정 부족: 댓글 {len(comments)}개인데 계정 {len(accounts)}개'}, 400)

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_event_loop()
        total = len(comments)
        from src.cafe_comment_bot import run_auto_comments

        yield _sse({'type': 'progress', 'msg': f'댓글 {total}개 자동 등록 시작...', 'cur': 0, 'total': total})

        results = await loop.run_in_executor(
            executor, run_auto_comments,
            post_url, comments, accounts[:len(comments)], None, None
        )

        for i, r in enumerate(results):
            status = '✅' if r.get('success') else '❌'
            yield _sse({'type': 'result', 'data': r, 'cur': i+1, 'total': total,
                         'msg': f'[{i+1}/{total}] {r.get("label","")} — {status} {r.get("error","")}'})
            if r.get('status_change') == '정지 의심':
                accs = _naver_load_accounts()
                for a in accs:
                    if a['id'] == r.get('account_id'):
                        a['status'] = '정지 의심'
                        break
                _naver_save_accounts(accs)
            if r.get('success'):
                accs = _naver_load_accounts()
                for a in accs:
                    if a['id'] == r.get('account_id'):
                        a['total_posts'] = a.get('total_posts', 0) + 1
                        a['last_used_at'] = datetime.now().isoformat()
                        break
                _naver_save_accounts(accs)

        success_count = sum(1 for r in results if r.get('success'))
        yield _sse({'type': 'complete', 'total': total, 'success': success_count, 'fail': total - success_count})
      except Exception as e:
        print(f"[cafe_auto_comment] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'댓글 자동 등록 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.get("/api/cafe/comment-history")
async def cafe_comment_history():
    from src.cafe_safety_rules import get_history
    return {'history': get_history(50)}

# ═══════════════════════════ PROMPT TESTER (프롬프트 테스트) ═══════════════════════════

PROMPT_OVERRIDES_FILE = os.path.join(os.path.dirname(__file__), "prompt_overrides.json")

def _prompt_load_overrides():
    if os.path.exists(PROMPT_OVERRIDES_FILE):
        try:
            with open(PROMPT_OVERRIDES_FILE, encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError):
            pass
    return {}

def _prompt_save_overrides(data):
    with open(PROMPT_OVERRIDES_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# 채널별 기본 프롬프트 매핑
def _get_default_prompt(channel):
    """채널별 기본 시스템 프롬프트 반환"""
    dummy_product = {'name':'테스트제품','brand_keyword':'테스트키워드','usp':'핵심특징','target':'타겟층','ingredients':'성분'}
    if channel == '블로그_제목':
        sys_p, _ = _build_blog_title_prompt('테스트', dummy_product)
        return sys_p
    elif channel == '블로그_본문':
        sys_p, _ = _build_blog_body_prompt('테스트', '', dummy_product, 10, 5, '테스트 제목')
        return sys_p
    elif channel == '카페SEO_제목':
        sys_p, _ = _build_cafe_title_prompt('테스트', '원본제목')
        return sys_p
    elif channel == '카페SEO_본문':
        sys_p, _ = _build_cafe_body_prompt('테스트', '제목', '', {}, dummy_product)
        return sys_p
    elif channel == '카페SEO_댓글':
        sys_p, _ = _build_cafe_comments_prompt('테스트', '본문', '브랜드', '')
        return sys_p
    elif channel == '지식인_질문제목':
        sys_p, _ = _build_jisikin_title_prompt('테스트', dummy_product)
        return sys_p
    elif channel == '지식인_질문본문':
        sys_p, _ = _build_jisikin_body_prompt('테스트', dummy_product)
        return sys_p
    elif channel == '지식인_답변':
        sys_p, _ = _build_jisikin_answers_prompt('테스트', '질문제목', '질문본문', dummy_product)
        return sys_p
    elif channel == '틱톡':
        sys_p, _ = _build_tiktok_prompt('테스트', '소구점', '구매원씽', dummy_product, '')
        return sys_p
    elif channel == '커뮤니티':
        sys_p, _ = _build_community_post_prompt('뽐뿌', '1', '테스트', '소구점', '구매원씽', dummy_product, '')
        return sys_p
    elif channel == '쓰레드_일상글':
        sys_p, _ = _build_threads_daily_prompt({'name':'테스트','age':'30','job':'직장인','tone':'친근','interests':['건강']}, [])
        return sys_p
    elif channel == '쓰레드_물길글_셔플':
        sys_p, _ = _build_threads_traffic_prompt('테스트', {'tone':'친근','job':'직장인','interests':['건강']}, dummy_product, '', 'shuffle')
        return sys_p
    elif channel == '쓰레드_물길글_연민':
        sys_p, _ = _build_threads_traffic_prompt('테스트', {'tone':'친근'}, dummy_product, '', 'sympathy')
        return sys_p
    elif channel == '쓰레드_물길글_후기':
        sys_p, _ = _build_threads_traffic_prompt('테스트', {'tone':'친근','job':'직장인','interests':['건강']}, dummy_product, '', 'review')
        return sys_p
    elif channel == '유튜브댓글':
        sys_p, _ = _build_youtube_comment_prompt('테스트 영상', '더보기 내용', '테스트키워드')
        return sys_p
    elif channel == '카페바이럴_일상글':
        sys_p, _ = _build_viral_stage1_prompt('타겟층', '타겟층', '일상 주제')
        return sys_p
    elif channel == '카페바이럴_고민글':
        sys_p, _ = _build_viral_stage2_prompt('타겟층', '고민키워드', '건강기능식품')
        return sys_p
    elif channel == '카페바이럴_침투글':
        sys_p, _ = _build_viral_stage3_prompt('타겟층', '고민키워드', '테스트키워드', '테스트제품', '핵심특징', '성분', '건강기능식품')
        return sys_p
    return ''

@app.get("/api/prompt-test/channels")
async def prompt_test_channels():
    """테스트 가능한 채널 목록"""
    channels = ['블로그_제목','블로그_본문','카페SEO_제목','카페SEO_본문','카페SEO_댓글','지식인_질문제목','지식인_질문본문','지식인_답변','유튜브댓글','틱톡','커뮤니티','카페바이럴_일상글','카페바이럴_고민글','카페바이럴_침투글','쓰레드_일상글','쓰레드_물길글_셔플','쓰레드_물길글_연민','쓰레드_물길글_후기']
    return {'channels': channels}

@app.get("/api/prompt-test/get")
async def prompt_test_get(channel: str = ''):
    """채널의 현재 프롬프트 반환 (오버라이드 있으면 오버라이드, 없으면 기본)"""
    overrides = _prompt_load_overrides()
    if channel in overrides:
        return {'prompt': overrides[channel], 'is_override': True}
    default = _get_default_prompt(channel)
    return {'prompt': default, 'is_override': False}

@app.post("/api/prompt-test/generate")
async def prompt_test_generate(request: Request):
    """커스텀 프롬프트로 테스트 생성"""
    body = await request.json()
    system_prompt = body.get('system_prompt', '')
    keyword = body.get('keyword', '테스트')
    product = body.get('product', {})
    temperature = body.get('temperature', 0.7)
    if not system_prompt:
        return JSONResponse({'error': '시스템 프롬프트 필요'}, 400)

    # user 프롬프트 구성
    user_prompt = f"키워드: {keyword}\n제품명: {product.get('name','')}\n나만의 키워드: {product.get('brand_keyword','')}\n핵심 특징: {product.get('usp','')}\n타겟층: {product.get('target','')}\n주요 성분: {product.get('ingredients','')}"

    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(executor, _call_claude, system_prompt, user_prompt, temperature)
    return {'result': result, 'char_count': len(result)}

@app.post("/api/prompt-test/save")
async def prompt_test_save(request: Request):
    """수정된 프롬프트를 오버라이드로 저장"""
    body = await request.json()
    channel = body.get('channel', '')
    prompt = body.get('prompt', '')
    if not channel or not prompt:
        return JSONResponse({'error': '채널과 프롬프트 필요'}, 400)
    overrides = _prompt_load_overrides()
    overrides[channel] = prompt
    _prompt_save_overrides(overrides)
    return {'ok': True, 'channel': channel}

@app.post("/api/prompt-test/reset")
async def prompt_test_reset(request: Request):
    """오버라이드 삭제 (기본 프롬프트로 복원)"""
    body = await request.json()
    channel = body.get('channel', '')
    overrides = _prompt_load_overrides()
    if channel in overrides:
        del overrides[channel]
        _prompt_save_overrides(overrides)
    return {'ok': True}

# ═══════════════════════════ NAVER ACCOUNTS (네이버 계정 관리) ═══════════════════════════

NAVER_ACCOUNTS_FILE = os.path.join(os.path.dirname(__file__), "naver_accounts.json")
import threading as _threading_naver
_naver_accounts_lock = _threading_naver.Lock()

def _naver_load_accounts():
    if os.path.exists(NAVER_ACCOUNTS_FILE):
        try:
            with open(NAVER_ACCOUNTS_FILE, encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError) as e:
            print(f"[naver_accounts] 로드 오류: {e}")
    return []

def _naver_save_accounts(accounts):
    with _naver_accounts_lock:
        tmp = NAVER_ACCOUNTS_FILE + '.tmp'
        with open(tmp, 'w', encoding='utf-8') as f:
            json.dump(accounts, f, ensure_ascii=False, indent=2)
        os.replace(tmp, NAVER_ACCOUNTS_FILE)

@app.get("/api/naver/accounts")
async def naver_accounts_list():
    accounts = _naver_load_accounts()
    safe = []
    for acc in accounts:
        safe.append({
            'id': acc.get('id', ''),
            'label': acc.get('label', ''),
            'naver_id': acc.get('naver_id', ''),
            'platform': acc.get('platform', ''),
            'purpose': acc.get('purpose', ''),
            'active_cafes': acc.get('active_cafes', []),
            'proxy': acc.get('proxy', ''),
            'status': acc.get('status', '활성'),
            'daily_limit': acc.get('daily_limit', 3),
            'min_interval_hours': acc.get('min_interval_hours', 2),
            'active_hours': acc.get('active_hours', [9, 22]),
            'total_posts': acc.get('total_posts', 0),
            'last_used_at': acc.get('last_used_at', ''),
            'created_at': acc.get('created_at', ''),
            'notes': acc.get('notes', ''),
        })
    return {'accounts': safe}

@app.post("/api/naver/accounts")
async def naver_accounts_add(request: Request):
    body = await request.json()
    accounts = _naver_load_accounts()
    acc_id = 'naver_' + str(uuid.uuid4())[:8]
    account = {
        'id': acc_id,
        'label': body.get('label', ''),
        'naver_id': body.get('naver_id', ''),
        'platform': body.get('platform', '블로그'),
        'purpose': body.get('purpose', '최적화'),
        'active_cafes': body.get('active_cafes', []),
        'proxy': body.get('proxy', ''),
        'status': body.get('status', '활성'),
        'daily_limit': body.get('daily_limit', 3),
        'min_interval_hours': body.get('min_interval_hours', 2),
        'active_hours': body.get('active_hours', [9, 22]),
        'total_posts': 0,
        'last_used_at': None,
        'created_at': datetime.now().strftime('%Y-%m-%d'),
        'notes': body.get('notes', ''),
    }
    accounts.append(account)
    _naver_save_accounts(accounts)
    return {'ok': True, 'id': acc_id}

@app.patch("/api/naver/accounts/{acc_id}")
async def naver_accounts_update(acc_id: str, request: Request):
    body = await request.json()
    accounts = _naver_load_accounts()
    for acc in accounts:
        if acc['id'] == acc_id:
            for key in ['label', 'naver_id', 'platform', 'purpose', 'active_cafes', 'proxy', 'status', 'daily_limit', 'min_interval_hours', 'active_hours', 'notes']:
                if key in body:
                    acc[key] = body[key]
            _naver_save_accounts(accounts)
            return {'ok': True}
    return JSONResponse({'ok': False, 'error': '계정 없음'}, 404)

@app.delete("/api/naver/accounts/{acc_id}")
async def naver_accounts_delete(acc_id: str):
    accounts = _naver_load_accounts()
    accounts = [a for a in accounts if a.get('id') != acc_id]
    _naver_save_accounts(accounts)
    return {'ok': True}

# ═══════════════════════════ BATCH GENERATE (일괄 생성) ═══════════════════════════

def _batch_save_to_notion(channel, keyword, page_id, title, body, account_id=''):
    """일괄 생성 후 자동 Notion 저장 (공통)"""
    headers_n = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    props = {
        '제목': {'title': [{'text': {'content': title}}]},
        '채널': {'select': {'name': channel}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body:
        props['본문'] = {'rich_text': [{'text': {'content': body[:2000]}}]}
    if page_id:
        props['키워드'] = {'relation': [{'id': page_id}]}
    if account_id:
        props['작업계정'] = {'select': {'name': account_id}}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    if body:
        children = []
        for para in [p.strip() for p in body.split('\n\n') if p.strip()][:100]:
            for k in range(0, len(para), 2000):
                children.append({'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}})
        payload['children'] = children[:100]
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        if r.status_code != 200:
            print(f"[batch_save_to_notion] 저장 실패 channel={channel} keyword={keyword}: {r.status_code} {r.text[:200]}")
        return r.status_code == 200
    except Exception as e:
        print(f"[batch_save_to_notion] 저장 에러 channel={channel} keyword={keyword}: {e}")
        return False

@app.get("/api/batch/keywords")
async def batch_keywords():
    """일괄 생성용: 배정완료 + 미사용 키워드 로드"""
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    payload = {
        'filter': {'and': [
            {'property': '상태', 'select': {'equals': '미사용'}},
            {'property': '배정 채널', 'multi_select': {'is_not_empty': True}},
        ]},
        'page_size': 100,
    }
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % KEYWORD_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code != 200:
            return {'keywords': []}
        keywords = []
        for page in r.json().get('results', []):
            props = page.get('properties', {})
            t = props.get('키워드', {}).get('title', [])
            kw = t[0]['text']['content'] if t else ''
            channels = [c['name'] for c in props.get('배정 채널', {}).get('multi_select', [])]
            channel = channels[0] if channels else ''
            stage_sel = props.get('구매여정_단계', {}).get('select')
            stage = stage_sel['name'] if stage_sel else ''
            if kw and channel:
                keywords.append({'keyword': kw, 'channel': channel, 'page_id': page['id'], 'stage': stage})
        return {'keywords': keywords}
    except Exception:
        return {'keywords': []}

@app.post("/api/batch/generate")
async def batch_generate(request: Request):
    """일괄 생성: 키워드별 배정 채널에 맞게 순차 생성 + 자동 Notion 저장"""
    body = await request.json()
    items = body.get('keywords', [])
    product = body.get('product', {})
    default_account = body.get('account_id', '')

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(items)
        for idx, item in enumerate(items, 1):
            if await request.is_disconnected():
                print("[batch_generate] 클라이언트 연결 끊김")
                return
            kw = item.get('keyword', '')
            channel = item.get('channel', '')
            page_id = item.get('page_id', '')
            stage = item.get('stage', '')
            acc_id = item.get('account_id', default_account)

            yield _sse({'type': 'progress', 'msg': f'[{idx}/{total}] {kw} — {channel} 생성 중...', 'cur': idx-1, 'total': total})

            title = ''
            body_text = ''
            extra = {}

            try:
                if channel == '블로그':
                    # STEP 1: 제목
                    overrides = _prompt_load_overrides()
                    t_sys = overrides.get('블로그_제목', None)
                    if t_sys:
                        t_usr = f"상위 노출 키워드: {kw}"
                    else:
                        t_sys, t_usr = _build_blog_title_prompt(kw, product)
                    title_raw = await loop.run_in_executor(executor, _call_claude, t_sys, t_usr)
                    title = title_raw.strip().replace('제목:', '').replace('제목 :', '').strip().split('\n')[0]
                    # STEP 2: 본문
                    b_sys = overrides.get('블로그_본문', None)
                    if b_sys:
                        b_usr = f"[시스템 자동 전달]\n제목: {title}\n\n[사용자 입력]\n상위 노출 키워드: {kw}\n제품명: {product.get('name','')}\n제품 USP (차별 포인트): {product.get('usp','')}\n타겟층: {product.get('target','')}\n주요 성분: {product.get('ingredients','')}\n나만의 키워드: {product.get('brand_keyword','')}\n구매여정 단계: {stage}\n사진 수: 10장\n키워드 반복 수: 5회\n\n위 정보를 기반으로, 제목과 맥락이 맞는 후기형 블로그 본문을 작성해주세요."
                    else:
                        b_sys, b_usr = _build_blog_body_prompt(kw, stage, product, 10, 5, title)
                    body_text = (await loop.run_in_executor(executor, _call_claude, b_sys, b_usr)).strip()

                elif channel in ('카페', '카페SEO'):
                    # STEP 1: 제목
                    sys1, usr1 = _build_cafe_title_prompt(kw, '')
                    title_raw = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
                    title = title_raw.strip().split('\n')[0].strip()
                    # STEP 2: 본문
                    sys2, usr2 = _build_cafe_body_prompt(kw, title, '', {}, product)
                    body_text = (await loop.run_in_executor(executor, _call_claude, sys2, usr2)).strip()
                    # STEP 3: 댓글
                    sys3, usr3 = _build_cafe_comments_prompt(kw, body_text, product.get('brand_keyword', ''), product.get('alternatives', ''))
                    comments = (await loop.run_in_executor(executor, _call_claude, sys3, usr3)).strip()
                    extra = {'comments': comments}

                elif channel == '지식인':
                    # 질문 제목
                    sys1, usr1 = _build_jisikin_title_prompt(kw, product)
                    q_title = await loop.run_in_executor(executor, _call_claude, sys1, usr1)
                    q_title = q_title.strip()
                    # 질문 본문
                    sys2, usr2 = _build_jisikin_body_prompt(kw, product)
                    q_body = await loop.run_in_executor(executor, _call_claude, sys2, usr2)
                    q_body = q_body.strip()
                    # 답변
                    sys3, usr3 = _build_jisikin_answers_prompt(kw, q_title, q_body, product)
                    answers = await loop.run_in_executor(executor, _call_claude, sys3, usr3)
                    title = q_title
                    body_text = f"[질문]\n{q_title}\n\n{q_body}\n\n[답변]\n{answers.strip()}"
                    extra = {'q_title': q_title, 'q_body': q_body, 'answers': answers.strip()}

                elif channel == '카페바이럴':
                    # 3단계: 일상글 → 고민글 → 침투글+댓글
                    target_concern = product.get('target_concern', kw)
                    brand_keyword = product.get('brand_keyword', '')
                    s1_sys, s1_usr = _build_viral_stage1_prompt('', product.get('target', ''), '')
                    raw1 = await loop.run_in_executor(executor, _call_claude, s1_sys, s1_usr)
                    s1 = _parse_viral_output(raw1)
                    s2_sys, s2_usr = _build_viral_stage2_prompt('', target_concern, product.get('product_category', ''))
                    raw2 = await loop.run_in_executor(executor, _call_claude, s2_sys, s2_usr)
                    s2 = _parse_viral_output(raw2)
                    s3_sys, s3_usr = _build_viral_stage3_prompt('', target_concern, brand_keyword, product.get('name', ''), product.get('usp', ''), product.get('ingredients', ''), product.get('product_category', ''))
                    raw3 = await loop.run_in_executor(executor, _call_claude, s3_sys, s3_usr)
                    s3 = _parse_viral_stage3(raw3)
                    title = s3.get('title', '') or s2.get('title', '') or s1.get('title', '')
                    body_text = f"[1단계 일상글]\n{s1.get('title','')}\n{s1.get('body','')}\n\n[2단계 고민글]\n{s2.get('title','')}\n{s2.get('body','')}\n\n[3단계 침투글]\n{s3.get('title','')}\n{s3.get('body','')}\n\n[댓글]\n{chr(10).join(s3.get('comments',[]))}"
                    extra = {'stage1': s1, 'stage2': s2, 'stage3': s3}

                else:
                    yield _sse({'type': 'progress', 'msg': f'[{idx}/{total}] {kw} — 미지원 채널: {channel}', 'cur': idx, 'total': total})
                    continue

                # 자동 Notion 저장
                saved = await loop.run_in_executor(executor, _batch_save_to_notion, channel, kw, page_id, title, body_text, acc_id)

                # 계정 사용 기록 업데이트
                if acc_id:
                    accs = _naver_load_accounts()
                    for a in accs:
                        if a['id'] == acc_id:
                            a['total_posts'] = a.get('total_posts', 0) + 1
                            a['last_used_at'] = datetime.now().isoformat()
                            break
                    _naver_save_accounts(accs)

                result = {
                    'keyword': kw, 'channel': channel, 'title': title,
                    'body_preview': body_text[:200], 'saved': saved,
                    'account_id': acc_id, **extra,
                }
                yield _sse({'type': 'result', 'data': result, 'cur': idx, 'total': total})

            except Exception as e:
                yield _sse({'type': 'result', 'data': {'keyword': kw, 'channel': channel, 'error': str(e), 'saved': False}, 'cur': idx, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[batch_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'일괄 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

# ═══════════════════════════ WEEKLY SCHEDULER (스케줄러) ═══════════════════════════

WEEKLY_SCHEDULE_FILE = os.path.join(os.path.dirname(__file__), "weekly_schedule.json")
_scheduler_notifications = []

def _sched_load():
    if os.path.exists(WEEKLY_SCHEDULE_FILE):
        try:
            return json.loads(open(WEEKLY_SCHEDULE_FILE, encoding='utf-8').read())
        except (json.JSONDecodeError, OSError):
            pass
    return {
        'daily': {
            'generate_remind': {'enabled': True, 'time': '09:00', 'label': '콘텐츠 생성'},
            'review_remind': {'enabled': True, 'time': '14:00', 'label': '콘텐츠 검수'},
            'deploy_remind': {'enabled': True, 'time': '17:00', 'label': '콘텐츠 배포'},
        },
        'weekly': {
            'keyword_analysis': {'enabled': True, 'day': 'mon', 'time': '09:00', 'auto_run': False, 'label': '키워드 분석'},
            'channel_assign': {'enabled': True, 'day': 'mon', 'time': '10:00', 'auto_run': False, 'label': '채널 배정'},
            'performance_collect': {'enabled': True, 'day': 'fri', 'time': '17:00', 'auto_run': True, 'label': '성과 수집'},
            'weekly_report': {'enabled': True, 'day': 'fri', 'time': '18:00', 'auto_run': True, 'label': '주간 리포트'},
        },
        'history': [],
    }

def _sched_save(data):
    tmp = WEEKLY_SCHEDULE_FILE + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, WEEKLY_SCHEDULE_FILE)

@app.get("/api/scheduler/config")
async def scheduler_config_get():
    return _sched_load()

@app.post("/api/scheduler/config")
async def scheduler_config_set(request: Request):
    body = await request.json()
    data = _sched_load()
    if 'daily' in body:
        data['daily'] = body['daily']
    if 'weekly' in body:
        data['weekly'] = body['weekly']
    _sched_save(data)
    return {'ok': True}

@app.get("/api/scheduler/notifications")
async def scheduler_notifications():
    global _scheduler_notifications
    notifs, _scheduler_notifications = _scheduler_notifications, []
    return {'notifications': notifs}

@app.get("/api/scheduler/history")
async def scheduler_history():
    data = _sched_load()
    return {'history': data.get('history', [])[-50:]}

_DAY_MAP = {'mon': 0, 'tue': 1, 'wed': 2, 'thu': 3, 'fri': 4, 'sat': 5, 'sun': 6}

async def _weekly_scheduler_loop():
    """주간 스케줄러: 1분마다 체크, 시간 매칭 시 알림 또는 자동 실행"""
    global _scheduler_notifications
    _last_fired = {}
    while True:
        try:
            await asyncio.sleep(60)
            now = datetime.now()
            today_key = now.strftime('%Y-%m-%d')
            current_time = now.strftime('%H:%M')
            current_weekday = now.weekday()
            data = _sched_load()

            # 매일 알림
            for task_id, task in data.get('daily', {}).items():
                if not task.get('enabled'):
                    continue
                fire_key = f'daily_{task_id}_{today_key}'
                if fire_key in _last_fired:
                    continue
                if task.get('time') == current_time:
                    _last_fired[fire_key] = True
                    label = task.get('label', task_id)
                    # 미사용 키워드 개수 조회
                    count_msg = ''
                    if task_id == 'generate_remind':
                        try:
                            r = await batch_keywords()
                            count_msg = f" ({len(r.get('keywords', []))}개 대기)"
                        except Exception:
                            pass
                    _scheduler_notifications.append({
                        'id': fire_key, 'type': 'remind', 'task': task_id,
                        'message': f'{label} 시간입니다{count_msg}',
                        'time': now.isoformat(),
                    })
                    data.setdefault('history', []).append({'task': task_id, 'type': 'remind', 'time': now.isoformat()})
                    _sched_save(data)

            # 주간 작업
            for task_id, task in data.get('weekly', {}).items():
                if not task.get('enabled'):
                    continue
                target_day = _DAY_MAP.get(task.get('day', ''), -1)
                if current_weekday != target_day:
                    continue
                fire_key = f'weekly_{task_id}_{today_key}'
                if fire_key in _last_fired:
                    continue
                if task.get('time') == current_time:
                    _last_fired[fire_key] = True
                    label = task.get('label', task_id)
                    auto_run = task.get('auto_run', False)

                    if auto_run:
                        # TODO: 자동 실행 구현 예정. 현재는 알림만.
                        _scheduler_notifications.append({
                            'id': fire_key, 'type': 'remind', 'task': task_id,
                            'message': f'{label} 시간입니다 — 대시보드에서 실행하세요',
                            'time': now.isoformat(),
                        })
                    else:
                        _scheduler_notifications.append({
                            'id': fire_key, 'type': 'remind', 'task': task_id,
                            'message': f'{label} 시간입니다',
                            'time': now.isoformat(),
                        })
                    data.setdefault('history', []).append({
                        'task': task_id, 'type': 'auto_run' if auto_run else 'remind',
                        'time': now.isoformat(),
                    })
                    if len(data['history']) > 100:
                        data['history'] = data['history'][-100:]
                    _sched_save(data)

        except Exception as e:
            print(f"[weekly_scheduler] 루프 에러: {e}")

@app.on_event("startup")
async def _start_weekly_scheduler():
    asyncio.create_task(_weekly_scheduler_loop())

# ═══════════════════════════ THREADS (쓰레드) ═══════════════════════════

_threads_lock = threading.Lock()
_threads_scheduler_running = False

def _threads_load_accounts():
    with _threads_lock:
        if os.path.exists(THREADS_ACCOUNTS_FILE):
            with open(THREADS_ACCOUNTS_FILE, 'r') as f:
                return json.load(f)
        return {'accounts': []}

def _threads_save_accounts(data):
    with _threads_lock:
        with open(THREADS_ACCOUNTS_FILE, 'w') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

def _threads_load_queue():
    with _threads_lock:
        if os.path.exists(THREADS_QUEUE_FILE):
            with open(THREADS_QUEUE_FILE, 'r') as f:
                return json.load(f)
        return {'queue': []}

def _threads_save_queue(data):
    with _threads_lock:
        with open(THREADS_QUEUE_FILE, 'w') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

def _threads_api(access_token, endpoint, method='GET', data=None):
    """Threads Graph API 래퍼"""
    base = 'https://graph.threads.net/v1.0'
    url = f'{base}/{endpoint}'
    headers = {'Authorization': f'Bearer {access_token}'}
    try:
        if method == 'GET':
            r = req.get(url, headers=headers, params=data, timeout=15)
        else:
            r = req.post(url, headers=headers, json=data, timeout=15)
        if r.status_code == 200:
            return {'ok': True, 'data': r.json()}
        return {'ok': False, 'error': r.text[:300], 'status': r.status_code}
    except Exception as e:
        return {'ok': False, 'error': str(e)}

# ────── 계정 관리 ──────

@app.get("/api/threads/accounts")
async def threads_accounts_list():
    data = _threads_load_accounts()
    safe = []
    for acc in data.get('accounts', []):
        safe.append({
            'id': acc['id'],
            'role': acc.get('role', 'support'),
            'persona': acc.get('persona', {}),
            'reference_accounts': acc.get('reference_accounts', []),
            'schedule': acc.get('schedule', {}),
            'connected': bool(acc.get('token', {}).get('access_token')),
            'username': acc.get('username', ''),
            'token_expires': acc.get('token', {}).get('expires_at', ''),
            'daily_count': acc.get('daily_count', 0),
        })
    return {'accounts': safe}

@app.post("/api/threads/accounts")
async def threads_accounts_add(request: Request):
    body = await request.json()
    data = _threads_load_accounts()
    acc_id = 'threads_' + str(uuid.uuid4())[:8]
    account = {
        'id': acc_id,
        'role': body.get('role', 'support'),
        'token': {},
        'username': '',
        'persona': body.get('persona', {'name': '', 'age': '', 'job': '', 'tone': '친근', 'interests': []}),
        'reference_accounts': body.get('reference_accounts', []),
        'schedule': body.get('schedule', {'daily_posts': 2, 'active_hours': [9, 22], 'min_interval_hours': 4}),
        'daily_count': 0,
        'daily_count_date': datetime.now().strftime('%Y-%m-%d'),
    }
    data['accounts'].append(account)
    _threads_save_accounts(data)
    return {'ok': True, 'id': acc_id}

@app.put("/api/threads/accounts/{acc_id}")
async def threads_accounts_update(acc_id: str, request: Request):
    body = await request.json()
    data = _threads_load_accounts()
    for acc in data['accounts']:
        if acc['id'] == acc_id:
            if 'persona' in body:
                acc['persona'] = body['persona']
            if 'role' in body:
                acc['role'] = body['role']
            if 'reference_accounts' in body:
                acc['reference_accounts'] = body['reference_accounts']
            if 'schedule' in body:
                acc['schedule'] = body['schedule']
            _threads_save_accounts(data)
            return {'ok': True}
    return JSONResponse({'ok': False, 'error': '계정 없음'}, 404)

@app.delete("/api/threads/accounts/{acc_id}")
async def threads_accounts_delete(acc_id: str):
    data = _threads_load_accounts()
    data['accounts'] = [a for a in data['accounts'] if a['id'] != acc_id]
    _threads_save_accounts(data)
    return {'ok': True}

# ────── OAuth ──────

@app.get("/api/threads/auth-url")
async def threads_auth_url(account_id: str = ''):
    if not THREADS_APP_ID:
        return JSONResponse({'error': 'THREADS_APP_ID 미설정'}, 400)
    redirect_uri = f'{REDIRECT_BASE_URL}/api/threads/callback'
    scope = 'threads_basic,threads_content_publish,threads_manage_insights,threads_manage_replies'
    state = account_id or 'default'
    url = (f'https://threads.net/oauth/authorize?client_id={THREADS_APP_ID}'
           f'&redirect_uri={quote(redirect_uri)}&scope={scope}'
           f'&response_type=code&state={state}')
    return {'url': url}

@app.get("/api/threads/callback")
async def threads_callback(code: str = '', state: str = ''):
    if not code:
        return JSONResponse({'error': 'code 없음'}, 400)
    redirect_uri = f'{REDIRECT_BASE_URL}/api/threads/callback'
    # 1단계: 단기 토큰 발급
    try:
        r = req.post('https://graph.threads.net/oauth/access_token', data={
            'client_id': THREADS_APP_ID,
            'client_secret': THREADS_APP_SECRET,
            'grant_type': 'authorization_code',
            'redirect_uri': redirect_uri,
            'code': code,
        }, timeout=15)
        if r.status_code != 200:
            return JSONResponse({'error': f'토큰 발급 실패: {r.text[:300]}'}, 400)
        short_token = r.json()
    except Exception as e:
        return JSONResponse({'error': str(e)}, 500)

    # 2단계: 장기 토큰 교환 (60일)
    try:
        r2 = req.get('https://graph.threads.net/access_token', params={
            'grant_type': 'th_exchange_token',
            'client_secret': THREADS_APP_SECRET,
            'access_token': short_token.get('access_token', ''),
        }, timeout=15)
        if r2.status_code == 200:
            long_token = r2.json()
            access_token = long_token.get('access_token', short_token.get('access_token', ''))
            expires_in = long_token.get('expires_in', 5184000)
        else:
            access_token = short_token.get('access_token', '')
            expires_in = 3600
    except Exception:
        access_token = short_token.get('access_token', '')
        expires_in = 3600

    expires_at = (datetime.now() + timedelta(seconds=expires_in)).isoformat()
    user_id = short_token.get('user_id', '')

    # 사용자 이름 가져오기
    username = ''
    if access_token:
        me = _threads_api(access_token, 'me', data={'fields': 'id,username'})
        if me.get('ok'):
            username = me['data'].get('username', '')
            user_id = me['data'].get('id', user_id)

    # 계정에 토큰 저장
    data = _threads_load_accounts()
    acc_id = state if state != 'default' else None
    found = False
    for acc in data['accounts']:
        if acc['id'] == acc_id:
            acc['token'] = {'access_token': access_token, 'expires_at': expires_at, 'user_id': user_id}
            acc['username'] = username
            found = True
            break
    if not found and data['accounts']:
        # state 매칭 안되면 첫 미연결 계정에 저장
        for acc in data['accounts']:
            if not acc.get('token', {}).get('access_token'):
                acc['token'] = {'access_token': access_token, 'expires_at': expires_at, 'user_id': user_id}
                acc['username'] = username
                found = True
                break
    if not found:
        # 계정이 아예 없으면 자동 생성
        new_id = 'threads_' + str(uuid.uuid4())[:8]
        data['accounts'].append({
            'id': new_id, 'role': 'main', 'token': {'access_token': access_token, 'expires_at': expires_at, 'user_id': user_id},
            'username': username, 'persona': {'name': '', 'age': '', 'job': '', 'tone': '친근', 'interests': []},
            'reference_accounts': [], 'schedule': {'daily_posts': 3, 'active_hours': [9, 22], 'min_interval_hours': 3},
            'daily_count': 0, 'daily_count_date': datetime.now().strftime('%Y-%m-%d'),
        })
    _threads_save_accounts(data)
    # 대시보드로 리다이렉트
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url=f'{REDIRECT_BASE_URL}/?menu=threads&auth=ok')

@app.get("/api/threads/status")
async def threads_status():
    data = _threads_load_accounts()
    statuses = []
    now = datetime.now()
    for acc in data.get('accounts', []):
        token = acc.get('token', {})
        connected = bool(token.get('access_token'))
        expires_at = token.get('expires_at', '')
        days_left = -1
        warning = False
        if expires_at:
            try:
                exp = datetime.fromisoformat(expires_at)
                days_left = (exp - now).days
                warning = days_left <= 7
            except Exception:
                pass
        statuses.append({
            'id': acc['id'], 'username': acc.get('username', ''),
            'connected': connected, 'days_left': days_left, 'warning': warning,
        })
    return {'statuses': statuses}

# ────── 참조계정 크롤링 ──────

_threads_crawl_semaphore = asyncio.Semaphore(1)

def _threads_crawl_account(username):
    """Playwright로 Threads 참조계정 최근 글 크롤링"""
    posts = []
    browser = None
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx = browser.new_context(
                user_agent='Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                viewport={'width': 430, 'height': 932},
                locale='ko-KR',
            )
            page = ctx.new_page()
            clean_name = username.lstrip('@')
            page.goto(f'https://www.threads.net/@{clean_name}', timeout=20000)
            page.wait_for_timeout(3000)
            # 스크롤해서 포스트 로드
            for _ in range(3):
                page.evaluate('window.scrollBy(0, 800)')
                page.wait_for_timeout(1500)
            # 포스트 텍스트 수집
            elements = page.query_selector_all('[data-pressable-container="true"]')
            for el in elements[:10]:
                text = el.inner_text().strip()
                if text and len(text) > 10:
                    lines = text.split('\n')
                    content = '\n'.join([l for l in lines if l.strip() and not l.strip().startswith('좋아요') and not l.strip().endswith('전')])
                    if content and len(content) > 10:
                        posts.append({'text': content[:500], 'length': len(content)})
    except Exception as e:
        print(f"[threads_crawl] 에러: {e}")
    finally:
        try:
            if browser:
                browser.close()
        except Exception:
            pass
    return posts

@app.post("/api/threads/crawl-reference")
async def threads_crawl_reference(request: Request):
    body = await request.json()
    account_id = body.get('account_id', '')
    data = _threads_load_accounts()
    acc = next((a for a in data['accounts'] if a['id'] == account_id), None)
    if not acc:
        return JSONResponse({'error': '계정 없음'}, 404)
    ref_accounts = acc.get('reference_accounts', [])
    if not ref_accounts:
        return JSONResponse({'error': '참조계정 미등록'}, 400)

    async with _threads_crawl_semaphore:
        loop = asyncio.get_running_loop()
        all_posts = []
        for ref in ref_accounts[:3]:
            if await request.is_disconnected():
                print("[threads_crawl] 클라이언트 연결 끊김, 크롤링 중단")
                break
            posts = await loop.run_in_executor(executor, _threads_crawl_account, ref)
            all_posts.extend([{**p, 'source': ref} for p in posts])
    return {'posts': all_posts, 'count': len(all_posts)}

# ────── 프롬프트 ──────

def _build_threads_daily_prompt(persona, ref_posts):
    """일상글 프롬프트 — 참조계정 스타일 카피"""
    persona_desc = f"이름: {persona.get('name','')}, 나이: {persona.get('age','')}, 직업: {persona.get('job','')}"
    interests = ', '.join(persona.get('interests', []))
    tone = persona.get('tone', '친근')

    ref_examples = ''
    if ref_posts:
        ref_examples = '참조 계정 글 예시:\n'
        for i, p in enumerate(ref_posts[:5], 1):
            ref_examples += f'---예시 {i}---\n{p["text"][:200]}\n'

    system = f"""역할: Threads(쓰레드)에서 자연스러운 일상글을 작성하는 SNS 사용자.

페르소나:
{persona_desc}
관심사: {interests}
말투: {tone}

목표: 참조 계정의 글 스타일(주제 선택, 문장 길이, 톤)을 분석하고, 이 페르소나의 말투로 변환하여 Threads에 맞는 일상글을 작성한다.

{ref_examples}

Threads 플랫폼 특성 (반드시 반영):
- 트위터와 비슷한 짧은 텍스트 중심 플랫폼
- 1~3문장 단문이 대세. 길어도 5문장 이내
- 줄임말, 구어체, 혼잣말 톤이 자연스러움 ("ㄹㅇ", "진짜", "아 맞다", "근데")
- 완결된 글보다 생각을 툭 던지는 느낌
- 이모티콘은 0~1개 (과하면 봇 느낌)
- 해시태그는 거의 안 씀 (쓰더라도 1개, 안 써도 됨)
- 블로그/카페처럼 정리된 글 금지 — 머릿속 생각을 그대로 옮긴 느낌

작성 규칙:
1. 참조 계정의 주제/감성을 따라하되, 말투는 페르소나에 맞게
2. 200자 이내 권장 (최대 500자)
3. 제품·브랜드·광고 언급 절대 금지
4. 문장 끝을 다양하게 ("~인 듯", "~하다가", "~했는데", "~임", "~ㅋㅋ")
5. 줄바꿈 적극 활용 (한 문장씩 끊기)

출력 형식:
[포스트]
(본문)

[해시태그]
(해시태그 또는 '없음')

★ 출력 규칙: 인사말, 부연 설명, 괄호 설명([어그로], [본문] 등) 절대 넣지 말고 위 형식 그대로만 출력할 것.

예시:
[포스트]
아 요즘 왜 이렇게 피곤한 건지
커피를 세 잔을 마셔도 눈이 안 떠짐
나만 그런 거 아니지..?

[해시태그]
없음"""

    user = f"위 참조 계정의 스타일을 분석해서, 페르소나({persona.get('name','')})의 말투로 일상글 1개를 작성해주세요."
    return system, user

def _build_threads_traffic_prompt(keyword, persona, product, forbidden='', selling_logic='shuffle'):
    """물길글 프롬프트 — 3가지 SL(셀링 로직) 지원"""
    brand_kw = product.get('brand_keyword', '')
    tone = persona.get('tone', '친근')
    persona_job = persona.get('job', '')
    persona_interests = ', '.join(persona.get('interests', []))

    threads_base = f"""Threads 플랫폼 특성 (반드시 반영):
- 짧은 텍스트 중심. 1~5문장.
- 줄임말, 구어체 자연스러움
- 줄바꿈으로 문장 끊기
- 이모티콘 0~1개, 해시태그 0~1개
- 정리된 글 금지 — 생각을 툭 던지는 느낌

★ 첫 줄 어그로 필수:
첫 문장에서 반드시 스크롤을 멈추게 해야 함. 아래 패턴 중 하나 활용:
- 부정 편향: "~하면 진짜 손해", "아직도 ~하는 사람 있음?"
- 숫자 충격: "3년 걸림", "100만원 날렸다"
- 권위: "N년차가 말하는데", "~해본 사람으로서"
- 호기심: "이거 나만 알고 싶은데", "솔직히 말할게"
- 공감: "나만 이런 거 아니지?", "다들 이거 겪어봤을 텐데"

페르소나 말투: {tone}"""

    if selling_logic == 'shuffle':
        system = f"""역할: Threads에서 경험을 공유하는 일반 사용자.

{threads_base}

셀링 로직: 정보 셔플식
대안들 사이에 자연스럽게 제품을 섞어 간접 노출한다.

규칙:
1. [어그로 첫 줄] → 대안 3~4개 짧게 나열 → "나는 개인적으로 ~가 맞았음"
2. 나만의 키워드를 대안 중 하나로 자연스럽게 배치 (1회만)
3. 비교표/번호 매기기 금지 — 대화체로 자연스럽게
4. 200자 이내 권장 (최대 500자)
5. 광고 어투 금지 ("강추", "인생템", "꼭 써보세요")
6. 제품명 직접 언급 금지 (나만의 키워드만 사용)

출력 형식:
[포스트]
(본문)

[해시태그]
(해시태그 또는 '없음')

★ 출력 규칙: 인사말, 부연 설명, 괄호 설명 절대 넣지 말고 위 형식 그대로만 출력할 것.

예시:
[포스트]
아직도 이거 모르는 사람 많던데
A도 써봤고 B도 써봤는데
솔직히 C가 제일 나았음
근데 개인적으로는 걍 (키워드)가 맞았다
뭐 사람마다 다르겠지만

[해시태그]
없음

금지 키워드: {forbidden}"""

    elif selling_logic == 'sympathy':
        system = f"""역할: Threads에서 자신의 사업/제품을 직접 홍보하는 판매자.

{threads_base}

셀링 로직: 연민 판매자식
불쌍한 사연을 깔고 대놓고 홍보한다. 동정심을 유발하되 진심이 느껴져야 함.

구조:
1. [어그로 첫 줄] — 충격적인 사연 한 줄 ("퇴사하고 이거 하나에 올인했는데", "3년 걸렸다 진짜로")
2. 힘든 사연 — 창업 고생, 실패 경험, 밤새 작업, 돈 다 날림 등
3. 이 제품을 왜 만들었는지, 얼마나 공들였는지
4. 대놓고 홍보 — "한번만 봐주세요", "검색해봐 주세요"

규칙:
1. 제품명/나만의 키워드 직접 언급 OK (대놓고 홍보니까)
2. 구걸 느낌 아니고 "진심을 담았다" 느낌
3. 200자 이내 권장 (최대 500자)
4. 줄바꿈 많이, 짧은 문장
5. 과장된 성공 스토리 금지 — 아직 힘든 상황이어야 동정 유발

출력 형식:
[포스트]
(본문)

[해시태그]
(해시태그 또는 '없음')

★ 출력 규칙: 인사말, 부연 설명, 괄호 설명 절대 넣지 말고 위 형식 그대로만 출력할 것.

예시:
[포스트]
퇴사하고 이거 하나에 올인한 지 8개월째인데
솔직히 아직도 통장 잔고 보면 한숨 나옴
근데 진짜 좋은 거 만들었다고 자신 있어서
한번만 검색해봐 주세요 (제품명)
제발ㅠ

[해시태그]
없음

금지 키워드: {forbidden}"""

    else:  # review
        system = f"""역할: Threads에서 전문가 포지션으로 후기를 공유하는 사용자.

{threads_base}

셀링 로직: 후기 신뢰식
이 분야의 전문가/경험자로서 제품 후기를 담백하게 공유한다. 신뢰가 핵심.

페르소나 직업: {persona_job}
페르소나 관심 분야: {persona_interests}

구조:
1. [어그로 첫 줄] — 권위/경험 기반 ("이 업계 N년차가 말하는데", "솔직히 다 거기서 거긴데 이건 좀 달랐음")
2. 전문가 포지셔닝 — 이 분야에서 뭘 해왔는지 간단히
3. 비교 대상 언급 — "A도 써봤고 B도 써봤는데"
4. 이 제품의 구체적 변화/결과 — 담백하고 건조하게

규칙:
1. 제품명/나만의 키워드 직접 언급 OK
2. 페르소나의 직업/관심사가 전문성 근거가 되어야 함
3. 과장 금지 — "인생이 바뀌었다" 같은 표현 금지
4. 구체적 수치/기간 포함 ("2주 써봤는데", "확실히 달라짐")
5. 200자 이내 권장 (최대 500자)

출력 형식:
[포스트]
(본문)

[해시태그]
(해시태그 또는 '없음')

★ 출력 규칙: 인사말, 부연 설명, 괄호 설명 절대 넣지 말고 위 형식 그대로만 출력할 것.

예시:
[포스트]
이 바닥 5년차가 말하는데
솔직히 다 고만고만함
A도 써봤고 B도 써봤는데
근데 (제품명)은 2주 쓰고 확실히 달라짐
과장 아니고 체감이 다름

[해시태그]
없음

금지 키워드: {forbidden}"""

    sl_label = {'shuffle': '정보 셔플식', 'sympathy': '연민 판매자식', 'review': '후기 신뢰식'}.get(selling_logic, '정보 셔플식')
    user = f"""메인 키워드: {keyword}
소구점: {product.get('usp', '')}
타겟층: {product.get('target', '')}
나만의 키워드: {brand_kw}
제품명: {product.get('name', '')}

위 정보를 바탕으로 [{sl_label}]으로 Threads 물길글 1개를 작성해주세요."""
    return system, user

def _build_threads_comment_prompt(post_content, persona):
    """댓글 프롬프트 — 게시물 맥락에 맞는 매력적 댓글"""
    tone = persona.get('tone', '친근')
    system = f"""역할: Threads에서 소통하는 일반 사용자.

말투: {tone}

목표: 아래 게시물에 달기 좋은 자연스러운 댓글 3개를 작성한다.

Threads 댓글 특성:
- 짧다. 한 줄이 대부분 ("ㄹㅇㅋㅋ", "이거 나도 그랬음", "오 대박")
- 공감형, 질문형, 경험공유형을 섞되 전부 짧게
- 이모티콘 0~1개
- 광고·홍보 느낌 금지
- "정보 감사합니다" 같은 정중한 톤 금지 — 친구한테 말하듯

작성 규칙:
1. 각 댓글 1문장 (30자 이내)
2. 3개 모두 톤이 달라야 함
3. 줄임말, 구어체 OK

출력 형식:
1. (댓글1)
2. (댓글2)
3. (댓글3)"""
    user = f"게시물 내용:\n{post_content[:500]}\n\n위 게시물에 달 매력적인 댓글 3개를 작성해주세요."
    return system, user

def _parse_threads_output(raw):
    """쓰레드 생성 결과 파싱: 포스트 본문 + 해시태그 분리 (LLM 인사말 방어 포함)"""
    text = raw.strip()
    hashtag = ''
    # LLM 인사말/부연 제거: [포스트] 태그 이전의 텍스트는 전부 버림
    if '[포스트]' in text:
        text = text[text.index('[포스트]'):]
    # 해시태그 분리
    if '[해시태그]' in text:
        parts = text.split('[해시태그]')
        text = parts[0].strip()
        ht = parts[1].strip()
        if ht and ht != '없음':
            hashtag = ht.strip()
    # [포스트] 태그 제거
    if text.startswith('[포스트]'):
        text = text[len('[포스트]'):].strip()
    # 괄호 설명 태그 잔여물 제거 ("[어그로 첫 줄]" 등)
    text = re.sub(r'\[어그로[^\]]*\]\s*', '', text)
    text = re.sub(r'\[본문[^\]]*\]\s*', '', text)
    text = re.sub(r'\[마무리[^\]]*\]\s*', '', text)
    # 해시태그가 본문 안에 있으면 분리
    if not hashtag:
        ht_match = re.findall(r'#\S+', text)
        if ht_match:
            hashtag = ht_match[0]
    return text.strip(), hashtag

# ────── 콘텐츠 생성 ──────

@app.get("/api/threads/notion-keywords")
async def threads_notion_keywords():
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    payload = {
        'filter': {'and': [
            {'property': '배정 채널', 'multi_select': {'contains': '쓰레드'}},
            {'property': '상태', 'select': {'equals': '미사용'}},
        ]},
        'page_size': 100,
    }
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % KEYWORD_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code != 200:
            return {'keywords': []}
        keywords = []
        for page in r.json().get('results', []):
            props = page.get('properties', {})
            t = props.get('키워드', {}).get('title', [])
            kw = t[0]['text']['content'] if t else ''
            if kw:
                keywords.append({'keyword': kw, 'page_id': page['id']})
        return {'keywords': keywords}
    except Exception:
        return {'keywords': []}

@app.post("/api/threads/generate")
async def threads_generate(request: Request):
    body = await request.json()
    post_type = body.get('type', 'daily')  # 'daily' or 'traffic'
    account_id = body.get('account_id', '')
    keywords = body.get('keywords', [])
    product = body.get('product', {})
    forbidden = body.get('forbidden', '')
    count = body.get('count', 3)
    ref_posts = body.get('ref_posts', [])
    selling_logic = body.get('selling_logic', 'shuffle')  # shuffle / sympathy / review

    data = _threads_load_accounts()
    acc = next((a for a in data['accounts'] if a['id'] == account_id), None)
    persona = acc.get('persona', {}) if acc else {'tone': '친근'}

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        if post_type == 'daily':
            total = count
            for idx in range(1, count + 1):
                if await request.is_disconnected():
                    print("[threads_generate] 클라이언트 연결 끊김, 생성 중단")
                    return
                yield _sse({'type': 'progress', 'msg': f'[{idx}/{total}] 일상글 생성 중...', 'cur': idx-1, 'total': total})
                sys_p, usr_p = _build_threads_daily_prompt(persona, ref_posts)
                raw = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p, 0.85)
                text, hashtag = _parse_threads_output(raw)
                full_text = f'{text}\n\n{hashtag}'.strip() if hashtag else text
                result = {'text': text, 'hashtag': hashtag, 'full_text': full_text, 'char_count': len(full_text), 'type': 'daily', 'num': idx}
                yield _sse({'type': 'result', 'data': result, 'cur': idx, 'total': total})
            yield _sse({'type': 'complete', 'total': total})
        else:
            total = len(keywords) * count
            idx = 0
            for kw_data in keywords:
                kw = kw_data if isinstance(kw_data, str) else kw_data.get('keyword', '')
                page_id = '' if isinstance(kw_data, str) else kw_data.get('page_id', '')
                for c in range(count):
                    if await request.is_disconnected():
                        print("[threads_generate] 클라이언트 연결 끊김, 생성 중단")
                        return
                    idx += 1
                    label = kw if count == 1 else f'{kw} (#{c+1})'
                    yield _sse({'type': 'progress', 'msg': f'[{idx}/{total}] {label} — 물길글 생성 중...', 'cur': idx-1, 'total': total})
                    sys_p, usr_p = _build_threads_traffic_prompt(kw, persona, product, forbidden, selling_logic)
                    raw = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p, 0.8)
                    text, hashtag = _parse_threads_output(raw)
                    full_text = f'{text}\n\n{hashtag}'.strip() if hashtag else text
                    result = {
                        'keyword': kw, 'text': text, 'hashtag': hashtag,
                        'full_text': full_text, 'char_count': len(full_text),
                        'type': 'traffic', 'page_id': page_id, 'num': c + 1,
                    }
                    yield _sse({'type': 'result', 'data': result, 'cur': idx, 'total': total})
            yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[threads_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'쓰레드 생성 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.post("/api/threads/generate-comment")
async def threads_generate_comment(request: Request):
    """댓글 텍스트 생성"""
    body = await request.json()
    post_content = body.get('post_content', '')
    account_id = body.get('account_id', '')
    if not post_content:
        return JSONResponse({'error': '게시물 내용 필요'}, 400)
    data = _threads_load_accounts()
    acc = next((a for a in data['accounts'] if a['id'] == account_id), None)
    persona = acc.get('persona', {}) if acc else {'tone': '친근'}
    loop = asyncio.get_running_loop()
    sys_p, usr_p = _build_threads_comment_prompt(post_content, persona)
    raw = await loop.run_in_executor(executor, _call_claude, sys_p, usr_p, 0.8)
    comments = []
    for line in raw.strip().split('\n'):
        line = line.strip()
        if line and line[0].isdigit() and '.' in line:
            comments.append(line.split('.', 1)[1].strip())
        elif line and not line.startswith('['):
            comments.append(line)
    return {'comments': comments[:5]}

# ────── 게시 ──────

def _threads_publish_post(access_token, user_id, text, reply_to=None):
    """Threads 공식 API 2단계 게시"""
    # 500자 제한
    if len(text) > 500:
        text = text[:497] + '...'
    # 1단계: 컨테이너 생성
    container_data = {'media_type': 'TEXT', 'text': text}
    if reply_to:
        container_data['reply_to_id'] = reply_to
    result = _threads_api(access_token, f'{user_id}/threads', 'POST', container_data)
    if not result.get('ok'):
        return result
    creation_id = result['data'].get('id', '')
    if not creation_id:
        return {'ok': False, 'error': 'container id 없음'}
    # 2초 대기 (컨테이너 처리)
    time.sleep(2)
    # 2단계: 게시
    pub_result = _threads_api(access_token, f'{user_id}/threads_publish', 'POST', {'creation_id': creation_id})
    return pub_result

@app.post("/api/threads/publish")
async def threads_publish(request: Request):
    body = await request.json()
    account_id = body.get('account_id', '')
    text = body.get('text', '')
    reply_to = body.get('reply_to', None)
    if not text:
        return JSONResponse({'error': '텍스트 필요'}, 400)
    data = _threads_load_accounts()
    acc = next((a for a in data['accounts'] if a['id'] == account_id), None)
    if not acc or not acc.get('token', {}).get('access_token'):
        return JSONResponse({'error': '계정 미연결'}, 400)
    token = acc['token']['access_token']
    user_id = acc['token'].get('user_id', 'me')
    # 일일 카운터 체크
    today = datetime.now().strftime('%Y-%m-%d')
    if acc.get('daily_count_date') != today:
        acc['daily_count'] = 0
        acc['daily_count_date'] = today
    if acc['daily_count'] >= 250:
        return JSONResponse({'error': '일일 게시 한도 초과 (250건)'}, 429)

    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(executor, _threads_publish_post, token, user_id, text, reply_to)
    if result.get('ok'):
        # Race Condition 방지: 최신 파일 읽고 해당 계정만 업데이트
        fresh = _threads_load_accounts()
        for fa in fresh['accounts']:
            if fa['id'] == account_id:
                today_str = datetime.now().strftime('%Y-%m-%d')
                if fa.get('daily_count_date') != today_str:
                    fa['daily_count'] = 0
                    fa['daily_count_date'] = today_str
                fa['daily_count'] = fa.get('daily_count', 0) + 1
                fa['last_publish_time'] = datetime.now().isoformat()
                break
        _threads_save_accounts(fresh)
    return result

# ────── Notion 저장 ──────

@app.post("/api/threads/save-notion")
async def threads_save_notion(request: Request):
    body = await request.json()
    headers_n = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    kw = body.get('keyword', '쓰레드')
    post_type = body.get('type', 'daily')
    type_label = '일상글' if post_type == 'daily' else '물길글'
    props = {
        '제목': {'title': [{'text': {'content': f'{kw} 쓰레드 {type_label}'}}]},
        '채널': {'select': {'name': '쓰레드'}},
        '생산 상태': {'select': {'name': '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    text = body.get('text', '')
    if text:
        props['본문'] = {'rich_text': [{'text': {'content': text[:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    if text:
        children = []
        for para in [p.strip() for p in text.split('\n\n') if p.strip()][:100]:
            for k in range(0, len(para), 2000):
                children.append({'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}})
        payload['children'] = children[:100]
    try:
        r = req.post('https://api.notion.com/v1/pages', headers=headers_n, json=payload, timeout=15)
        return {'success': r.status_code == 200, 'error': '' if r.status_code == 200 else r.text[:300]}
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ────── 스케줄러 ──────

@app.post("/api/threads/schedule")
async def threads_schedule_post(request: Request):
    """게시물을 스케줄 큐에 등록"""
    body = await request.json()
    queue = _threads_load_queue()
    item = {
        'id': str(uuid.uuid4())[:8],
        'account_id': body.get('account_id', ''),
        'text': body.get('text', ''),
        'type': body.get('type', 'daily'),
        'keyword': body.get('keyword', ''),
        'page_id': body.get('page_id', ''),
        'created_at': datetime.now().isoformat(),
        'scheduled_at': None,  # 스케줄러가 자동 배정
        'status': 'pending',  # pending / published / failed
    }
    queue['queue'].append(item)
    _threads_save_queue(queue)
    return {'ok': True, 'queue_id': item['id']}

@app.get("/api/threads/queue")
async def threads_queue_list():
    queue = _threads_load_queue()
    return {'queue': queue.get('queue', [])}

@app.delete("/api/threads/queue/{queue_id}")
async def threads_queue_delete(queue_id: str):
    queue = _threads_load_queue()
    queue['queue'] = [q for q in queue['queue'] if q['id'] != queue_id]
    _threads_save_queue(queue)
    return {'ok': True}

def _threads_refresh_token(acc):
    """만료 15일 전 장기 토큰 자동 갱신"""
    token = acc.get('token', {})
    access_token = token.get('access_token', '')
    expires_at = token.get('expires_at', '')
    if not access_token or not expires_at:
        return False
    try:
        exp = datetime.fromisoformat(expires_at)
        days_left = (exp - datetime.now()).days
        if days_left > 15:
            return False  # 아직 갱신 불필요
        # 장기 토큰 갱신 API 호출
        r = req.get('https://graph.threads.net/refresh_access_token', params={
            'grant_type': 'th_refresh_token',
            'access_token': access_token,
        }, timeout=15)
        if r.status_code == 200:
            new_data = r.json()
            acc['token']['access_token'] = new_data.get('access_token', access_token)
            new_expires = new_data.get('expires_in', 5184000)
            acc['token']['expires_at'] = (datetime.now() + timedelta(seconds=new_expires)).isoformat()
            print(f"[threads] 토큰 갱신 완료: {acc.get('username','')} (새 만료: {acc['token']['expires_at']})")
            return True
        else:
            print(f"[threads] 토큰 갱신 실패: {r.status_code} {r.text[:200]}")
            return False
    except Exception as e:
        print(f"[threads] 토큰 갱신 에러: {e}")
        return False

async def _threads_scheduler_loop():
    """백그라운드 스케줄러: 1분마다 큐 확인 → 조건 맞으면 자동 게시 + 토큰 갱신"""
    global _threads_scheduler_running
    _threads_scheduler_running = True
    _token_check_counter = 0
    while _threads_scheduler_running:
        try:
            await asyncio.sleep(60)
            now = datetime.now()
            hour = now.hour
            queue = _threads_load_queue()
            accounts_data = _threads_load_accounts()

            # 토큰 갱신 체크 (60분마다 1회)
            _token_check_counter += 1
            if _token_check_counter >= 60:
                _token_check_counter = 0
                loop = asyncio.get_running_loop()
                token_changed = False
                for acc in accounts_data.get('accounts', []):
                    if acc.get('token', {}).get('access_token'):
                        refreshed = await loop.run_in_executor(executor, _threads_refresh_token, acc)
                        if refreshed:
                            token_changed = True
                if token_changed:
                    _threads_save_accounts(accounts_data)
            changed = False

            for item in queue.get('queue', []):
                if item['status'] != 'pending':
                    continue
                acc_id = item['account_id']
                acc = next((a for a in accounts_data['accounts'] if a['id'] == acc_id), None)
                if not acc or not acc.get('token', {}).get('access_token'):
                    continue
                schedule = acc.get('schedule', {})
                active_start, active_end = schedule.get('active_hours', [9, 22])
                daily_max = schedule.get('daily_posts', 3)
                min_interval = schedule.get('min_interval_hours', 3)

                # 활동 시간대 체크
                if hour < active_start or hour >= active_end:
                    continue
                # 일일 한도 체크
                today = now.strftime('%Y-%m-%d')
                if acc.get('daily_count_date') != today:
                    acc['daily_count'] = 0
                    acc['daily_count_date'] = today
                if acc.get('daily_count', 0) >= daily_max:
                    continue
                # 최소 간격 체크 (마지막 게시 시간)
                last_pub = acc.get('last_publish_time', '')
                if last_pub:
                    try:
                        last_dt = datetime.fromisoformat(last_pub)
                        elapsed = (now - last_dt).total_seconds() / 3600
                        # 랜덤 지터 추가 (봇 탐지 방지)
                        jitter = random.uniform(0, 0.5)
                        if elapsed < min_interval + jitter:
                            continue
                    except Exception:
                        pass

                # 게시 실행
                token = acc['token']['access_token']
                user_id = acc['token'].get('user_id', 'me')
                loop = asyncio.get_running_loop()
                result = await loop.run_in_executor(executor, _threads_publish_post, token, user_id, item['text'])
                if result.get('ok'):
                    item['status'] = 'published'
                    item['published_at'] = now.isoformat()
                    item['post_id'] = result.get('data', {}).get('id', '')
                    acc['daily_count'] = acc.get('daily_count', 0) + 1
                    acc['last_publish_time'] = now.isoformat()
                    changed = True
                    print(f"[threads_scheduler] 게시 완료: {acc.get('username','')} - {item['text'][:30]}...")
                else:
                    item['status'] = 'failed'
                    item['error'] = result.get('error', '')[:200]
                    changed = True

            if changed:
                _threads_save_queue(queue)
                # Race Condition 방지: 최신 파일 읽고 변경된 필드만 병합
                fresh_data = _threads_load_accounts()
                for acc in accounts_data.get('accounts', []):
                    for fresh_acc in fresh_data.get('accounts', []):
                        if fresh_acc['id'] == acc['id']:
                            # 스케줄러가 변경하는 필드만 덮어쓰기
                            fresh_acc['daily_count'] = acc.get('daily_count', 0)
                            fresh_acc['daily_count_date'] = acc.get('daily_count_date', '')
                            fresh_acc['last_publish_time'] = acc.get('last_publish_time', '')
                            if acc.get('token', {}).get('access_token'):
                                fresh_acc['token'] = acc['token']
                            break
                _threads_save_accounts(fresh_data)
        except asyncio.CancelledError:
            raise
        except Exception as e:
            print(f"[threads_scheduler] 루프 에러: {e}")

@app.on_event("startup")
async def _start_threads_scheduler():  # noqa: deprecation — lifespan 전환은 별도 리팩터링
    asyncio.create_task(_threads_scheduler_loop())

# ────── 인사이트 ──────

@app.get("/api/threads/insights")
async def threads_insights(post_id: str = '', account_id: str = ''):
    if not post_id or not account_id:
        return JSONResponse({'error': 'post_id, account_id 필요'}, 400)
    data = _threads_load_accounts()
    acc = next((a for a in data['accounts'] if a['id'] == account_id), None)
    if not acc or not acc.get('token', {}).get('access_token'):
        return JSONResponse({'error': '계정 미연결'}, 400)
    token = acc['token']['access_token']
    result = _threads_api(token, f'{post_id}/insights', data={
        'metric': 'views,likes,replies,reposts,quotes,shares'
    })
    return result

# ───────────────────────────── PERFORMANCE ─────────────────────────────

PERF_FILE = os.path.join(os.path.dirname(__file__), "performance_history.json")
PERF_SCHEDULE_FILE = os.path.join(os.path.dirname(__file__), "perf_schedule.json")
_perf_schedule = {"enabled": False, "interval_hours": 24}
_perf_task = None  # asyncio.Task

def _perf_load():
    if os.path.exists(PERF_FILE):
        try:
            with open(PERF_FILE, 'r') as f:
                return json.load(f)
        except Exception:
            pass
    return {"records": [], "last_checked": "", "total_checks": 0}

def _perf_save(data):
    data["records"] = data["records"][-10000:]  # 최근 10000건 유지
    with open(PERF_FILE, 'w') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def _check_exposure_enhanced(keyword, deploy_url, channel=''):
    """네이버 검색에서 URL 노출 여부 + 섹션명 + 정확한 순위"""
    if not deploy_url:
        return {'exposure': '-', 'rank': 0, 'section': ''}
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get('https://search.naver.com/search.naver?query=%s&where=nexearch' % quote(keyword), headers=headers, timeout=10)
        html = r.text
        if deploy_url not in html:
            return {'exposure': '미노출', 'rank': 0, 'section': ''}
        soup = BeautifulSoup(html, 'html.parser')
        # 섹션별 순위 파악
        section_name = ''
        rank_in_section = 0
        # SERP API 방식: data-cr-area 속성으로 섹션 식별
        containers = soup.find_all(attrs={'data-cr-area': True})
        for container in containers:
            area_code = container.get('data-cr-area', '')
            if deploy_url in str(container):
                section_name = SECTION_MAP.get(area_code, area_code)
                # 섹션 내 링크 순위
                links = container.find_all('a', href=True)
                for idx, a in enumerate(links):
                    if deploy_url in a.get('href', ''):
                        rank_in_section = idx + 1
                        break
                break
        # fallback: 전체 페이지에서 순위
        if not rank_in_section:
            for i, a in enumerate(soup.find_all('a', href=True)):
                if deploy_url in a.get('href', ''):
                    rank_in_section = i + 1
                    break
        return {'exposure': '노출중', 'rank': rank_in_section, 'section': section_name}
    except Exception as e:
        print(f"[perf] exposure check error: {e}")
        return {'exposure': '-', 'rank': 0, 'section': ''}

def _fetch_blog_stats(deploy_url):
    """네이버 블로그 모바일 페이지에서 조회수/댓글수/공감수 크롤링"""
    result = {'views': 0, 'comments': 0, 'likes': 0}
    if not deploy_url:
        return result
    try:
        m = re.search(r'blog\.naver\.com/([^/?]+)/(\d+)', deploy_url)
        if not m:
            return result
        blog_id, log_no = m.group(1), m.group(2)
        mobile_url = f"https://m.blog.naver.com/{blog_id}/{log_no}"
        headers = {'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.0 Mobile/15E148 Safari/604.1'}
        r = req.get(mobile_url, headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        # 조회수
        view_el = soup.select_one('.u_cnt._count, .view_count .u_cnt, .blog_count, [class*="view"] [class*="count"], .info_count .num')
        if view_el:
            view_text = re.sub(r'[^\d]', '', view_el.get_text())
            if view_text:
                result['views'] = int(view_text)
        # 댓글수
        comment_el = soup.select_one('.comment_count, .btn_comment .u_cnt, [class*="comment"] [class*="count"], .comment_area .num')
        if comment_el:
            comment_text = re.sub(r'[^\d]', '', comment_el.get_text())
            if comment_text:
                result['comments'] = int(comment_text)
        # 공감수
        like_el = soup.select_one('.sympathy_count, .btn_like .u_cnt, [class*="like"] [class*="count"], .sympathy_area .num')
        if like_el:
            like_text = re.sub(r'[^\d]', '', like_el.get_text())
            if like_text:
                result['likes'] = int(like_text)
    except Exception as e:
        print(f"[perf] blog stats error: {e}")
    return result

def _run_performance_collect_sync(items):
    """성과 수집 실행 (동기, 스레드에서 호출)"""
    results = []
    for item in items:
        kw = item.get('keyword', '')
        url = item.get('deploy_url', '')
        channel = item.get('channel', '')
        title = item.get('title', '')
        deploy_date = item.get('deploy_date', '')
        # 노출 체크
        exp = _check_exposure_enhanced(kw, url, channel)
        # 블로그면 반응 데이터도 수집
        stats = {'views': 0, 'comments': 0, 'likes': 0}
        if channel == '블로그' and 'blog.naver.com' in (url or ''):
            stats = _fetch_blog_stats(url)
        record = {
            'checked_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%S'),
            'keyword': kw,
            'channel': channel,
            'title': title,
            'deploy_url': url,
            'deploy_date': deploy_date,
            'work_account': item.get('work_account', ''),
            'exposure': exp['exposure'],
            'rank': exp['rank'],
            'section': exp['section'],
            'views': stats['views'],
            'comments': stats['comments'],
            'likes': stats['likes'],
        }
        results.append(record)
        time.sleep(1.5)  # 차단 방지
    return results

@app.post("/api/performance/collect")
async def performance_collect(request: Request):
    """성과 수집 실행 (SSE 스트리밍)"""
    body = await request.json()
    mode = body.get('mode', 'all')  # all | selected
    selected_items = body.get('items', [])

    def _sse(obj):
        return "data: " + json.dumps(obj, ensure_ascii=False) + "\n\n"

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        items = []
        if mode == 'all':
            # 배포완료 콘텐츠 전체 조회
            yield _sse({'type': 'progress', 'msg': '노션에서 배포 콘텐츠 조회 중...', 'cur': 0, 'total': 0})
            filter_obj = {'property': '발행_상태', 'select': {'equals': '발행완료'}}
            pages = await loop.run_in_executor(executor, _notion_query_all, CONTENT_DB_ID, filter_obj)
            for page in pages:
                props = page.get('properties', {})
                kw_rels = _extract_prop(props, '키워드', 'relation')
                kw_name = ''
                if kw_rels:
                    # 키워드 이름 가져오기
                    try:
                        kw_r = req.get('https://api.notion.com/v1/pages/%s' % kw_rels[0],
                                       headers={'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Notion-Version': '2022-06-28'}, timeout=10)
                        if kw_r.status_code == 200:
                            kw_props = kw_r.json().get('properties', {})
                            kw_name = _extract_prop(kw_props, '키워드', 'title')
                    except Exception:
                        pass
                work_acc = _extract_prop(props, '작업계정', 'select')
                if not work_acc:
                    work_acc_rt = props.get('작업계정', {}).get('rich_text', [])
                    work_acc = work_acc_rt[0]['text']['content'] if work_acc_rt else ''
                items.append({
                    'keyword': kw_name,
                    'title': _extract_prop(props, '제목', 'title'),
                    'channel': _extract_prop(props, '채널', 'select'),
                    'deploy_url': _extract_prop(props, '발행_URL', 'url'),
                    'deploy_date': _extract_prop(props, '생성일', 'date'),
                    'work_account': work_acc,
                })
            items = [it for it in items if it['deploy_url']]
        else:
            items = selected_items

        total = len(items)
        if total == 0:
            yield _sse({'type': 'complete', 'total': 0, 'results': [], 'message': '배포된 콘텐츠가 없습니다.'})
            return

        yield _sse({'type': 'progress', 'msg': f'총 {total}개 콘텐츠 성과 수집 시작', 'cur': 0, 'total': total})

        results = []
        for i, item in enumerate(items):
            kw = item.get('keyword', '')
            url = item.get('deploy_url', '')
            channel = item.get('channel', '')
            yield _sse({'type': 'progress', 'msg': f'[{i+1}/{total}] {kw or item.get("title","")} 수집 중...', 'cur': i+1, 'total': total})

            exp = await loop.run_in_executor(executor, _check_exposure_enhanced, kw, url, channel)
            stats = {'views': 0, 'comments': 0, 'likes': 0}
            if channel == '블로그' and 'blog.naver.com' in (url or ''):
                stats = await loop.run_in_executor(executor, _fetch_blog_stats, url)

            record = {
                'checked_at': datetime.now().strftime('%Y-%m-%dT%H:%M:%S'),
                'keyword': kw,
                'channel': channel,
                'title': item.get('title', ''),
                'deploy_url': url,
                'deploy_date': item.get('deploy_date', ''),
                'exposure': exp['exposure'],
                'rank': exp['rank'],
                'section': exp['section'],
                'views': stats['views'],
                'comments': stats['comments'],
                'likes': stats['likes'],
            }
            results.append(record)
            yield _sse({'type': 'result', 'data': record, 'cur': i+1, 'total': total})
            await asyncio.sleep(1.5)

        # 히스토리 저장
        perf_data = _perf_load()
        perf_data['records'].extend(results)
        perf_data['last_checked'] = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
        perf_data['total_checks'] = perf_data.get('total_checks', 0) + 1
        _perf_save(perf_data)

        yield _sse({'type': 'complete', 'total': total, 'results': results, 'message': f'{total}개 콘텐츠 성과 수집 완료'})
      except Exception as e:
        print(f"[performance] collect error: {e}")
        yield _sse({'type': 'error', 'message': f'성과 수집 중 오류: {e}'})

    return StreamingResponse(generate(), media_type="text/event-stream")

@app.get("/api/performance/history")
async def performance_history(keyword: str = '', channel: str = '', days: int = 30):
    """성과 히스토리 조회"""
    perf_data = _perf_load()
    records = perf_data.get('records', [])
    cutoff = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%dT%H:%M:%S')
    filtered = [r for r in records if r.get('checked_at', '') >= cutoff]
    if keyword:
        filtered = [r for r in filtered if keyword in r.get('keyword', '')]
    if channel:
        filtered = [r for r in filtered if r.get('channel', '') == channel]
    return {'records': filtered, 'total': len(filtered), 'last_checked': perf_data.get('last_checked', '')}

@app.get("/api/performance/dashboard-data")
async def performance_dashboard_data(days: int = 30):
    """대시보드 요약 집계"""
    perf_data = _perf_load()
    records = perf_data.get('records', [])
    cutoff = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%dT%H:%M:%S')
    recent = [r for r in records if r.get('checked_at', '') >= cutoff]

    # 최신 체크 기준으로 키워드별 최신 레코드 추출
    latest_by_key = {}
    for r in recent:
        key = (r.get('keyword', ''), r.get('deploy_url', ''))
        if key not in latest_by_key or r.get('checked_at', '') > latest_by_key[key].get('checked_at', ''):
            latest_by_key[key] = r
    latest = list(latest_by_key.values())

    # 요약 카드
    total_content = len(latest)
    exposed = sum(1 for r in latest if r.get('exposure') == '노출중')
    ranks = [r['rank'] for r in latest if r.get('exposure') == '노출중' and r.get('rank', 0) > 0]
    avg_rank = round(sum(ranks) / len(ranks), 1) if ranks else 0
    total_views = sum(r.get('views', 0) for r in latest)

    # 채널별 노출률
    channel_stats = {}
    for r in latest:
        ch = r.get('channel', '기타')
        if ch not in channel_stats:
            channel_stats[ch] = {'total': 0, 'exposed': 0}
        channel_stats[ch]['total'] += 1
        if r.get('exposure') == '노출중':
            channel_stats[ch]['exposed'] += 1
    channel_rates = []
    for ch, st in sorted(channel_stats.items()):
        rate = round(st['exposed'] / st['total'] * 100) if st['total'] > 0 else 0
        channel_rates.append({'channel': ch, 'total': st['total'], 'exposed': st['exposed'], 'rate': rate})

    # 일별 노출 추이 (최근 N일)
    daily_trend = {}
    for r in recent:
        day = r.get('checked_at', '')[:10]
        if day not in daily_trend:
            daily_trend[day] = {'total': 0, 'exposed': 0}
        daily_trend[day]['total'] += 1
        if r.get('exposure') == '노출중':
            daily_trend[day]['exposed'] += 1
    trend = [{'date': d, 'total': v['total'], 'exposed': v['exposed']} for d, v in sorted(daily_trend.items())]

    # 순위 변동 (최신 vs 이전)
    rank_changes = []
    for r in latest:
        key = (r.get('keyword', ''), r.get('deploy_url', ''))
        # 이전 기록 찾기
        prev_records = [p for p in recent if (p.get('keyword', ''), p.get('deploy_url', '')) == key and p.get('checked_at', '') < r.get('checked_at', '')]
        prev_rank = 0
        if prev_records:
            prev_records.sort(key=lambda x: x.get('checked_at', ''), reverse=True)
            prev_rank = prev_records[0].get('rank', 0)
        change = prev_rank - r.get('rank', 0) if prev_rank > 0 and r.get('rank', 0) > 0 else 0
        rank_changes.append({
            'keyword': r.get('keyword', ''),
            'channel': r.get('channel', ''),
            'title': r.get('title', ''),
            'deploy_url': r.get('deploy_url', ''),
            'deploy_date': r.get('deploy_date', ''),
            'current_rank': r.get('rank', 0),
            'prev_rank': prev_rank,
            'change': change,
            'section': r.get('section', ''),
            'exposure': r.get('exposure', ''),
            'views': r.get('views', 0),
            'comments': r.get('comments', 0),
            'likes': r.get('likes', 0),
        })

    # 블로그 반응 TOP (조회수 순)
    blog_stats = [r for r in rank_changes if r.get('channel') == '블로그' and r.get('views', 0) > 0]
    blog_stats.sort(key=lambda x: x.get('views', 0), reverse=True)

    return {
        'summary': {
            'total_content': total_content,
            'exposed': exposed,
            'avg_rank': avg_rank,
            'total_views': total_views,
        },
        'channel_rates': channel_rates,
        'trend': trend,
        'rank_changes': rank_changes,
        'blog_stats': blog_stats[:20],
        'last_checked': perf_data.get('last_checked', ''),
        'total_checks': perf_data.get('total_checks', 0),
    }

# 자동 수집 스케줄러
async def _perf_auto_collect():
    """자동 성과 수집 (백그라운드)"""
    try:
        filter_obj = {'property': '발행_상태', 'select': {'equals': '발행완료'}}
        loop = asyncio.get_running_loop()
        pages = await loop.run_in_executor(executor, _notion_query_all, CONTENT_DB_ID, filter_obj)
        items = []
        for page in pages:
            props = page.get('properties', {})
            kw_rels = _extract_prop(props, '키워드', 'relation')
            kw_name = ''
            if kw_rels:
                try:
                    kw_r = req.get('https://api.notion.com/v1/pages/%s' % kw_rels[0],
                                   headers={'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Notion-Version': '2022-06-28'}, timeout=10)
                    if kw_r.status_code == 200:
                        kw_props = kw_r.json().get('properties', {})
                        kw_name = _extract_prop(kw_props, '키워드', 'title')
                except Exception:
                    pass
            url = _extract_prop(props, '발행_URL', 'url')
            if url:
                items.append({
                    'keyword': kw_name,
                    'title': _extract_prop(props, '제목', 'title'),
                    'channel': _extract_prop(props, '채널', 'select'),
                    'deploy_url': url,
                    'deploy_date': _extract_prop(props, '생성일', 'date'),
                })
        if items:
            loop = asyncio.get_running_loop()
            results = await loop.run_in_executor(executor, _run_performance_collect_sync, items)
            perf_data = _perf_load()
            perf_data['records'].extend(results)
            perf_data['last_checked'] = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
            perf_data['total_checks'] = perf_data.get('total_checks', 0) + 1
            _perf_save(perf_data)
            print(f"[perf-auto] {len(results)}개 콘텐츠 성과 수집 완료")
    except Exception as e:
        print(f"[perf-auto] error: {e}")

async def _perf_scheduler_loop():
    """자동 수집 스케줄러 루프"""
    global _perf_schedule
    while _perf_schedule.get("enabled", False):
        await _perf_auto_collect()
        hours = _perf_schedule.get("interval_hours", 24)
        await asyncio.sleep(hours * 3600)

@app.post("/api/performance/schedule")
async def performance_schedule_set(request: Request):
    """자동 수집 스케줄 설정"""
    global _perf_schedule, _perf_task
    body = await request.json()
    enabled = body.get('enabled', False)
    interval = body.get('interval_hours', 24)
    _perf_schedule = {"enabled": enabled, "interval_hours": interval}
    # 설정 저장
    with open(PERF_SCHEDULE_FILE, 'w') as f:
        json.dump(_perf_schedule, f)
    # 태스크 관리
    if _perf_task and not _perf_task.done():
        _perf_task.cancel()
        _perf_task = None
    if enabled:
        _perf_task = asyncio.create_task(_perf_scheduler_loop())
    return {'success': True, 'schedule': _perf_schedule}

@app.get("/api/performance/schedule")
async def performance_schedule_get():
    """현재 스케줄 상태 조회"""
    running = _perf_task is not None and not _perf_task.done() if _perf_task else False
    return {'schedule': _perf_schedule, 'running': running}

# 서버 시작 시 스케줄 복원
@app.on_event("startup")
async def _perf_restore_schedule():
    global _perf_schedule, _perf_task
    if os.path.exists(PERF_SCHEDULE_FILE):
        try:
            with open(PERF_SCHEDULE_FILE, 'r') as f:
                _perf_schedule = json.load(f)
            if _perf_schedule.get("enabled", False):
                _perf_task = asyncio.create_task(_perf_scheduler_loop())
                print("[perf] auto-collect schedule restored")
        except Exception:
            pass

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
