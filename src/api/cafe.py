"""카페SEO 원고 API 라우터"""
import asyncio
import json
import os
import re
import shutil
import subprocess
import time
import threading
from datetime import datetime
from urllib.parse import quote

import cv2
import requests as req
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Request
from fastapi.responses import FileResponse, JSONResponse, Response
from src.services.sse_helper import sse_dict, SSEResponse

from src.services.config import (
    executor, KEYWORD_DB_ID, CONTENT_DB_ID, NOTION_TOKEN, GEMINI_API_KEY, BASE_DIR,
)
from src.services.common import error_response
from src.services.ai_client import call_claude
from src.services.review_service import review_and_save

router = APIRouter()

# ── 경로 ──
TEMP_PHOTO_DIR = os.path.join(BASE_DIR, "temp_photos")
OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(TEMP_PHOTO_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)

NAVER_ACCOUNTS_FILE = os.path.join(BASE_DIR, "naver_accounts.json")
_naver_accounts_lock = threading.Lock()

XHS_PATH = os.environ.get('XHS_PATH', shutil.which('xhs') or '/Users/iconlms/Library/Python/3.11/bin/xhs')


# ── 네이버 계정 관리 ──

def _naver_load_accounts():
    if os.path.exists(NAVER_ACCOUNTS_FILE):
        try:
            with open(NAVER_ACCOUNTS_FILE, encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, OSError) as e:
            print(f"[naver_accounts] 로드 오류: {e}")
    return []


def _naver_save_accounts(accounts):
    with _naver_accounts_lock:
        tmp = NAVER_ACCOUNTS_FILE + '.tmp'
        with open(tmp, 'w', encoding='utf-8') as f:
            json.dump(accounts, f, ensure_ascii=False, indent=2)
        os.replace(tmp, NAVER_ACCOUNTS_FILE)


# ── 상위글 분석 ──

def _analyze_cafe_article(url, keyword):
    """개별 카페 글 분석: 사진수, 키워드반복수, 글자수 (3-tier 폴백)"""
    from src.services.cafe_crawler import crawl_cafe_article as _crawl_impl
    result = _crawl_impl(url)
    return result.to_analysis_dict(keyword)


def _extract_cafe_urls(soup, urls, top_titles, seen_articles, max_count):
    """BeautifulSoup에서 카페 URL 추출 (광고 필터링 포함)"""
    added = 0
    for a in soup.find_all('a', href=lambda h: h and 'cafe.naver.com' in h):
        if added >= max_count:
            break
        href = a.get('href', '')
        text = a.get_text(strip=True)

        # RE(댓글) 제외, 짧은 텍스트 제외
        if text.startswith('RE') or len(text) < 10:
            continue

        # 글 번호가 포함된 URL만
        article_match = re.search(r'cafe\.naver\.com/[^/]+/(\d+)', href)
        if not article_match:
            continue
        article_id = article_match.group(1)
        if article_id in seen_articles:
            continue

        # 광고 필터링: 부모 요소 class에 ad/sponsor 포함 시 제외
        parent_div = a.find_parent(['li', 'div', 'section'])
        if parent_div:
            classes = ' '.join(parent_div.get('class', []))
            if re.search(r'ad_|_ad|sponsor', classes):
                continue
            # 컨테이너 앞부분에 "광고" 텍스트 있으면 제외
            first_text = parent_div.get_text()[:30]
            if '광고' in first_text:
                continue

        seen_articles.add(article_id)
        top_titles.append(text)
        urls.append(href.split('?')[0])
        added += 1


def _analyze_top_for_cafe(keyword):
    """통합검색 우선 → 카페탭 보충, 상위글 5개 제목+URL 수집 → 상위 3개 본문 분석"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    top_titles = []
    urls = []
    seen_articles = set()

    try:
        # Phase 1: 통합검색(nexearch) — 카페 글 최대 3개
        try:
            r = req.get(f"https://search.naver.com/search.naver?query={quote(keyword)}&where=nexearch",
                        headers=headers, timeout=10)
            soup = BeautifulSoup(r.text, 'html.parser')
            _extract_cafe_urls(soup, urls, top_titles, seen_articles, max_count=3)
        except Exception:
            pass

        # Phase 2: 카페탭(article) — 나머지 보충 (총 5개까지)
        remaining = 5 - len(top_titles)
        if remaining > 0:
            try:
                r = req.get(f"https://search.naver.com/search.naver?query={quote(keyword)}&where=article",
                            headers=headers, timeout=10)
                soup = BeautifulSoup(r.text, 'html.parser')
                _extract_cafe_urls(soup, urls, top_titles, seen_articles, max_count=remaining)
            except Exception:
                pass

        if not urls:
            return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500, 'top_titles': []}

        results = []
        articles = []
        for url in urls[:3]:
            a = _analyze_cafe_article(url, keyword)
            if a:
                results.append(a)
                articles.append({'url': url, 'char_count': a['char_count'], 'photo_count': a['photo_count'], 'keyword_repeat': a['keyword_repeat']})
            time.sleep(0.5)
        if not results:
            return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500, 'top_titles': top_titles, 'articles': []}
        return {
            'photo_count': max(round(sum(r['photo_count'] for r in results) / len(results)), 3),
            'keyword_repeat': max(round(sum(r['keyword_repeat'] for r in results) / len(results)), 3),
            'char_count': max(round(sum(r['char_count'] for r in results) / len(results)), 800),
            'top_titles': top_titles,
            'articles': articles,
        }
    except Exception as e:
        print(f"[_analyze_top_for_cafe] 오류: {e}")
        return {'photo_count': 8, 'keyword_repeat': 5, 'char_count': 1500, 'top_titles': []}


# ── 크롤링 ──

def _search_cafe_top_article(keyword):
    """네이버 검색에서 카페 상위 1위 글 URL+제목 찾기"""
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36'}
    try:
        r = req.get('https://search.naver.com/search.naver?query=%s&where=nexearch' % quote(keyword), headers=headers, timeout=10)
        soup = BeautifulSoup(r.text, 'html.parser')
        # 카페 링크 찾기
        for a in soup.find_all('a', href=re.compile(r'cafe\.naver\.com')):
            href = a.get('href', '')
            if '/ArticleRead' in href or re.search(r'cafe\.naver\.com/[^/]+/\d+', href):
                title_text = a.get_text(strip=True)
                return {'url': href, 'title': title_text}
        # fallback: 카페 검색 탭
        r2 = req.get('https://search.naver.com/search.naver?query=%s&where=article' % quote(keyword), headers=headers, timeout=10)
        soup2 = BeautifulSoup(r2.text, 'html.parser')
        for a in soup2.find_all('a', href=re.compile(r'cafe\.naver\.com')):
            href = a.get('href', '')
            if '/ArticleRead' in href or re.search(r'cafe\.naver\.com/[^/]+/\d+', href):
                title_text = a.get_text(strip=True)
                return {'url': href, 'title': title_text}
    except Exception as e:
        print("Cafe search error: %s" % e)
    return {'url': '', 'title': ''}


def _crawl_cafe_article(url):
    """네이버 카페 글 크롤링 (3-tier 폴백) — 기존 시그니처 호환"""
    from src.services.cafe_crawler import crawl_cafe_article as _crawl_impl
    result = _crawl_impl(url)
    return result.to_legacy_dict()


# ── 이미지 수집 ──

def _translate_ko_to_zh(text):
    """한국어 → 중국어 번역 (Gemini Flash)"""
    try:
        url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key=" + GEMINI_API_KEY
        payload = {
            "contents": [{"parts": [{"text": "다음 한국어를 중국어(간체)로 번역해줘. 번역 결과만 출력하고 다른 설명은 하지 마.\n\n" + text}]}]
        }
        r = req.post(url, json=payload, timeout=10)
        data = r.json()
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except Exception:
        return text


def _crawl_xhs_images(query_zh, count):
    """샤오홍슈 이미지 수집 (xiaohongshu-cli 사용)"""
    results = []
    try:
        result = subprocess.run(
            [XHS_PATH, "search", query_zh, "--sort", "popular", "--json"],
            capture_output=True, text=True, timeout=30
        )
        data = json.loads(result.stdout)
        items = data.get("data", {}).get("items", [])

        for item in items:
            if len(results) >= count:
                break
            nc = item.get("note_card", {})
            # 이미지 타입 게시물만 (영상 제외)
            if nc.get("type") == "video":
                continue
            for img in nc.get("image_list", []):
                if len(results) >= count:
                    break
                # WB_DFT = 고해상도
                url = ""
                for info in img.get("info_list", []):
                    if info.get("image_scene") == "WB_DFT":
                        url = info.get("url", "")
                        break
                if not url:
                    url = img.get("info_list", [{}])[0].get("url", "") if img.get("info_list") else ""
                if url and url.startswith("http"):
                    results.append(url)
    except Exception as e:
        print(f"[xhs-cli] search error: {e}")
    return results


def _mosaic_faces(image_path):
    """OpenCV로 얼굴 감지 후 모자이크 처리"""
    try:
        img = cv2.imread(image_path)
        if img is None:
            return False
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        faces = cascade.detectMultiScale(gray, 1.1, 4, minSize=(30, 30))
        for (x, y, w, h) in faces:
            roi = img[y:y+h, x:x+w]
            small = cv2.resize(roi, (max(1, w//10), max(1, h//10)))
            mosaic = cv2.resize(small, (w, h), interpolation=cv2.INTER_NEAREST)
            img[y:y+h, x:x+w] = mosaic
        cv2.imwrite(image_path, img)
        return len(faces) > 0
    except Exception as e:
        print(f"[mosaic] error: {e}")
        return False


def _collect_xhs_for_cafe(target, count):
    """타겟층 기반으로 샤오홍슈 이미지 수집 + 모자이크"""
    # 타겟층 → 중국어 검색 키워드 변환
    zh_query = _translate_ko_to_zh(target + " 일상")
    print(f"[xhs-cafe] target='{target}' → query='{zh_query}'")

    urls = _crawl_xhs_images(zh_query, count + 3)
    if not urls:
        print("[xhs-cafe] no images found, trying simpler query")
        zh_query = _translate_ko_to_zh(target)
        urls = _crawl_xhs_images(zh_query, count + 3)

    downloaded = []
    headers = {"User-Agent": "Mozilla/5.0", "Referer": "https://www.xiaohongshu.com/"}
    for i, url in enumerate(urls):
        if len(downloaded) >= count:
            break
        try:
            r = req.get(url, headers=headers, timeout=15)
            if r.status_code == 200 and len(r.content) > 2000:
                fname = f"xhs_cafe_{int(time.time()*1000)}_{i}.jpg"
                fpath = os.path.join(TEMP_PHOTO_DIR, fname)
                with open(fpath, 'wb') as f:
                    f.write(r.content)
                _mosaic_faces(fpath)
                downloaded.append(fpath)
        except Exception:
            pass
        time.sleep(2)  # 차단 방지
    return downloaded


# ── DOCX 생성 ──

def _create_cafe_docx(title, body_text, image_paths, keyword):
    """카페 원고 docx 생성"""
    doc = DocxDocument()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    img_idx = 0
    for line in body_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if '[어울릴 사진' in line or '[이미지' in line:
            if img_idx < len(image_paths) and os.path.exists(image_paths[img_idx]):
                try:
                    doc.add_picture(image_paths[img_idx], width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(line)
                img_idx += 1
            else:
                doc.add_paragraph(line)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_kw = re.sub(r'[\\/*?:"<>|]', '', keyword)
    fname = f"{safe_kw}_카페원고_{ts}.docx"
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return fpath, fname


# ── 프롬프트 빌더 ──

def _build_cafe_title_prompt(keyword, original_title, top_titles=None):
    system = """📌 역할:
너는 네이버 카페 상위노출을 위한 제목 리라이팅 전문가야.

[목표]
기존에 상위노출된 카페 제목을 기반으로, 메인 키워드를 유지하면서 마지막 말(어미·단어)만 자연스럽게 바꿔주는 리라이팅 작업을 수행한다.
광고처럼 보이지 않고, 궁금형/공감형 톤으로 바꿔 클릭을 유도할 수 있게 한다.

[지침]
1. 제목의 전체 구조는 유지
2. 마지막 말(어미 or 단어)만 자연스럽게 바꿔줘
3. 클릭을 유도할 수 있는 궁금형 or 공감형 문장으로 바꾸기
4. 광고 느낌 X ("추천", "인생템", "꼭 사세요" 등 금지)
5. 이모티콘 사용 금지
6. 문장 길이 25~35자 이내
7. 정보 탐색형 문장으로 구성 — '궁금증'을 자극하면 클릭률(CTR)이 올라감

[제목 유형 참고]
- 고민형: "(30대 여자) 출산탈모 때문에 미치겠어요ㅠㅠ"
- 직관형: "40대 되면 출산탈모 심해지나요?"
- 궁금형: "다이어트 영양제 진짜 효과 있나요?"
- 결핍형: "다이어트 영양제 먹고 붓기 빠졌다는 후기 봤는데 진짜인가요?"

[잘못된 제목 vs 올바른 제목]
- ❌ 다이어트영양제 추천드려요 → ✅ 다이어트 영양제 진짜 효과 있나요?
- ❌ 체지방 줄이는 약 후기 → ✅ 체지방 감소 도움된다는 영양제, 효과 있을까요?
- ❌ 출산탈모 걱정이네요 → ✅ (30대 여자) 출산탈모 때문에 미치겠어요ㅠㅠ

[출력 규칙]
- 출력은 제목 1줄만
- 불필요한 부가 문구 없이 제공
- 예시: 입력 "다이어트 자전거 도움되나요?" → 출력 "다이어트 자전거 효과 있나요?\""""
    titles_text = ''
    if top_titles:
        titles_text = '\n[카페 탭 상위 노출 제목들 — 톤과 패턴 참고용]\n' + '\n'.join(f'- {t}' for t in top_titles[:5])

    if original_title:
        user = "- 기존 제목: %s\n- 메인 키워드: %s → 제목에 그대로 유지%s" % (original_title, keyword, titles_text)
    else:
        user = "- 기존 제목: (없음 - 새로 생성)\n- 메인 키워드: %s → 제목에 그대로 유지%s" % (keyword, titles_text)
    return system, user


def _build_cafe_body_prompt(keyword, title, original_body, settings, product):
    repeat_count = settings.get('repeat_count', 5)
    sub_keywords = settings.get('sub_keywords', '')
    forbidden = product.get('forbidden', '')
    photo_count = settings.get('photo_count', 8)
    char_target = settings.get('char_target', 1500)
    ref_text = original_body if original_body else '(없음 - 새로 생성)'

    system = """📌 역할:
너는 네이버 카페에 고민을 털어놓는 실제 사용자야.
제품 후기를 쓰는 게 아니라, 진짜 힘든 상황을 호소하면서 댓글로 도움을 구하는 글을 써.

⚠️ 핵심: 이 글은 "고민글"이지 "후기글"이 아니야.
- 본문에서 특정 제품명, 성분명, 효과 체험을 절대 언급하지 마.
- 제품/해결책은 댓글에서 다른 사람들이 알려주는 구조야.
- 글쓴이는 아직 해결 못 한 사람이야.

[본문 구조]
1. 인사 + 고민 꺼내기 — "안녕하세요 ~~ 관련해서 질문 드리려고 합니다"
2. 내 상황 설명 — 나이, 직업, 생활패턴 등 구체적 페르소나 + 증상 호소
   예: "요즘 컨디션이 너무 안좋아져서 스트레스가 심해요ㅠㅠ"
3. 이것저것 알아본 경험 — 병원, 약, 다른 방법 알아봤지만 부담/걱정
   예: "팔팔정가격 알아보다가 생각보다 부담스러워서 고민되더라고요.."
4. 현재 막막한 심정 — 근본적 해결이 필요한데 뭘 해야 할지 모르겠음
   예: "일시적으로 좋아지는 건 의미 없잖아요. 근본적으로 개선되는 방법이 필요한데 그게 쉽지가 않네요ㅠ"
5. 도움 요청 — 경험담 공유 부탁
   예: "혹시 건강한 방식으로 좋아지신 분 계신가요??"

[키워드 규칙]
- 메인 키워드: 지정된 반복횟수 +1회 본문에 삽입
- 서브 키워드: 지정된 횟수 +1회 삽입
- 키워드 누락 시 실패
- 금칙어 포함 시 실패

[사진 표시 규칙]
- 문단과 문단 사이 또는 문장 중간에 어색하지 않게 [어울릴 사진] 삽입이라고 표기
- 어떤 이미지를 넣으면 좋을지 짧은 설명도 함께 작성
- 사진 수: 지정된 수 +1개 삽입
- 사진 유형: 고민하는 모습, 병원/약국 관련, 검색하는 모습 등 고민 상황에 맞는 사진
- 제품 촬영 컷, 홍보용 배너 이미지 금지
- 예시:
  - [어울릴 사진: 소파에 앉아 고민하는 중년 남성의 일상 컷]
  - [어울릴 사진: 스마트폰으로 건강 정보를 검색하는 모습]

[문체 규칙]
- 줄바꿈은 2~4줄 간격. 빈 줄로 문단 구분.
- 실제 카페에 고민 올리는 사람의 말투. 존댓말 기반이지만 편한 톤.
- 이모티콘(ㅠㅠ, ㅜㅜ, .., ??) 자연스럽게 사용
- 광고 어투 절대 금지
- 제품 추천, 성분 설명, 효과 체험 서술 금지 — 이건 댓글에서 할 일

[출력 규칙]
- 출력 시 제목 없이 바로 본문만
- 글자 수: 지정된 글자수 ±100자 (공백 제외)

🚫 절대 금지:
- 본문에서 특정 제품명 언급
- 본문에서 "이걸 먹어봤는데", "효과가 있었다" 식의 후기 서술
- 본문에서 특정 성분 효능 설명
- 해시태그 사용"""

    user = """- 제목: %s
- 참고 경쟁사글 (없다면 무시하고 새로 작성): %s
- 메인 키워드: %s
- 메인 키워드 반복횟수: %d → 총 %d에서 +1회 본문 삽입
- 서브 키워드: %s → 서브 키워드 +1회 삽입
- 글자 수: %d ±100자 (공백 제외)
- 금칙어 목록: %s → 포함 시 실패
- 사진 수: %d +1개 삽입""" % (title, ref_text, keyword, repeat_count, repeat_count, sub_keywords, char_target, forbidden, photo_count)

    return system, user


def _build_cafe_comments_prompt(keyword, body_text, brand_keyword, alternatives=''):
    alt_section = ""
    if alternatives:
        alt_section = "\n사용자가 입력한 차선책 정보:\n\n" + alternatives
    else:
        alt_section = """
아래 기본 차선책을 참고하여 댓글을 작성하세요.

자연 식이요법 (식단 관리 및 음식으로 개선)
장점: 건강한 식습관으로 항산화 성분 및 영양 섭취, 비교적 안전함
단점: 꾸준히 지키기 어려움, 식단만으로 체지방 감소나 노화 개선 효과에는 한계
메시지 포인트: "음식만으로는 체지방을 줄이고 활력을 되찾기엔 한계가 있더라."

기본 영양제 및 비타민 (멀티비타민, 콜라겐 등)
장점: 필요한 기본 영양 보충 가능, 간편함
단점: 특허 성분이나 체지방 감소 등 특화 기능성은 부족, 떨어진 컨디션을 끌어올리기에는 역부족
메시지 포인트: "일반 영양제만으론 예전 같지 않은 체력이나 피부 탄력까지 끌어올리기엔 부족하더라구요."

전통 다이어트 방법 (한약, 디톡스 주스 등)
장점: 몸에 좋다고 전해지는 방법으로 일부 디톡스 효과 기대, 심리적 위안
단점: 과학적·임상적 검증 부족, 고가의 한약은 경제적 부담, 디톡스 주스만으로는 영양 불균형 초래
메시지 포인트: "민간요법이라고 좋다 해서 해봤지만, 솔직히 큰 효과는 못 느꼈고 돈만 들었어요."

피부과 시술 및 휴식 (전문의 관리, 일시적 휴식)
장점: 전문 관리와 휴식을 통해 일시적으로 피부나 컨디션 개선에 도움
단점: 단기간 일시적 케어에 그침 – 근본적인 체질 개선이나 체중 관리는 어려움, 지속 시 비용 부담 큼
메시지 포인트: "클리닉 가서 시술도 받아봤지만 잠깐일 뿐이고, 근본적으로 달라지는 건 없더라구요."

운동 및 식단 다이어트 (자체 체중 감량 시도)
장점: 체중 감량과 건강 회복에 도움, 특별한 비용 없음
단점: 바쁜 일상 속 운동 시간 확보 어려움, 무리한 식단 조절은 건강을 해칠 수 있음, 꾸준한 실천 어려움
메시지 포인트: "운동이나 식이조절이 좋다는 건 알지만, 막상 바쁘게 지내다 보면 실천이 거의 불가능하더라구요."

일반 다이어트 보조제 (다이어트 약, 카페인 함유 제품 등)
장점: 빠른 체중 감량 효과 기대, 복용 간편
단점: 카페인 등 성분으로 부작용 위험 (두근거림, 불면 등), 건강에 무리가 갈 수 있음, 요요 등 지속적 관리 어려움
메시지 포인트: "효과 빠르다는 다이어트 약은 부작용도 걱정되고 솔직히 무서워서 시도도 못 하겠더라구요."
"""

    system = """역할:
너는 네이버 카페 글에 달릴 자연스러운 댓글 10개를 생성하는 작가야.
댓글은 실제 사용자의 다양한 반응처럼 보여야 하고, 제품 삽입이 자연스러워야 해.

[작성 정보]

제품의 차선책 (경쟁 대안)
{alt_section}

[지시사항]

상황: 바쁜 일상과 스트레스 탓인지 피로, 기억력 저하, 체중 증가와 피부 탄력 저하로 고민하는 사람이 쓴 게시글에 달린 댓글들을 작성합니다. 각 댓글은 서로 다른 사람이 작성한 것으로, 앞 사람의 의견에 공감하거나 자신의 경험을 추가하면서 자연스럽게 이어지는 흐름이어야 합니다. 모든 댓글은 일상적인 존댓말 체로 쓰되, 너무 딱딱하지 않게 편안하고 친근한 어조를 사용하세요. 필요한 경우 감정을 나타내는 이모티콘이나 표현(예: ㅠㅠ, ㅎㅎ, ;;, 👍 등)을 적절히 섞어 자연스러운 분위기를 만드세요.

1~2번째 댓글 (공감 및 차선책 언급): 문제 상황에 깊이 공감하며 비슷한 어려움을 겪어봤다는 경험담으로 시작하세요. 많은 사람들이 흔히 시도하는 대안들을 언급하고, 본인도 열심히 해봤지만 효과를 거의 못 봤거나 일시적이었다는 식으로 써주세요. (주의: 이 단계에서는 특정 제품 언급이나 광고 티가 나면 안 됩니다!)

3~4번째 댓글 (해결책 제시 – 제품명 언급 없이 장점 소개): 새로운 해결 방법을 제시하는 흐름으로 전환합니다. 구체적으로 특화된 건강기능식품이 요즘 나와 있다는 정보를 공유하되, 아직 제품명은 직접 언급하지 마세요. (여전히 제품명이나 구매 유도는 하지 않습니다.)

5~6번째 댓글 (제품명 공개 및 경험담으로 추천): 앞에서 언급한 보조제가 무엇인지 자연스럽게 제품명으로 연결합니다. 마치 이미 써본 사람들처럼, 제품 복용 후 구체적인 변화를 경험담 형태로 작성하세요. 한 댓글에서는 제품을 쓰고 효과를 체감한 내용을 담습니다. 또 다른 댓글에서는 제품의 강점(특허 성분 함유, 가성비 등)과 함께 주효과를 언급하세요. 각각 다양한 사용자의 입장에서 제품을 사용해 만족했다는 흐름을 보여주세요. (링크, 구매 강조 ❌ 단, 제품명이 노출됩니다.)

7번째 댓글 (효과 언급): 이어서 제품명을 한 번 더 등장시켜, 고민 해결 경험담을 써주세요. 제품 복용 후 효과를 본 사례를 구체적으로 적습니다. 안전하게 효과를 봤다는 점을 강조하세요. 추가로 개인 소감을 덧붙여줍니다.

8~10번째 댓글 (추가 공감, 조언 및 마무리): 마지막 댓글들에서는 앞선 대화를 자연스럽게 이어받으며 마무리합니다. 다른 사람들도 등장하여 영양제 관리에 공감하고 꾸준한 자기관리의 중요성을 강조하세요. "방치하면 더 악화된다"는 경고를 살짝 넣어 건강 관리의 필요성을 상기시키고, "꾸준히 하면 좋아진다"는 희망적인 조언을 덧붙이세요. 사회적 증거와 긍정적 분위기로 끝맺습니다. 독자들이 "나도 해봐야겠다"는 마음이 들도록 유도하며 대화를 마무리하세요.

각 댓글에 메인 키워드 1개 이상은 포함

광고 어투 X (강추, 무조건, 최고 등 과장 표현 및 구매 유도 금지)

[출력형식]
한줄에 한개의 댓글을 모두 작성하며, 각각의 댓글 사이의 줄바꿈은 1회만 할 것. 절대 줄바꿈을 2회하지 않을 것 (중요).

[형식]
맞아요... 저도 요즘 부쩍 피로도 심하고 몸이 예전같지 않아서 좋다는 음식이며 영양제까지 열심히 챙겨먹어봤는데도 피로 회복이나 체중감량엔 효과가 거의 없었어요 ㅠㅠ 역시 식단 조절만으로는 한계가 있나봐요...
저도 휴가도 내서 푹 쉬어보고 좋다는 영양제도 찾아 먹고, 한약이랑 디톡스 주스까지 열심히 따라해봤는데 바쁘게 지내다 보니 금세 지치고 기억력 떨어진 건 그대로이더라구요 ㅠㅠ 피부 푸석해진 게 신경 쓰여서 피부과 시술까지 받아봤지만 솔직히 잠깐 나아진 듯한 느낌만 들고 비용만 많이 들었어요;; 운동이나 다이어트도 해보려 했지만 일하고 집안일 하랴 꾸준히 하기가 제 마음처럼 잘 안 되더라구요. 정말 쉽지 않아요.
요즘은 집에서 꾸준히 챙겨먹을 수 있는 건강기능식품들이 잘 나와 있더라구요. 단순 종합비타민제가 아니라 떨어진 체력이나 기억력, 푸석해진 피부 회복부터 체중감량까지 도움을 주는 특허 성분의 영양제들이요.
저도 최근에 그런 다이어트 영양제를 하나 먹기 시작했는데, 확실히 전보다 덜 피곤하고 멍했던 정신도 맑아지는 느낌이에요 👍 제품 고를 땐 원료가 특허받았는지, 또 식약처 기능성 인정을 받은 건지 꼭 확인해야 된다더라구요!
맞아요, 항산화 다이어트 영양제 중에서는 후기가 제일 좋길래 저도 다이어트 시작하면서 함께 먹어봤거든요? 꾸준히 먹었더니 몸도 한결 가벼워지고 피부톤도 환해지는 것 같아요. 괜히 많이들 찾는 게 아니더라구요.
국내 유일 특허 조합이 들어간 다이어트 영양제인데 가격도 생각보다 괜찮아서 좋았어요ㅠㅠ 저는 사실 나이 들면서 늘어난 뱃살 때문에 걱정이 많았는데, 먹고 나서 주변에서 얼굴 좋아졌단 소리도 듣고 체지방률도 조금 내려가서 만족 중이에요 🙂
저는 살 찐 게 가장 고민이었는데 식약처에서 기능성 인정받은 다이어트 영양제가 있어서 시작했어요. 카페인 들어간 다이어트제가 아니라 국내 특허 받은 폴리페놀 성분으로 만들었다고 해서 그런지 카페인 민감한 저도 부담 없이 천천히 살이 빠지더라구요. 요요 없고 안전하게 빼고 있어서 너무 만족해요!
사실 다이어트 약이나 병원 치료 같은 데에 의존하는 건 부담도 크고 장기적으로도 어렵잖아요… 이렇게 집에서 꾸준히 영양제로 관리하는 게 확실히 답인 것 같아요.
몸 컨디션 떨어지고 체중도 늘고 피부 노화 오는 걸 그냥 방치하면 나중에 체중감량이나 피부 회복이 더 힘들어진다고 하더라구요. 그러니까 영양제라도 미리 챙겨두는 게 낫다고 해서 저도 열심히 챙겨보려구요.
많은 분들이 이렇게 꾸준히 영양제 챙겨드시면서 기운도 되찾고 체중감량에도 성공하셨다니까 저도 계속 챙겨보려구요. 병원 안 가고도 이렇게 관리할 수 있다는 게 진짜 다행이고 신기하네요ㅎㅎ""".format(alt_section=alt_section)

    alt_display = alternatives if alternatives else '(AI 자동 판단)'
    user = f"""- 실제 본문: {body_text[:2000]}
- 메인 키워드: {keyword}
- 제품명 or 간접 언급: {brand_keyword or keyword}
- 차선책 (경쟁 대안, 선택사항): {alt_display}"""

    return system, user


def _build_cafe_replies_prompt(keyword, body_text, comments_text, brand_keyword):
    """답글 10개 생성 프롬프트 — 게시글 작성자 관점에서 각 댓글에 1:1 답글."""
    system = """역할:
너는 네이버 카페 게시글 작성자야. 네 글에 달린 댓글 하나하나에 자연스럽고 친근한 답글을 작성해.

[핵심 원칙]
1. 답글 작성자 = 게시글 작성자. 게시글의 흐름, 감정, 메시지와 싱크(일관성)를 맞추는 것이 가장 중요.
2. 톤: 30~40대 여성 커뮤니티 말투. 부드럽고 진솔하며 친근하게.
3. 이모티콘(ㅎㅎ, ㅠㅠ, 🙂 등)은 상황에 맞게 적절히 사용.
4. 답글당 2~4문장. 짧고 자연스럽게.
5. "나만의 키워드"는 답글 전체 중 1~2회만 자연스럽게 노출. 과하면 광고 티남.
6. 광고 어투 금지: "강추", "꼭 사세요", "인생템" 등 과장 표현 X.
7. 게시글에서 "힘들다"고 했는데 답글에서 "이거 써보세요~ 진짜 좋아요!"처럼 앞뒤 안 맞는 답글 금지.

[답글 유형별 톤]
- 공감 댓글(1~2번) → "맞아요ㅠㅠ 저도 그랬어요" 식으로 공감 + 경험 공유
- 정보 제공 댓글(3~4번) → "오 그런 게 있나요? 저도 찾아봐야겠어요" 식으로 관심 표현
- 제품 언급 댓글(5~7번) → "그거 어디서 구매하셨어요?" / "저도 요즘 챙겨먹고 있어요" 식으로 경험 나눔
- 마무리 댓글(8~10번) → "댓글 보니까 힘이 나네요ㅎㅎ" / "다들 건강 챙기시는 거 보니 저도 힘내볼게요" 식으로 감사+의지

[출력 규칙]
- 댓글 원문을 먼저 쓰고, 바로 다음 줄에 답글을 작성
- 형식:
1. 댓글 원문
→ 답글 내용

2. 댓글 원문
→ 답글 내용

(총 10개)"""

    user = f"""[게시글 본문]
{body_text[:2000]}

--
[댓글 목록]
{comments_text[:3000]}

- 메인 키워드: {keyword}
- 나만의 키워드 (답글 중 1~2회만 노출): {brand_keyword or keyword}"""

    return system, user


def _build_cafe_polish_prompt(keyword, title, body_text, comments_text, brand_keyword, forbidden_words=''):
    """금칙어 대체 + 타사→자사 치환 + 문체 보정 프롬프트."""
    system = """역할:
너는 네이버 카페 원고 교정 전문가야. AI가 생성한 초안을 사람이 쓴 것처럼 자연스럽게 다듬는 작업을 수행해.

[작업 항목]

1. 금칙어 검사 + 자동 대체
   - 금칙어가 발견되면 자연스러운 대체어로 바꿔줘
   - 단, 상위노출 키워드가 금칙어에 해당하면 그대로 유지
   - 자주 나오는 금칙어 대체 예시:
     병원 → 동네 / 진료 → 진찰 / 가격 → 금액
     하더라구요 → 했어요 / 바르고 자도 → 바르구 자도

2. 타사 제품/브랜드 → 자사 키워드 치환
   - 본문/댓글에서 타사 제품명이나 브랜드명이 있으면 "나만의 키워드"로 자연스럽게 치환
   - 예: "혹시 메디힐 시카크림 써보신 분?" → "요즘 (나만의 키워드) 좋다는데 써보신 분?"

3. 문체 보정
   - 광고 어투 → 자연스러운 후기/체험담 톤으로
   - 딱딱한 문장 → 편안한 대화체로
   - 키워드 반복 구조는 유지하되 문장 흐름만 자연스럽게

4. 제목 강화
   - 직관적+감정형으로 보강
   - 나이/성별 표시 가능 (예: "(30대 여자) 출산탈모 때문에 미치겠어요ㅠㅠ")

[출력 형식]
JSON으로 출력:
{
  "title": "보정된 제목",
  "body": "보정된 본문",
  "comments": "보정된 댓글",
  "changes": [
    {"type": "금칙어", "before": "원문", "after": "수정문", "reason": "사유"},
    {"type": "타사치환", "before": "원문", "after": "수정문", "reason": "사유"},
    {"type": "문체", "before": "원문", "after": "수정문", "reason": "사유"}
  ]
}
- changes에 변경 내역을 모두 기록
- 변경사항이 없으면 changes를 빈 배열로"""

    user = f"""[원고 정보]
- 메인 키워드: {keyword}
- 나만의 키워드 (자사 제품): {brand_keyword or keyword}
- 금칙어 목록: {forbidden_words or '(없음)'}

[제목]
{title}

[본문]
{body_text[:3000]}

[댓글]
{comments_text[:2000]}"""

    return system, user


# ── 엔드포인트 ──

@router.get("/notion-keywords")
async def cafe_notion_keywords():
    """노션 키워드 DB에서 카페SEO 배정 키워드 조회"""
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    payload = {
        'filter': {
            'and': [
                {'property': '배정 채널', 'multi_select': {'contains': '카페SEO'}},
                {'property': '상태', 'select': {'equals': '미사용'}},
            ]
        },
        'page_size': 100,
    }
    try:
        from src.services.notion_client import query_database
        data = query_database(KEYWORD_DB_ID, filter_obj=payload['filter'], page_size=100)
        keywords = []
        for page in data.get('results', []):
            props = page.get('properties', {})
            title_prop = props.get('키워드', {}).get('title', [])
            kw = title_prop[0]['text']['content'] if title_prop else ''
            if kw:
                keywords.append({'keyword': kw, 'page_id': page['id']})
        return {'keywords': keywords}
    except Exception as e:
        return {'keywords': [], 'error': str(e)}


@router.post("/build-prompt")
async def cafe_build_prompt(request: Request):
    """카페SEO 프롬프트만 생성 (크롤링+분석, 본문/댓글은 claude.ai용)"""
    body = await request.json()
    keywords = body.get('keywords', [])
    urls = body.get('urls', [])
    product = body.get('product', {})
    settings = body.get('settings', {})
    user_title = body.get('user_title', '')
    loop = asyncio.get_running_loop()
    results = []
    url_list = [u.strip() for u in urls if u.strip()] if urls else []

    for i, kw_data in enumerate(keywords):
        kw = kw_data['keyword']
        url = url_list[i] if i < len(url_list) else ''
        original_title = ''
        original_body = ''

        # 상위글 분석
        cafe_analysis = await loop.run_in_executor(executor, _analyze_top_for_cafe, kw)
        kw_settings = {**settings}
        kw_settings['photo_count'] = cafe_analysis['photo_count']
        kw_settings['repeat_count'] = cafe_analysis['keyword_repeat']
        kw_settings['char_target'] = cafe_analysis['char_count']

        # 경쟁사 크롤링
        if url:
            crawled = await loop.run_in_executor(executor, _crawl_cafe_article, url)
            original_title = crawled['title']
            original_body = crawled['body']
        else:
            found = await loop.run_in_executor(executor, _search_cafe_top_article, kw)
            if found['url']:
                crawled = await loop.run_in_executor(executor, _crawl_cafe_article, found['url'])
                original_title = crawled['title'] or found['title']
                original_body = crawled['body']

        # 제목 프롬프트 조립 (API 호출 X — claude.ai에서 직접 테스트)
        sys1, usr1 = _build_cafe_title_prompt(kw, original_title, top_titles=cafe_analysis.get('top_titles', []))
        combined_title = f"다음 시스템 프롬프트의 역할을 수행해주세요.\n\n---\n\n{sys1}\n\n---\n\n{usr1}"
        title = user_title if user_title else '(제목란에 입력하세요)'

        # 본문 프롬프트 조립
        sys2, usr2 = _build_cafe_body_prompt(kw, title, original_body, kw_settings, product)
        combined_body = f"다음 시스템 프롬프트의 역할을 수행해주세요.\n\n---\n\n{sys2}\n\n---\n\n{usr2}"

        # 댓글 프롬프트 조립
        sys3, usr3 = _build_cafe_comments_prompt(kw, '(본문은 위에서 생성한 결과를 넣어주세요)', product.get('brand_keyword', ''), product.get('alternatives', ''))
        combined_comments = f"다음 시스템 프롬프트의 역할을 수행해주세요.\n\n---\n\n{sys3}\n\n---\n\n{usr3}"

        # 답글 프롬프트 조립
        sys4, usr4 = _build_cafe_replies_prompt(kw, '(본문+댓글 생성 후 입력)', '(댓글 생성 결과를 넣어주세요)', product.get('brand_keyword', ''))
        combined_replies = f"다음 시스템 프롬프트의 역할을 수행해주세요.\n\n---\n\n{sys4}\n\n---\n\n{usr4}"

        results.append({
            'keyword': kw,
            'title': title,
            'analysis': {'photo_count': kw_settings['photo_count'], 'keyword_repeat': kw_settings['repeat_count'], 'char_count': kw_settings['char_target']},
            'crawled': {'top_titles': cafe_analysis.get('top_titles', []), 'original_title': original_title, 'original_body': original_body[:2000], 'articles': cafe_analysis.get('articles', [])},
            'title_prompt': {'system_prompt': sys1, 'user_prompt': usr1, 'combined': combined_title},
            'body_prompt': {'system_prompt': sys2, 'user_prompt': usr2, 'combined': combined_body},
            'comments_prompt': {'system_prompt': sys3, 'user_prompt': usr3, 'combined': combined_comments},
            'replies_prompt': {'system_prompt': sys4, 'user_prompt': usr4, 'combined': combined_replies},
        })

    return {'results': results}


@router.post("/generate")
async def cafe_generate(request: Request):
    """카페SEO 원고 생성 (SSE): 제목→본문→댓글"""
    body = await request.json()
    keywords = body.get('keywords', [])
    urls = body.get('urls', [])
    product = body.get('product', {})
    settings = body.get('settings', {})

    _sse = sse_dict

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(keywords)
        url_list = [u.strip() for u in urls if u.strip()] if urls else []

        for i, kw_data in enumerate(keywords):
            kw = kw_data['keyword']
            url = url_list[i] if i < len(url_list) else ''
            original_title = ''
            original_body = ''

            # 상위글 분석 (사진수/키워드반복수/글자수 자동 계산)
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 상위글 분석 중 (사진수/반복수/글자수)...' % (i+1, total, kw), 'cur': i, 'total': total})
            cafe_analysis = await loop.run_in_executor(executor, _analyze_top_for_cafe, kw)
            kw_settings = {**settings}  # 키워드별 로컬 복사
            kw_settings['photo_count'] = cafe_analysis['photo_count']
            kw_settings['repeat_count'] = cafe_analysis['keyword_repeat']
            kw_settings['char_target'] = cafe_analysis['char_count']
            kw_settings['top_titles'] = cafe_analysis.get('top_titles', [])

            # 경쟁사 글 크롤링 (URL 직접 입력 or 자동 검색)
            if url:
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 경쟁사 글 크롤링 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                crawled = await loop.run_in_executor(executor, _crawl_cafe_article, url)
                original_title = crawled['title']
                original_body = crawled['body']
            else:
                # URL 없으면 네이버 검색으로 카페 상위글 자동 크롤링
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 카페 상위글 검색 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                found = await loop.run_in_executor(executor, _search_cafe_top_article, kw)
                if found['url']:
                    yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 상위글 크롤링 중...' % (i+1, total, kw), 'cur': i, 'total': total})
                    crawled = await loop.run_in_executor(executor, _crawl_cafe_article, found['url'])
                    original_title = crawled['title'] or found['title']
                    original_body = crawled['body']

            # STEP 1: 제목
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 제목 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys1, usr1 = _build_cafe_title_prompt(kw, original_title, top_titles=kw_settings.get('top_titles', []))
            title = await loop.run_in_executor(executor, call_claude, sys1, usr1)
            title = title.strip().split('\n')[0].strip()

            # STEP 2: 본문
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 본문 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys2, usr2 = _build_cafe_body_prompt(kw, title, original_body, kw_settings, product)
            cafe_body = await loop.run_in_executor(executor, call_claude, sys2, usr2)
            cafe_body = cafe_body.strip()

            # STEP 3: 댓글 10개
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 댓글 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys3, usr3 = _build_cafe_comments_prompt(kw, cafe_body, product.get('brand_keyword', ''), product.get('alternatives', ''))
            comments = await loop.run_in_executor(executor, call_claude, sys3, usr3)
            comments = comments.strip()

            # STEP 3b: 답글 10개
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 답글 생성 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys_r, usr_r = _build_cafe_replies_prompt(kw, cafe_body, comments, product.get('brand_keyword', ''))
            replies = await loop.run_in_executor(executor, call_claude, sys_r, usr_r)
            replies = replies.strip()

            # STEP 4: 타겟 맞춤 사진 수집 (샤오홍슈)
            target = product.get('target', '')
            photo_count = kw_settings.get('photo_count', 8)
            # 본문에서 [어울릴 사진] 개수 세기
            img_slots = len(re.findall(r'\[어울릴 사진[^\]]*\]|\[이미지\d+\]', cafe_body))
            need_count = max(img_slots, photo_count) if img_slots else photo_count
            image_paths = []
            if target:
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 사진 수집 중 (샤오홍슈)...' % (i+1, total, kw), 'cur': i, 'total': total})
                image_paths = await loop.run_in_executor(executor, _collect_xhs_for_cafe, target, need_count)
                yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 사진 %d장 수집+모자이크 완료' % (i+1, total, kw, len(image_paths)), 'cur': i, 'total': total})

            # 이미지 파일명 목록 (프론트에 전달)
            image_filenames = [os.path.basename(p) for p in image_paths]

            # STEP 5: Polish (금칙어 대체 + 타사치환 + 문체 보정)
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 원고 보정(Polish) 중...' % (i+1, total, kw), 'cur': i, 'total': total})
            sys_pol, usr_pol = _build_cafe_polish_prompt(
                kw, title, cafe_body, comments,
                product.get('brand_keyword', ''), product.get('forbidden', ''),
            )
            polish_raw = await loop.run_in_executor(executor, call_claude, sys_pol, usr_pol)
            polish_changes = []
            try:
                json_match = re.search(r'\{[\s\S]*\}', polish_raw)
                if json_match:
                    polish_data = json.loads(json_match.group())
                    polish_changes = polish_data.get('changes', [])
                    # Polish 결과 적용 (변경사항이 있을 때만)
                    if polish_data.get('title'):
                        title = polish_data['title']
                    if polish_data.get('body'):
                        cafe_body = polish_data['body']
                    if polish_data.get('comments'):
                        comments = polish_data['comments']
            except (json.JSONDecodeError, ValueError):
                pass  # polish 실패 시 원본 유지
            yield _sse({'type': 'progress', 'msg': '[%d/%d] %s — 보정 완료 (변경 %d건)' % (i+1, total, kw, len(polish_changes)), 'cur': i, 'total': total})

            result = {
                'keyword': kw, 'title': title, 'body': cafe_body, 'comments': comments,
                'replies': replies,
                'original_title': original_title, 'original_body': original_body,
                'page_id': kw_data.get('page_id', ''),
                'photo_count': kw_settings.get('photo_count', 8),
                'repeat_count': kw_settings.get('repeat_count', 5),
                'images': image_filenames,
                'polish_changes': polish_changes,
            }

            # ── 검수 단계 ──
            yield _sse({'type': 'progress', 'msg': f'[{i+1}/{total}] {kw} — 검수 중...', 'cur': i, 'total': total})
            review_result = await loop.run_in_executor(
                executor, review_and_save, "cafe-seo", result, kw,
            )
            for ev in review_result.get("events", []):
                yield _sse(ev)
            result['review_status'] = review_result["status"]
            result['review_passed'] = review_result["passed"]
            result['revision_count'] = review_result["revision_count"]
            result['project_id'] = review_result["project_id"]

            yield _sse({'type': 'result', 'data': result, 'cur': i+1, 'total': total})

        yield _sse({'type': 'complete', 'total': total})
      except Exception as e:
        print(f"[cafe_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'카페SEO 원고 생성 중 오류: {e}'})

    return SSEResponse(generate())


@router.post("/save-notion")
async def cafe_save_notion(request: Request):
    """카페SEO 원고를 노션 콘텐츠 DB에 저장"""
    body = await request.json()
    headers = {
        'Authorization': 'Bearer %s' % NOTION_TOKEN,
        'Content-Type': 'application/json',
        'Notion-Version': '2022-06-28',
    }
    props = {
        '제목': {'title': [{'text': {'content': body.get('title', '')}}]},
        '채널': {'select': {'name': '카페'}},
        '생산 상태': {'select': {'name': '승인됨' if body.get('review_status') == 'approved' else '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    if body.get('body_summary'):
        props['본문'] = {'rich_text': [{'text': {'content': body['body_summary'][:2000]}}]}
    if body.get('page_id'):
        props['키워드'] = {'relation': [{'id': body['page_id']}]}
    if body.get('photo_count') is not None:
        props['사진수'] = {'number': body['photo_count']}

    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}

    # 본문 + 댓글 + 답글 + Polish 변경점을 페이지 children으로
    text_blocks = [
        body.get('body', ''),
        '---\n댓글 10개:\n' + body.get('comments', ''),
    ]
    if body.get('replies'):
        text_blocks.append('---\n답글 10개:\n' + body['replies'])
    if body.get('polish_changes'):
        changes_text = '---\nPolish 변경점:\n'
        for ch in body['polish_changes']:
            if isinstance(ch, dict):
                changes_text += f"- [{ch.get('type','')}] {ch.get('before','')} → {ch.get('after','')} ({ch.get('reason','')})\n"
            else:
                changes_text += f"- {ch}\n"
        text_blocks.append(changes_text)

    children = []
    for text_block in text_blocks:
        paragraphs = [p.strip() for p in text_block.split('\n\n') if p.strip()]
        for para in paragraphs:
            for k in range(0, len(para), 2000):
                children.append({
                    'object': 'block', 'type': 'paragraph',
                    'paragraph': {'rich_text': [{'type': 'text', 'text': {'content': para[k:k+2000]}}]}
                })
    payload['children'] = children[:100]

    try:
        from src.services.notion_client import create_page
        result = create_page(CONTENT_DB_ID, props, children=payload.get('children'))
        return {'success': result['success'], 'error': result.get('error', '')}
    except Exception as e:
        return {'success': False, 'error': str(e)}


@router.get("/temp-image/{filename}")
async def cafe_temp_image(filename: str):
    """카페 원고용 임시 이미지 서빙"""
    fpath = os.path.join(TEMP_PHOTO_DIR, filename)
    if os.path.exists(fpath):
        return FileResponse(fpath, media_type="image/jpeg")
    return Response(status_code=404)


@router.post("/polish")
async def cafe_polish(request: Request):
    """금칙어 대체 + 타사→자사 치환 + 문체 보정"""
    body = await request.json()
    keyword = body.get('keyword', '')
    title = body.get('title', '')
    body_text = body.get('body', '')
    comments = body.get('comments', '')
    product = body.get('product', {})
    brand_keyword = product.get('brand_keyword', '')
    forbidden = product.get('forbidden', '')

    loop = asyncio.get_running_loop()
    sys_p, usr_p = _build_cafe_polish_prompt(
        keyword, title, body_text, comments, brand_keyword, forbidden,
    )
    try:
        raw = await loop.run_in_executor(executor, call_claude, sys_p, usr_p)
    except Exception as e:
        print(f"[cafe_polish] AI 호출 오류: {e}")
        return {
            'polished_title': title, 'polished_body': body_text,
            'polished_comments': comments, 'changes': [],
            'original_title': title, 'original_body': body_text,
            'original_comments': comments, 'error': str(e),
        }

    # JSON 파싱 시도
    try:
        # Claude 출력에서 JSON 블록 추출
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if json_match:
            result = json.loads(json_match.group())
        else:
            result = {'title': title, 'body': body_text, 'comments': comments, 'changes': []}
    except (json.JSONDecodeError, ValueError):
        result = {'title': title, 'body': body_text, 'comments': comments, 'changes': []}

    return {
        'polished_title': result.get('title', title),
        'polished_body': result.get('body', body_text),
        'polished_comments': result.get('comments', comments),
        'changes': result.get('changes', []),
        'original_title': title,
        'original_body': body_text,
        'original_comments': comments,
    }


@router.post("/match-photos")
async def cafe_match_photos(request: Request):
    """본문의 [어울릴 사진] 태그 파싱 → XHS 이미지 수집 매칭"""
    body = await request.json()
    body_text = body.get('body', '')
    target = body.get('target', '')
    existing_images = body.get('existing_images', [])

    # 본문에서 [어울릴 사진: ...] 태그 파싱
    photo_tags = re.findall(r'\[어울릴 사진[：:]?\s*([^\]]*)\]', body_text)
    if not photo_tags:
        photo_tags = re.findall(r'\[이미지\d*[：:]?\s*([^\]]*)\]', body_text)

    need_count = max(len(photo_tags) - len(existing_images), 0)
    new_images = []

    if need_count > 0 and target:
        try:
            loop = asyncio.get_running_loop()
            new_images_paths = await loop.run_in_executor(
                executor, _collect_xhs_for_cafe, target, need_count,
            )
            new_images = [os.path.basename(p) for p in new_images_paths]
        except Exception as e:
            print(f"[cafe_match_photos] 이미지 수집 오류: {e}")
            new_images = []

    all_images = existing_images + new_images
    # 태그별 매칭
    matched = []
    for i, tag in enumerate(photo_tags):
        matched.append({
            'tag_index': i,
            'description': tag.strip(),
            'image': all_images[i] if i < len(all_images) else None,
        })

    return {
        'photo_tags': photo_tags,
        'matched': matched,
        'new_images': new_images,
        'total_images': len(all_images),
    }


@router.post("/build-reply")
async def cafe_build_reply(request: Request):
    """답글 10개 생성"""
    body = await request.json()
    keyword = body.get('keyword', '')
    body_text = body.get('body', '')
    comments = body.get('comments', '')
    product = body.get('product', {})
    brand_keyword = product.get('brand_keyword', '')

    loop = asyncio.get_running_loop()
    sys_r, usr_r = _build_cafe_replies_prompt(keyword, body_text, comments, brand_keyword)
    try:
        replies = await loop.run_in_executor(executor, call_claude, sys_r, usr_r)
    except Exception as e:
        print(f"[cafe_build_reply] AI 호출 오류: {e}")
        return {'replies': '', 'error': str(e)}

    return {'replies': replies.strip()}


@router.post("/docx")
async def cafe_docx(request: Request):
    """카페 원고를 docx 파일로 생성 + 다운로드"""
    body = await request.json()
    title = body.get('title', '')
    body_text = body.get('body', '')
    keyword = body.get('keyword', '')
    image_filenames = body.get('images', [])

    image_paths = [os.path.join(TEMP_PHOTO_DIR, f) for f in image_filenames]
    loop = asyncio.get_running_loop()
    fpath, fname = await loop.run_in_executor(executor, _create_cafe_docx, title, body_text, image_paths, keyword)

    safe_fname = quote(fname)
    return FileResponse(
        fpath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename*=UTF-8''" + safe_fname}
    )


@router.post("/auto-comment")
async def cafe_auto_comment(request: Request):
    """카페 댓글 자동 등록 (SSE 스트리밍)"""
    body = await request.json()
    post_url = body.get('post_url', '')
    comments = body.get('comments', [])
    account_ids = body.get('account_ids', [])

    if not post_url or not comments:
        return JSONResponse({'error': '게시글 URL과 댓글 필요'}, 400)

    naver_accounts = _naver_load_accounts()
    accounts = []
    for acc_id in account_ids:
        acc = next((a for a in naver_accounts if a.get('id') == acc_id), None)
        if acc:
            accounts.append(acc)

    if len(accounts) < len(comments):
        return JSONResponse({'error': f'계정 부족: 댓글 {len(comments)}개인데 계정 {len(accounts)}개'}, 400)

    _sse = sse_dict

    async def generate():
      try:
        loop = asyncio.get_running_loop()
        total = len(comments)
        from src.cafe_comment_bot import run_auto_comments

        yield _sse({'type': 'progress', 'msg': f'댓글 {total}개 자동 등록 시작...', 'cur': 0, 'total': total})

        results = await loop.run_in_executor(
            executor, run_auto_comments,
            post_url, comments, accounts[:len(comments)], None, None
        )

        for i, r in enumerate(results):
            status = '✅' if r.get('success') else '❌'
            yield _sse({'type': 'result', 'data': r, 'cur': i+1, 'total': total,
                         'msg': f'[{i+1}/{total}] {r.get("label","")} — {status} {r.get("error","")}'})
            # read-modify-write를 하나의 lock 블록으로 보호
            if r.get('status_change') == '정지 의심' or r.get('success'):
                with _naver_accounts_lock:
                    # lock 내부에서는 save가 중첩 lock을 걸지 않도록 load만 직접 수행
                    import json as _json
                    try:
                        with open(NAVER_ACCOUNTS_FILE, 'r', encoding='utf-8') as f:
                            accs = _json.load(f)
                    except (FileNotFoundError, _json.JSONDecodeError):
                        accs = []
                    for a in accs:
                        if a['id'] == r.get('account_id'):
                            if r.get('status_change') == '정지 의심':
                                a['status'] = '정지 의심'
                            if r.get('success'):
                                a['total_posts'] = a.get('total_posts', 0) + 1
                                a['last_used_at'] = datetime.now().isoformat()
                            break
                    # save도 lock 내부에서 직접 수행 (중첩 lock 방지)
                    tmp = NAVER_ACCOUNTS_FILE + '.tmp'
                    with open(tmp, 'w', encoding='utf-8') as f:
                        _json.dump(accs, f, ensure_ascii=False, indent=2)
                    os.replace(tmp, NAVER_ACCOUNTS_FILE)

        success_count = sum(1 for r in results if r.get('success'))
        yield _sse({'type': 'complete', 'total': total, 'success': success_count, 'fail': total - success_count})
      except Exception as e:
        print(f"[cafe_auto_comment] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'댓글 자동 등록 중 오류: {e}'})

    return SSEResponse(generate())


@router.get("/comment-history")
async def cafe_comment_history():
    from src.cafe_safety_rules import get_history
    return {'history': get_history(50)}
