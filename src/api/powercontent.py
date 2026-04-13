"""파워컨텐츠 API 라우터"""
import asyncio
import json
import os
import re
from datetime import datetime

import requests as req
from bs4 import BeautifulSoup
from fastapi import APIRouter, Request
from fastapi.responses import FileResponse
from src.services.sse_helper import sse_dict, SSEResponse

from src.services.config import executor, BASE_DIR, CONTENT_DB_ID, NOTION_TOKEN
from src.services.common import error_response
from src.services.ai_client import call_claude
from src.services.review_service import review_and_save
from src.services.notion_client import notion_headers

router = APIRouter()

OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(OUTPUTS_DIR, exist_ok=True)


# ═══════════════════════════ HELPERS ═══════════════════════════

def _crawl_power_content(url):
    """파워컨텐츠 랜딩페이지 크롤링"""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'}
        r = req.get(url, headers=headers, timeout=15, allow_redirects=True)
        soup = BeautifulSoup(r.text, 'html.parser')
        # 본문 추출 시도
        for sel in ['.se-main-container', '.content_view', '#content', '.article_body', '.post_ct', 'article', 'main']:
            el = soup.select_one(sel)
            if el and len(el.get_text(strip=True)) > 200:
                return el.get_text('\n', strip=True)[:10000]
        # 폴백: body 전체
        body = soup.find('body')
        if body:
            for tag in body.find_all(['script','style','nav','header','footer']):
                tag.decompose()
            return body.get_text('\n', strip=True)[:10000]
        return ''
    except Exception as e:
        print("[power-content] crawl error: %s" % e)
        return ''


def _build_pc_analysis_prompt(ref_text):
    system = "당신은 파워컨텐츠(설득형 랜딩페이지) 구조 분석 전문가입니다."
    user = """아래 파워컨텐츠 본문을 분석해.

%s

분석 항목:
1. total_chars: 전체 글자수 (공백 포함)
2. sections: 오프닝(1막)/솔루션(2막)/클로징(3막) 각각의 글자수와 비율
3. opening_method: 오프닝 전개 방식
4. solution_method: 솔루션 전개 방식 + 소거 개수
5. closing_method: 클로징 전개 방식
6. image_count: 이미지 삽입 위치와 개수
7. cta_count: CTA 위치와 개수
8. hooking_type: 후킹 방식 (부정편향/비밀폭로/비교검증)
9. keyword_repeats: 키워드 반복 횟수 추정
10. persuasion_flow: 전체 설득 흐름 (한 줄 요약)

JSON으로 응답해.""" % ref_text
    return system, user


def _build_pc_ad_prompt(keyword, appeal, buying_thing, product, hooking_type, forbidden):
    system = "당신은 네이버 파워컨텐츠 CTR 최적화 전문가입니다."
    user = """키워드와 소구점을 기반으로 네이버 파워컨텐츠 광고 소재를 생성해.

## 입력
- 메인 키워드: %s
- 소구점: %s
- 구매원씽: %s
- 제품 정보: 제품명=%s, 성분=%s, USP=%s
- 레퍼런스 후킹 방식: %s

## 출력 규격
1. 제목: 최대 28자 (공백 포함), 메인 키워드 반드시 포함 (문장이 다소 어색하더라도 반드시 삽입)
2. 설명: 최대 110자 (공백 포함), 메인 키워드 반드시 포함, 궁금증 유발 + 혜택/근거/가치

제목 유형 (레퍼런스와 같은 유형):
- 부정편향: "절대 사면 안 되는 OOO", "효과 없었던 진짜 이유"
- 비밀폭로: "업계가 숨기는 OOO", "전문가는 안 알려주는"
- 비교검증: "3개 써보고 솔직 비교", "성분 분석해봤더니"

출력 형식:
제목: (제목)
설명: (설명)

%s
## 필수 규칙
- 제목과 도입부를 소구점(%s)의 결핍 언어로 시작한다. 고객이 "이건 완전 내 얘기잖아?"라고 느끼게 해야 한다.
- 금지: 허위·과장, 근거 없는 효능, 낚시""" % (
        keyword, appeal, buying_thing,
        product.get('name',''), product.get('ingredients',''), product.get('usp',''),
        hooking_type, forbidden, appeal
    )
    return system, user


def _build_pc_body_prompt(keyword, stage, appeal, buying_thing, deficit_level, product, ad_title, ad_desc, analysis_json):
    system = "당신은 네이버 파워컨텐츠 랜딩 본문 전문 카피라이터입니다. 멘토 3단 공식 + BA 설득 기법으로 작성합니다."

    if stage in ('0_무지','1_인지','2_호기심','3_정보습득'):
        template = """[템플릿 A -- 구매여정 0~3]

1막 오프닝: 문제를 '심각한 질병'으로 격상
- ⛔ 도입부에서 제품 이야기를 절대 먼저 꺼내지 않는다. 고객의 결핍부터 파고든다.
- BA 카모플라주: 건강정보/전문가 분석 포맷으로 진입
- BA 점진화: 독자가 동의하는 증상에서 출발, 동의 3~5회 축적
- 결핍 마케팅: 방치 시 미래를 구체적으로 묘사
- 권위: "해외 논문", "전문가 발견" 등

2막 솔루션: 기존 상식 파괴 + 새 길 제시
- BA 집중화: 좌절 서사로 경쟁 경로 차단
- 소거법: 기존 대안 소거 → 우리 제품만 남김
- 성분 비교표 등으로 새 기준 증명

3막 클로징: 첫 번째 선택지 되기
- BA 메커니제이션: 작동 원리를 판매 언어로
- BA 재정의 가격축소: "하루 OOO원"
- 위장 진정성: "저는 광고 목적이 아닌, 진짜 효과 본 정보를 공유하고 싶었을 뿐입니다" 식으로 상업적 의도 방어벽을 무너뜨린다
- CTA: "현명한 첫걸음을 내딛으세요" + [CTA: %s]""" % product.get('url','')
    else:
        template = """[템플릿 B -- 구매여정 4~5]

1막 오프닝: 잘못된 선택의 위험성 경고
- ⛔ 도입부에서 제품 이야기를 절대 먼저 꺼내지 않는다. 잘못된 선택의 위험부터 경고한다.
- BA 카모플라주: "직접 써보고 비교한 분석가" 포맷
- BA 점진화: 효과 없는 제품에 돈/시간 낭비가 진짜 문제

2막 솔루션: 경쟁자 압살
- BA 집중화: 경쟁사별 구체적 약점
- 소거법: A사 핵심 성분 없음, B사 함량 미달 → 우리만 충족

3막 클로징: 즉시 구매 유도
- FOMO + 가성비 증명
- 위장 진정성: "저는 광고 목적이 아닌, 진짜 효과 본 정보를 공유하고 싶었을 뿐입니다" 식으로 상업적 의도 방어벽을 무너뜨린다
- CTA: "계속 돈을 낭비하시겠습니까?" + [CTA: %s]""" % product.get('url','')

    user = """아래 조건으로 파워컨텐츠 랜딩 본문을 작성해.

## 입력
- 광고 소재: 제목=%s / 설명=%s
- 메인 키워드: %s
- 구매여정: %s
- 소구점: %s / 구매원씽: %s / 결핍강도 목표: %s
- 제품: 이름=%s, 성분=%s, USP=%s, 타겟=%s, URL=%s

## 레퍼런스 구조 (이 구조를 그대로 따르되 내용만 교체)
%s

## 템플릿
%s

## 충돌방지 3규칙
1. 결핍강도 -3~-5, 한 문단 -2 이상 급락 금지
2. BA 재정의 우선, 공략집 프레이밍은 참고만
3. 섹션 순서 = 멘토 3단 공식 고정

## 공통 규칙
- 레퍼런스 전체 글자수와 비슷하게 (최소 2,500자, ±200자)
- 각 막 비율도 레퍼런스와 유사
- 메인 키워드 5~8회 자연스럽게 삽입
- 이미지/CTA 위치·개수도 레퍼런스와 동일
- [이미지N - 설명], [CTA: URL], [공정위 문구 삽입 위치] 표기

## 출력: 본문만 출력 (부연 설명 없이)""" % (
        ad_title, ad_desc, keyword, stage, appeal, buying_thing, deficit_level,
        product.get('name',''), product.get('ingredients',''), product.get('usp',''),
        product.get('target',''), product.get('url',''),
        analysis_json, template
    )
    return system, user


# ═══════════════════════════ ENDPOINTS ═══════════════════════════

@router.post("/analyze")
async def pc_analyze(request: Request):
    """레퍼런스 크롤링 + 구조 분석"""
    body = await request.json()
    url = body.get('url', '')
    raw_text = body.get('raw_text', '')
    loop = asyncio.get_running_loop()

    if url and not raw_text:
        raw_text = await loop.run_in_executor(executor, _crawl_power_content, url)
    if not raw_text:
        return {'error': '본문을 가져올 수 없습니다', 'analysis': ''}

    sys_p, usr_p = _build_pc_analysis_prompt(raw_text)
    analysis = await loop.run_in_executor(executor, call_claude, sys_p, usr_p)
    return {'raw_text': raw_text[:5000], 'analysis': analysis}


@router.post("/build-prompt")
async def pc_build_prompt(request: Request):
    """파워컨텐츠 프롬프트만 생성 (광고소재까지 서버, 본문은 claude.ai용)"""
    body = await request.json()
    keyword = body.get('keyword', '')
    stage = body.get('stage', '3_정보습득')
    appeal = body.get('appeal', '')
    buying_thing = body.get('buying_thing', '')
    deficit_level = body.get('deficit_level', '-4')
    product = body.get('product', {})
    analysis = body.get('analysis', '')
    forbidden = product.get('forbidden', '')
    loop = asyncio.get_running_loop()

    # 구조분석에서 hooking_type/total_chars 추출
    hooking_type = '부정편향'
    try:
        m = re.search(r'"hooking_type"\s*:\s*"([^"]+)"', analysis)
        if m: hooking_type = m.group(1)
    except (AttributeError, ValueError): pass
    target_chars = 3000
    try:
        m = re.search(r'"total_chars"\s*:\s*(\d+)', analysis)
        if m: target_chars = max(2500, int(m.group(1)))
    except (AttributeError, ValueError): pass

    # 광고 소재 생성 (API — 짧아서 비용 미미)
    sys2, usr2 = _build_pc_ad_prompt(keyword, appeal, buying_thing, product, hooking_type, forbidden)
    ad_raw = await loop.run_in_executor(executor, call_claude, sys2, usr2)
    ad_title = ad_desc = ''
    for line in ad_raw.split('\n'):
        line = line.strip()
        if line.startswith('제목:') or line.startswith('제목 :'):
            ad_title = line.split(':', 1)[1].strip()
        elif line.startswith('설명:') or line.startswith('설명 :'):
            ad_desc = line.split(':', 1)[1].strip()
    if not ad_title:
        ad_title = ad_raw.split('\n')[0].strip()[:28]

    # 본문 프롬프트 조립 (API 호출 X)
    sys3, usr3 = _build_pc_body_prompt(keyword, stage, appeal, buying_thing, deficit_level, product, ad_title, ad_desc, analysis)
    combined = f"다음 시스템 프롬프트의 역할을 수행해주세요.\n\n---\n\n{sys3}\n\n---\n\n{usr3}"

    return {
        'ad_title': ad_title,
        'ad_desc': ad_desc,
        'target_chars': target_chars,
        'body_prompt': {
            'system_prompt': sys3,
            'user_prompt': usr3,
            'combined': combined,
        }
    }


@router.post("/generate")
async def pc_generate(request: Request):
    """파워컨텐츠 생성 (SSE): 구조분석 완료 후 -> 광고소재 -> 본문"""
    body = await request.json()
    keyword = body.get('keyword', '')
    stage = body.get('stage', '3_정보습득')
    appeal = body.get('appeal', '')
    buying_thing = body.get('buying_thing', '')
    deficit_level = body.get('deficit_level', '-4')
    product = body.get('product', {})
    analysis = body.get('analysis', '')
    forbidden = product.get('forbidden', '')

    # 구조분석에서 hooking_type 추출
    hooking_type = '부정편향'
    try:
        m = re.search(r'"hooking_type"\s*:\s*"([^"]+)"', analysis)
        if m: hooking_type = m.group(1)
    except (AttributeError, ValueError): pass

    # 구조분석에서 total_chars 추출
    target_chars = 3000
    try:
        m = re.search(r'"total_chars"\s*:\s*(\d+)', analysis)
        if m: target_chars = max(2500, int(m.group(1)))
    except (AttributeError, ValueError): pass

    _sse = sse_dict

    async def generate():
      try:
        loop = asyncio.get_running_loop()

        # STEP 2: 광고 소재 생성
        yield _sse({'type':'progress','msg':'광고 소재 생성 중 (제목 28자 + 설명 110자)...'})
        sys2, usr2 = _build_pc_ad_prompt(keyword, appeal, buying_thing, product, hooking_type, forbidden)
        ad_raw = await loop.run_in_executor(executor, call_claude, sys2, usr2)
        ad_title = ad_desc = ''
        for line in ad_raw.split('\n'):
            line = line.strip()
            if line.startswith('제목:') or line.startswith('제목 :'):
                ad_title = line.split(':', 1)[1].strip()
            elif line.startswith('설명:') or line.startswith('설명 :'):
                ad_desc = line.split(':', 1)[1].strip()
        if not ad_title:
            ad_title = ad_raw.split('\n')[0].strip()[:28]

        yield _sse({'type':'ad','title':ad_title,'desc':ad_desc})

        # STEP 3: 랜딩 본문 생성 (글자수 검증 포함, 최대 3회 재시도)
        yield _sse({'type':'progress','msg':'랜딩 본문 생성 중 (목표 %d자)...' % target_chars})
        sys3, usr3 = _build_pc_body_prompt(keyword, stage, appeal, buying_thing, deficit_level, product, ad_title, ad_desc, analysis)
        body_text = ''
        for attempt in range(3):
            body_text = await loop.run_in_executor(executor, call_claude, sys3, usr3)
            char_count = len(body_text)
            if char_count >= 2500 and abs(char_count - target_chars) <= 400:
                break
            # 재시도
            yield _sse({'type':'progress','msg':'글자수 부족 (%d자). 재생성 중 (%d/3)...' % (char_count, attempt+2)})
            usr3 = "글자수가 부족합니다. 현재 %d자입니다. %d자 이상으로 다시 작성해주세요. 이전 내용의 구조를 유지하면서 각 단락을 더 풍부하게 확장해주세요.\n\n" % (char_count, target_chars) + usr3

        char_count = len(body_text)

        # ── 검수 단계 ──
        yield _sse({'type': 'progress', 'msg': '검수 중...'})
        pc_content = {'ad_title': ad_title, 'ad_desc': ad_desc, 'body': body_text, 'char_count': char_count}
        review_result = await loop.run_in_executor(
            executor, review_and_save, "powercontent", pc_content, keyword,
        )
        for ev in review_result.get("events", []):
            yield _sse(ev)
        review_status = review_result["status"]
        review_passed = review_result["passed"]

        yield _sse({'type':'result','ad_title':ad_title,'ad_desc':ad_desc,'body':body_text,'char_count':char_count,'target_chars':target_chars,'review_status':review_status,'review_passed':review_passed})
        yield _sse({'type':'complete'})
      except Exception as e:
        print(f"[powercontent_generate] 에러: {e}")
        yield _sse({'type': 'error', 'message': f'파워컨텐츠 생성 중 오류: {e}'})

    return SSEResponse(generate())


@router.post("/docx")
async def pc_docx(request: Request):
    """파워컨텐츠 docx 생성"""
    body = await request.json()
    ad_title = body.get('ad_title', '')
    ad_desc = body.get('ad_desc', '')
    body_text = body.get('body', '')
    keyword = body.get('keyword', '')

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 광고 소재
    h = doc.add_heading('파워컨텐츠 광고 소재', level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run('제목: ' + ad_title)
    r1.bold = True
    r1.font.size = Pt(14)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('설명: ' + ad_desc)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph('─' * 40)
    doc.add_heading('랜딩 본문', level=1)

    for line in body_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('[이미지') or line.startswith('[CTA') or line.startswith('[공정위'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 120, 200)
                run.italic = True
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_kw = re.sub(r'[\\/*?:"<>|]', '', keyword)
    fname = "%s_파워컨텐츠_%s.docx" % (safe_kw, ts)
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return FileResponse(fpath, filename=fname, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.post("/save-notion")
async def pc_save_notion(request: Request):
    body = await request.json()
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    props = {
        '제목': {'title': [{'text': {'content': body.get('ad_title', '')}}]},
        '채널': {'select': {'name': '파워컨텐츠'}},
        '생산 상태': {'select': {'name': '승인됨' if body.get('review_status') == 'approved' else '초안'}},
        '발행_상태': {'select': {'name': '미발행'}},
    }
    summary = body.get('body', '')[:300]
    props['본문'] = {'rich_text': [{'text': {'content': summary}}]}
    payload = {'parent': {'database_id': CONTENT_DB_ID}, 'properties': props}
    # 본문을 페이지 children으로
    body_text = body.get('body', '')
    if body_text:
        paras = [p.strip() for p in body_text.split('\n\n') if p.strip()]
        children = []
        for para in paras[:100]:
            children.append({'object':'block','type':'paragraph','paragraph':{'rich_text':[{'type':'text','text':{'content':para[:2000]}}]}})
        payload['children'] = children
    try:
        from src.services.notion_client import create_page
        result = create_page(CONTENT_DB_ID, props, children=payload.get('children'))
        return {'success': result['success'], 'error': result.get('error', '')}
    except Exception as e:
        return {'success': False, 'error': str(e)}
