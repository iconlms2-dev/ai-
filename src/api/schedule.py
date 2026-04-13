"""스케줄/리포트/스케줄러 API 라우터"""
import asyncio
import json
import os
from datetime import datetime, timedelta, date as dt_date

import requests as req
from fastapi import APIRouter, Request
from fastapi.responses import JSONResponse

from src.services.config import BASE_DIR, NOTION_TOKEN, CONTENT_DB_ID, executor
from src.services.common import error_response
from src.services.ai_client import call_claude

router = APIRouter()
report_router = APIRouter()
scheduler_router = APIRouter()

OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(OUTPUTS_DIR, exist_ok=True)

WEEKLY_SCHEDULE_FILE = os.path.join(BASE_DIR, "weekly_schedule.json")
_scheduler_notifications: list = []


# ═══════════════════════════ HELPERS ═══════════════════════════

def _notion_query_by_date(date_str, channel=None):
    """노션 콘텐츠 DB에서 특정 날짜의 콘텐츠 집계"""
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    filters = [{'property': '생성일', 'date': {'equals': date_str}}]
    if channel:
        filters.append({'property': '채널', 'select': {'equals': channel}})
    payload = {'filter': {'and': filters}, 'page_size': 100}
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % CONTENT_DB_ID, headers=headers, json=payload, timeout=15)
        if r.status_code == 200:
            return r.json().get('results', [])
        return []
    except Exception:
        return []


def _count_by_channel(results):
    """노션 결과를 채널별 카운트"""
    counts = {}
    for page in results:
        props = page.get('properties', {})
        ch = props.get('채널', {}).get('select', {})
        ch_name = ch.get('name', '기타') if ch else '기타'
        counts[ch_name] = counts.get(ch_name, 0) + 1
    return counts


def _notion_query_range(start_date, end_date, db_id=None):
    """노션 DB에서 기간 내 콘텐츠 조회"""
    headers = {'Authorization': 'Bearer %s' % NOTION_TOKEN, 'Content-Type': 'application/json', 'Notion-Version': '2022-06-28'}
    payload = {
        'filter': {'and': [
            {'property': '생성일', 'date': {'on_or_after': start_date}},
            {'property': '생성일', 'date': {'on_or_before': end_date}},
        ]},
        'page_size': 100
    }
    target_db = db_id or CONTENT_DB_ID
    all_results = []
    try:
        r = req.post('https://api.notion.com/v1/databases/%s/query' % target_db, headers=headers, json=payload, timeout=20)
        if r.status_code == 200:
            data = r.json()
            all_results.extend(data.get('results', []))
            while data.get('has_more'):
                payload['start_cursor'] = data['next_cursor']
                r = req.post('https://api.notion.com/v1/databases/%s/query' % target_db, headers=headers, json=payload, timeout=20)
                if r.status_code != 200:
                    break
                data = r.json()
                all_results.extend(data.get('results', []))
    except Exception as e:
        print("[report] query error: %s" % e)
    return all_results


def _build_report_data(start, end, results):
    """리포트 데이터 구축"""
    channels = ['블로그','카페','지식인','카페바이럴','커뮤니티','파워컨텐츠','유튜브','틱톡','쓰레드','메타광고']
    production = {}
    published = {}
    for page in results:
        props = page.get('properties', {})
        ch_sel = props.get('채널', {}).get('select')
        ch = ch_sel.get('name', '기타') if ch_sel else '기타'
        production[ch] = production.get(ch, 0) + 1
        deploy_sel = props.get('발행_상태', {}).get('select')
        deploy = deploy_sel.get('name', '') if deploy_sel else ''
        if deploy == '발행완료' or deploy == '발행완료':
            published[ch] = published.get(ch, 0) + 1
    rows = []
    total_prod = total_pub = 0
    for ch in channels:
        p = production.get(ch, 0)
        pub = published.get(ch, 0)
        unpub = p - pub
        rate = round(pub / p * 100) if p > 0 else 0
        if p > 0:
            rows.append({'channel': ch, 'produced': p, 'published': pub, 'unpublished': unpub, 'rate': rate})
            total_prod += p
            total_pub += pub
    total_rate = round(total_pub / total_prod * 100) if total_prod > 0 else 0
    return {
        'period': '%s ~ %s' % (start, end),
        'rows': rows,
        'total': {'produced': total_prod, 'published': total_pub, 'unpublished': total_prod - total_pub, 'rate': total_rate}
    }


# ═══════════════════════════ SCHEDULER HELPERS ═══════════════════════════

def _sched_load():
    if os.path.exists(WEEKLY_SCHEDULE_FILE):
        try:
            return json.loads(open(WEEKLY_SCHEDULE_FILE, encoding='utf-8').read())
        except (json.JSONDecodeError, OSError):
            pass
    return {
        'daily': {
            'generate_remind': {'enabled': True, 'time': '09:00', 'label': '콘텐츠 생성'},
            'review_remind': {'enabled': True, 'time': '14:00', 'label': '콘텐츠 검수'},
            'deploy_remind': {'enabled': True, 'time': '17:00', 'label': '콘텐츠 배포'},
        },
        'weekly': {
            'keyword_analysis': {'enabled': True, 'day': 'mon', 'time': '09:00', 'auto_run': False, 'label': '키워드 분석'},
            'channel_assign': {'enabled': True, 'day': 'mon', 'time': '10:00', 'auto_run': False, 'label': '채널 배정'},
            'performance_collect': {'enabled': True, 'day': 'fri', 'time': '17:00', 'auto_run': True, 'label': '성과 수집'},
            'weekly_report': {'enabled': True, 'day': 'fri', 'time': '18:00', 'auto_run': True, 'label': '주간 리포트'},
        },
        'history': [],
    }


def _sched_save(data):
    tmp = WEEKLY_SCHEDULE_FILE + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, WEEKLY_SCHEDULE_FILE)


_DAY_MAP = {'mon': 0, 'tue': 1, 'wed': 2, 'thu': 3, 'fri': 4, 'sat': 5, 'sun': 6}


# ═══════════════════════════ SCHEDULE ENDPOINTS ═══════════════════════════

@router.get("/today")
async def schedule_today():
    """오늘 생산 현황"""
    today = datetime.now().strftime('%Y-%m-%d')
    loop = asyncio.get_running_loop()
    results = await loop.run_in_executor(executor, _notion_query_by_date, today)
    counts = _count_by_channel(results)
    return {'date': today, 'counts': counts, 'total': len(results)}


@router.get("/week")
async def schedule_week():
    """이번 주 월~금 일별 생산 현황"""
    today = datetime.now().date()
    monday = today - timedelta(days=today.weekday())
    loop = asyncio.get_running_loop()
    week_data = []
    for i in range(7):
        d = monday + timedelta(days=i)
        ds = d.isoformat()
        results = await loop.run_in_executor(executor, _notion_query_by_date, ds)
        counts = _count_by_channel(results)
        week_data.append({'date': ds, 'weekday': ['월','화','수','목','금','토','일'][i], 'counts': counts, 'total': len(results)})
    return {'week': week_data, 'today': today.isoformat()}


# ═══════════════════════════ REPORT ENDPOINTS ═══════════════════════════

@report_router.post("/generate")
async def report_generate(request: Request):
    """주간 리포트 생성"""
    body = await request.json()
    start = body.get('start', '')
    end = body.get('end', '')
    loop = asyncio.get_running_loop()

    # 현재 기간 데이터
    results = await loop.run_in_executor(executor, _notion_query_range, start, end)
    report = _build_report_data(start, end, results)

    # 이전 기간 데이터 (같은 길이만큼 이전)
    d_start = dt_date.fromisoformat(start)
    d_end = dt_date.fromisoformat(end)
    period_days = (d_end - d_start).days + 1
    prev_end = d_start - timedelta(days=1)
    prev_start = prev_end - timedelta(days=period_days - 1)
    prev_results = await loop.run_in_executor(executor, _notion_query_range, prev_start.isoformat(), prev_end.isoformat())
    prev_report = _build_report_data(prev_start.isoformat(), prev_end.isoformat(), prev_results)

    # 증감 계산
    for row in report['rows']:
        prev_row = next((r for r in prev_report['rows'] if r['channel'] == row['channel']), None)
        if prev_row:
            row['prev_produced'] = prev_row['produced']
            row['prev_published'] = prev_row['published']
            row['diff_produced'] = row['produced'] - prev_row['produced']
            row['diff_published'] = row['published'] - prev_row['published']
        else:
            row['prev_produced'] = 0
            row['prev_published'] = 0
            row['diff_produced'] = row['produced']
            row['diff_published'] = row['published']
    report['prev_total'] = prev_report['total']

    return report


@report_router.post("/ai-actions")
async def report_ai_actions(request: Request):
    """AI 추천 액션 생성"""
    body = await request.json()
    report_data = body.get('report', {})
    viral_data = body.get('viral', [])

    summary = "채널별 성과:\n"
    for row in report_data.get('rows', []):
        summary += "- %s: 생산 %d, 발행 %d, 달성률 %d%%\n" % (row['channel'], row['produced'], row['published'], row['rate'])
    summary += "\n총 생산: %d, 발행: %d, 달성률: %d%%\n" % (
        report_data.get('total', {}).get('produced', 0),
        report_data.get('total', {}).get('published', 0),
        report_data.get('total', {}).get('rate', 0)
    )
    if viral_data:
        summary += "\n카페바이럴 현황:\n"
        for v in viral_data:
            summary += "- %s: 1단계 %s, 완료=%s\n" % (v.get('cafe',''), v.get('date1',''), v.get('done3', False))

    system = "당신은 마케팅 성과 분석가입니다."
    user = """아래 주간 성과 데이터를 분석해서 다음 주 액션 3~5개를 제안해줘.

%s

제안 기준:
1. 노출률/달성률 낮은 채널 → 원인 분석 + 개선 방안
2. 미발행 콘텐츠 → 빠른 발행 독촉
3. 성과 좋은 채널 → 확장 제안
4. 바이럴 예정일 경과 → 발행 알림
5. 전체 생산성 개선 아이디어

구체적이고 실행 가능한 액션으로. 막연한 조언 금지.""" % summary

    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(executor, call_claude, system, user)
    return {'actions': result}


@report_router.post("/docx")
async def report_docx(request: Request):
    """주간 리포트 docx 생성"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import FileResponse

    body = await request.json()
    report = body.get('report', {})
    actions = body.get('actions', '')

    doc = Document()
    doc.add_heading('주간 성과 리포트', level=0)
    doc.add_paragraph('기간: ' + report.get('period', ''))

    # 생산 현황 테이블
    doc.add_heading('생산 현황', level=1)
    rows_data = report.get('rows', [])
    if rows_data:
        table = doc.add_table(rows=len(rows_data)+2, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(['채널','생산','발행','미발행','달성률']):
            hdr[i].text = h
        for i, row in enumerate(rows_data):
            cells = table.rows[i+1].cells
            cells[0].text = row['channel']
            cells[1].text = str(row['produced'])
            cells[2].text = str(row['published'])
            cells[3].text = str(row['unpublished'])
            cells[4].text = '%d%%' % row['rate']
        # 합계
        t = report.get('total', {})
        last = table.rows[-1].cells
        last[0].text = '합계'
        last[1].text = str(t.get('produced', 0))
        last[2].text = str(t.get('published', 0))
        last[3].text = str(t.get('unpublished', 0))
        last[4].text = '%d%%' % t.get('rate', 0)

    # AI 추천 액션
    if actions:
        doc.add_heading('추천 액션', level=1)
        doc.add_paragraph(actions)

    period = report.get('period', '').replace(' ~ ', '_').replace('/', '-')
    fname = '주간리포트_%s.docx' % period
    fpath = os.path.join(OUTPUTS_DIR, fname)
    doc.save(fpath)
    return FileResponse(fpath, filename=fname, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ═══════════════════════════ SCHEDULER ENDPOINTS ═══════════════════════════

@scheduler_router.get("/config")
async def scheduler_config_get():
    return _sched_load()


@scheduler_router.post("/config")
async def scheduler_config_set(request: Request):
    body = await request.json()
    data = _sched_load()
    if 'daily' in body:
        data['daily'] = body['daily']
    if 'weekly' in body:
        data['weekly'] = body['weekly']
    _sched_save(data)
    # APScheduler job 재등록
    _sync_weekly_jobs(data)
    return {'ok': True}


@scheduler_router.get("/notifications")
async def scheduler_notifications():
    global _scheduler_notifications
    notifs, _scheduler_notifications = _scheduler_notifications, []
    return {'notifications': notifs}


@scheduler_router.get("/history")
async def scheduler_history():
    data = _sched_load()
    return {'history': data.get('history', [])[-50:]}


# ═══════════════════════════ WEEKLY SCHEDULER (APScheduler) ═══════════════════════════

async def _fire_daily_task(task_id: str):
    """APScheduler cron job 콜백 — 매일 알림 발행."""
    global _scheduler_notifications
    data = _sched_load()
    task = data.get('daily', {}).get(task_id)
    if not task or not task.get('enabled'):
        return
    now = datetime.now()
    label = task.get('label', task_id)
    count_msg = ''
    if task_id == 'generate_remind':
        try:
            from src.api.batch import batch_keywords
            r = await batch_keywords()
            count_msg = f" ({len(r.get('keywords', []))}개 대기)"
        except Exception:
            pass
    fire_key = f'daily_{task_id}_{now.strftime("%Y-%m-%d")}'
    _scheduler_notifications.append({
        'id': fire_key, 'type': 'remind', 'task': task_id,
        'message': f'{label} 시간입니다{count_msg}',
        'time': now.isoformat(),
    })
    data.setdefault('history', []).append({'task': task_id, 'type': 'remind', 'time': now.isoformat()})
    if len(data['history']) > 100:
        data['history'] = data['history'][-100:]
    _sched_save(data)


async def _fire_weekly_task(task_id: str):
    """APScheduler cron job 콜백 — 주간 작업 알림/실행."""
    global _scheduler_notifications
    data = _sched_load()
    task = data.get('weekly', {}).get(task_id)
    if not task or not task.get('enabled'):
        return
    now = datetime.now()
    label = task.get('label', task_id)
    auto_run = task.get('auto_run', False)
    fire_key = f'weekly_{task_id}_{now.strftime("%Y-%m-%d")}'
    if auto_run:
        _scheduler_notifications.append({
            'id': fire_key, 'type': 'remind', 'task': task_id,
            'message': f'{label} 시간입니다 — 대시보드에서 실행하세요',
            'time': now.isoformat(),
        })
    else:
        _scheduler_notifications.append({
            'id': fire_key, 'type': 'remind', 'task': task_id,
            'message': f'{label} 시간입니다',
            'time': now.isoformat(),
        })
    data.setdefault('history', []).append({
        'task': task_id, 'type': 'auto_run' if auto_run else 'remind',
        'time': now.isoformat(),
    })
    if len(data['history']) > 100:
        data['history'] = data['history'][-100:]
    _sched_save(data)


def _sync_weekly_jobs(data=None):
    """weekly_schedule.json → APScheduler cron job 동기화."""
    from src.services.scheduler_service import scheduler

    if data is None:
        data = _sched_load()

    # 기존 weekly_ / daily_ job 제거
    for job in scheduler.get_jobs():
        if job.id.startswith('daily_') or job.id.startswith('weekly_'):
            job.remove()

    # daily job 등록
    for task_id, task in data.get('daily', {}).items():
        if not task.get('enabled'):
            continue
        t = task.get('time', '09:00')
        hour, minute = int(t.split(':')[0]), int(t.split(':')[1])
        scheduler.add_job(
            _fire_daily_task, 'cron',
            id=f'daily_{task_id}', hour=hour, minute=minute,
            args=[task_id], replace_existing=True,
            misfire_grace_time=300,
        )

    # weekly job 등록
    for task_id, task in data.get('weekly', {}).items():
        if not task.get('enabled'):
            continue
        t = task.get('time', '09:00')
        hour, minute = int(t.split(':')[0]), int(t.split(':')[1])
        day_of_week = task.get('day', 'mon')
        scheduler.add_job(
            _fire_weekly_task, 'cron',
            id=f'weekly_{task_id}', day_of_week=day_of_week, hour=hour, minute=minute,
            args=[task_id], replace_existing=True,
            misfire_grace_time=300,
        )


async def start_weekly_scheduler():
    """앱 startup 이벤트에서 호출 — JSON → APScheduler job 등록."""
    _sync_weekly_jobs()
